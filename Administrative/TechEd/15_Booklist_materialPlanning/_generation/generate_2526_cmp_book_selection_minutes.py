"""Generate 2025-2026 junior CMP textbook selection meeting minutes (DOCX)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "2022-2023科學科初中綜合科學選書會議記錄.docx"
OUT = ROOT / "2025-2026電腦科初中選書會議記錄.docx"


def _add(doc: Document, text: str, *, style: str = "Normal", bullet: bool = False) -> None:
    p = doc.add_paragraph(text, style="List Paragraph" if bullet else style)
    if bullet:
        p.paragraph_format.left_indent = Pt(18)


def build() -> Document:
    doc = Document(str(TEMPLATE))
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    _add(doc, "迦密聖道中學")
    _add(doc, "電腦科")
    _add(doc, "2025 – 2026年度初中電腦科選書會議記錄")
    _add(doc, "日期： 28-5-2026 (星期三)", style="Default")
    _add(doc, "時間： 4:00pm – 4:30pm", style="Default")
    _add(doc, "地點： 電腦室", style="Default")
    _add(doc, "主席： 陳家倫老師", style="Default")
    _add(doc, "文書： 陳卓文老師", style="Default")
    _add(doc, "出席： 陳卓文老師、甘濠銘老師、陳家倫老師", style="Default")
    doc.add_paragraph("")

    _add(
        doc,
        "出版社A：香港大學電機電子工程系／香港大學電子學習發展實驗室 — "
        "明德電腦系列（iClass eBook）",
    )
    _add(doc, "出版社B：卓思出版社有限公司（iClass 平台）")
    doc.add_paragraph("")

    _add(
        doc,
        "會議討論 2026-2027 學年中二及中三電腦科學生用書。中二、中三本年度（2025-2026）"
        "採用校本教材。科技教育課程引入生成式人工智能、Python 與人工智能等內容，"
        "發展迅速，教材送審與批核需時。參考教育局適用書目表，出版社A及出版社B雖為"
        "科技教育相關出版社，但就上述新興課題，兩者目前均未有載列於適用書目表、且能"
        "完全配合本校課程編排之獲批課本。審書小組透過 iClass 平台檢視電子課本樣本，"
        "並就課程需要、章節編排、例子與實習指引、學生自學適切性等作出比較。",
    )
    doc.add_paragraph("")

    _add(doc, "中二級選用《明德電腦——生成式人工智能》（2026版本）原因如下：")
    bullets_s2 = [
        "課程需涵蓋生成式人工智能概述及常見應用、文本生成與提示語技巧（R-I-C-C-O）、"
        "利用 AI 閱讀及整理資訊、影像生成，以及負責任使用 AI（幻覺、交叉核對、私隱及學術誠信）等，"
        "需要一本課本供上課及課後閱讀。",
        "就「生成式人工智能」課題，出版社A及出版社B均未有適用書目表內之完全匹配獲批課本；"
        "《明德電腦——生成式人工智能》（2026版本）內容切合本校編排，為現時可取得之最新版本。",
        "課本章節符合課程需要；章節設有詞彙定義及實驗站供學生自檢；例子及指引清晰，"
        "學生自行閱讀亦能理解及跟從實習。",
    ]
    for b in bullets_s2:
        _add(doc, b, bullet=True)
    doc.add_paragraph("")

    _add(doc, "中三級選用《初中 AI and Python》（2023版本）原因如下：")
    bullets_s3 = [
        "課程需系統化推行「Python 與人工智能」單元，包括 Python 函數庫、文字轉語音／"
        "語音轉文字、電腦視覺、物件追蹤、運用 AI 輔助編程，以及問題解決與除錯等，"
        "需要一本課本供上課、課後閱讀及溫習。",
        "就「Python 與人工智能」課題，出版社A及出版社B均未有適用書目表內之完全匹配初中獲批課本；"
        "《初中 AI and Python》（2023版本）章節及實驗活動與校本教學進度相若。",
        "課本章節符合課程需要；章節編排清晰，配合 Python 編程實習；任務及例子清楚，"
        "學生自行閱讀亦能理解及跟從。",
        "送呈 2026-2027 書單時，出版社欄填寫「香港大學電子學習發展實驗室」，"
        "與本校其他明德電腦／明德資訊及通訊科技教科書之填報方式一致；"
        "課本以 iClass 電子課本形式供學生使用。",
    ]
    for b in bullets_s3:
        _add(doc, b, bullet=True)
    doc.add_paragraph("")

    _add(
        doc,
        "審書成員均聲明與上述出版社及 iClass 平台沒有任何直接或間接關係，"
        "沒有造成利益衝突。",
    )
    doc.add_paragraph("")

    _add(doc, "結論", bullet=True)
    _add(
        doc,
        "議決通過：2026-2027 學年中二級採用《明德電腦——生成式人工智能》（2026版本）"
        "（iClass eBook，香港大學電機電子工程系／香港大學電子學習發展實驗室）；"
        "中三級採用《初中 AI and Python》（2023版本）（iClass eBook，"
        "送呈書單之出版社為香港大學電子學習發展實驗室）。",
        bullet=True,
    )
    doc.add_paragraph("")

    _add(doc, "附錄：審書成員意見摘要", bullet=True)
    appendix = [
        "（陳卓文老師）課本章節符合課程需要；章節中除了有特定詞彙的定義外，"
        "亦提供實驗站供學生自檢所學；課本內的例子，學生自行閱讀亦能理解及跟從實習。",
        "（甘濠銘老師）同意所選課本；內容相對豐富，例子及指引比較清晰。",
        "（陳家倫老師）中三課本章節編排清晰，配合 Python 編程實習；"
        "課本內的任務及例子，學生自行閱讀亦能理解及跟從。",
        "（全體）就適用書目表未有完全匹配之新興課題，同意在審批完成前採用上述 iClass 電子課本，"
        "以配合校本課程及科技教育發展需要。",
    ]
    for line in appendix:
        _add(doc, line, bullet=True)

    return doc


def main() -> None:
    doc = build()
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
