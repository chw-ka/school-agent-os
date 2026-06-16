import argparse
from pathlib import Path

from docx import Document


def extract_docx_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    lines: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    # Tables often contain agenda items.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    # Normalize multi-line cell content.
                    for line in (x.strip() for x in t.splitlines()):
                        if line:
                            lines.append(line)

    return "\n".join(lines).strip() + "\n"


def main() -> int:
    ap = argparse.ArgumentParser(description="Extract plain text from a DOCX file.")
    ap.add_argument("docx", type=Path, help="Path to .docx file")
    ap.add_argument("--out", type=Path, required=True, help="Output text file path")
    args = ap.parse_args()

    text = extract_docx_text(args.docx)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(text, encoding="utf-8")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

