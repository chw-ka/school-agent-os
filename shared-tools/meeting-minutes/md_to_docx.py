"""Convert meeting minutes markdown to DOCX (微軟正黑體 12pt)."""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def set_run_font(run, *, size_pt: float = 12, bold: bool = False) -> None:
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_paragraph(doc: Document, text: str, *, bold: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, bold=bold)


def parse_table_lines(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(set(c) <= {"-", ":"} for c in cells):
            continue
        rows.append(cells)
    return rows


def md_to_docx(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")
    style.font.size = Pt(12)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            add_paragraph(doc, line[2:].strip(), bold=True, center=True)
            i += 1
            continue

        if line.startswith("## "):
            add_paragraph(doc, line[3:].strip(), bold=True)
            i += 1
            continue

        if line.startswith("### "):
            add_paragraph(doc, line[4:].strip(), bold=True)
            i += 1
            continue

        if line.startswith("#### "):
            add_paragraph(doc, line[5:].strip(), bold=True)
            i += 1
            continue

        if line.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table_lines(table_lines)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        run = p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text))
                        set_run_font(run, bold=(ri == 0))
            continue

        if line.startswith("- "):
            content = line[2:].strip()
            content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(content)
            set_run_font(run)
            i += 1
            continue

        if re.match(r"^\d+\.\s", line):
            content = re.sub(r"^\d+\.\s", "", line)
            content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(content)
            set_run_font(run)
            i += 1
            continue

        if line.startswith("> "):
            add_paragraph(doc, line[2:].strip())
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        plain = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        add_paragraph(doc, plain)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    ap = argparse.ArgumentParser(description="Convert meeting minutes .md to .docx")
    ap.add_argument("md", type=Path)
    ap.add_argument("--out", type=Path, required=True)
    args = ap.parse_args()
    md_to_docx(args.md.expanduser().resolve(), args.out.expanduser().resolve())
    print(f"Wrote: {args.out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
