"""Extract marking spec from teacher-finalized S3 CMP written exam DOCX (not generated draft)."""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document

_REPO = Path(__file__).resolve().parents[2]
_QC = _REPO / "shared-tools" / "question-quality-check"
if str(_QC) not in sys.path:
    sys.path.insert(0, str(_QC))

from exam_spec import build_spec, make_item  # noqa: E402

_SUB = ("(a)", "(b)", "(c)", "(d)", "(e)")


def _parse_rubric_abcde(lines: list[str]) -> list[list[str]]:
    blocks: list[list[str]] = []
    current: list[str] = []
    for line in lines:
        if re.match(r"^\d+\.", line.strip()):
            if current:
                blocks.append(current)
            current = []
            m = re.search(r"\(a\)\s+(\S+)", line)
            if m:
                current.append(m.group(1))
            continue
        m = re.match(r"\(([a-e])\)\s+(\S+)", line.strip())
        if m:
            current.append(m.group(2))
    if current:
        blocks.append(current)
    return blocks


def _split_question_and_marking(paras: list[str]) -> tuple[list[str], list[str]]:
    """Question paper ends before second 評分準則 / 答題紙 duplicate."""
    join_idx = None
    for i, t in enumerate(paras):
        if "評分準則" in t and i > 20:
            join_idx = i
            break
    if join_idx is None:
        for i, t in enumerate(paras):
            if t.startswith("答題紙") and i > 50:
                join_idx = i
                break
    if join_idx is None:
        raise ValueError("Cannot find marking scheme section in DOCX")
    return paras[:join_idx], paras[join_idx:]


def extract_answers(marking_paras: list[str]) -> dict[str, Any]:
    sections: dict[str, list[str]] = {}
    cur: str | None = None
    buf: list[str] = []
    for line in marking_paras:
        if "甲部" in line and "多項" in line:
            if cur and buf:
                sections[cur] = buf
            cur, buf = "mcq", []
            continue
        if "乙部" in line and "配對" in line:
            if cur and buf:
                sections[cur] = buf
            cur, buf = "match", []
            continue
        if "丙部" in line and "是非" in line:
            if cur and buf:
                sections[cur] = buf
            cur, buf = "tf", []
            continue
        if "丁部" in line and "填充" in line:
            if cur and buf:
                sections[cur] = buf
            cur, buf = "fill", []
            continue
        if "戊部" in line and ("短答" in line or "問答" in line):
            if cur and buf:
                sections[cur] = buf
            cur, buf = "sa", []
            continue
        if cur:
            if line.strip():
                buf.append(line.strip())
    if cur and buf:
        sections[cur] = buf

    mcq = None
    for line in sections.get("mcq", []):
        s = re.sub(r"\s+", "", line.strip())
        if len(s) == 20 and all(c in "ABCD" for c in s):
            mcq = s
            break

    matching = _parse_rubric_abcde(sections.get("match", []))
    tf_blocks = _parse_rubric_abcde(sections.get("tf", []))
    tf = "".join(tf_blocks[0]) if tf_blocks else ""
    fill = _parse_rubric_abcde(sections.get("fill", []))

    return {
        "mcq_answers": mcq,
        "matching_answers": ["".join(b) for b in matching],
        "tf_answers": tf,
        "fill_answers": fill,
    }


def _word_banks_from_tables(doc: Document) -> list[list[str]]:
    banks: list[list[str]] = []
    for table in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in table.rows]
        flat = " ".join(" ".join(r) for r in rows)
        if rows and len(rows[0]) >= 5 and all(
            w in flat
            for w in ("gTTS", "playsound", "opencv-python", "quiz_data.json", "tkinter")
        ):
            if len(rows) == 1 and len(rows[0]) == 5:
                banks.append(rows[0])
            elif len(rows) == 1 and "\t" in rows[0][0]:
                banks.append([w for w in re.split(r"\s+", rows[0][0]) if w])
    # tables 3 and 4 in typical layout
    if len(doc.tables) >= 5:
        t3 = [c.text.strip() for c in doc.tables[3].rows[0].cells]
        t4raw = doc.tables[4].rows[0].cells[0].text.strip()
        t4 = [w for w in re.split(r"\s+", t4raw) if w]
        if t3 and t4:
            return [t3, t4]
    return banks


def _sa_rubric_from_table(doc: Document) -> list[dict[str, Any]]:
    if len(doc.tables) < 8:
        return []
    rows = doc.tables[7].rows
    out: list[dict[str, Any]] = []
    for row in rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 3:
            out.append({"part": cells[0], "rubric": cells[1], "marks": int(cells[2])})
    return out


def build_items_from_docx(doc: Document, answers: dict[str, Any]) -> list[dict]:
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    qparas, _ = _split_question_and_marking(paras)
    items: list[dict] = []

    # MCQ stems are before 乙部 — keep placeholders; stems not needed for B–E marking
    if answers.get("mcq_answers"):
        for i, letter in enumerate(answers["mcq_answers"], start=1):
            items.append(
                make_item(f"mcq-{i:02d}", "mcq", f"MCQ {i}", marks=1, answer=letter)
            )

    match_keys = answers.get("matching_answers") or []
    for bi, key in enumerate(match_keys, start=1):
        items.append(
            make_item(
                f"b-match-{bi}",
                "section_b",
                f"乙部配對題 {bi}",
                marks=5,
                answer=key,
            )
        )

    tf = answers.get("tf_answers") or ""
    for i, ch in enumerate(tf, start=1):
        items.append(
            make_item(f"c-tf-{i:02d}", "section_c", f"丙部是非 {i}", marks=1, answer=ch)
        )

    fill_blocks = answers.get("fill_answers") or []
    for bi, block in enumerate(fill_blocks, start=1):
        for j, word in enumerate(block, start=1):
            items.append(
                make_item(
                    f"d-fill-{bi}-{j:02d}",
                    "section_d",
                    f"丁部填充 {bi}({ _SUB[j-1] })",
                    marks=1,
                    answer=word,
                )
            )

    for rub in _sa_rubric_from_table(doc):
        part = rub["part"].strip("()")
        items.append(
            make_item(
                f"e-sa-{part}",
                "section_e",
                rub["rubric"],
                marks=rub["marks"],
            )
        )

    return items


def extract_spec(docx_path: Path, *, mcq_answers: str | None = None) -> dict[str, Any]:
    doc = Document(docx_path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    _, marking = _split_question_and_marking(paras)
    answers = extract_answers(marking)
    if mcq_answers:
        answers["mcq_answers"] = re.sub(r"\s+", "", mcq_answers)

    meta: dict[str, Any] = {
        "title": "25-26 S3 CMP Term 2 Written Exam",
        "subject": "S3 CMP",
        "level": "中三級",
        "total_marks": 50,
        "academic_year": "2025-2026",
        "source_docx": str(docx_path).replace("\\", "/"),
        "mcq_answers": answers.get("mcq_answers"),
        "mcq_answers_pending": answers.get("mcq_answers") is None,
        "matching_answers": answers.get("matching_answers"),
        "tf_answers": answers.get("tf_answers"),
        "fill_answers": answers.get("fill_answers"),
        "fill_word_banks": _word_banks_from_tables(doc),
        "sa_rubric": _sa_rubric_from_table(doc),
        "footer": {
            "academic_year": "2025-2026",
            "level": "中三級",
            "term_exam": "下學期考試",
            "subject": "電腦認知",
        },
    }
    items = build_items_from_docx(doc, answers)
    return build_spec(meta, items)


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("docx", type=Path, help="Teacher-finalized exam DOCX (assessments/WrittenExam/)")
    ap.add_argument("-o", "--output", type=Path, required=True)
    ap.add_argument("--mcq-answers", default=None, help="Override MCQ key (20 letters)")
    args = ap.parse_args(argv)

    spec = extract_spec(args.docx.expanduser().resolve(), mcq_answers=args.mcq_answers)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {args.output} ({len(spec['items'])} items)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
