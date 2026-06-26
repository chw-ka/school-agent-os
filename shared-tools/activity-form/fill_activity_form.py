"""Fill 校內/外活動通告申請表 from a JSON spec (copy official template, then populate)."""

from __future__ import annotations

import argparse
import json
import re
import shutil
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document

OPTIONAL_ROWS = frozenset({"escort_teachers"})

DEFAULT_TEMPLATE = Path(
    r"S:\01_School Administration\11_Useful Forms\校內外活動申請表_20240819.docx"
)

# Label in column 0 (after normalize) -> spec key
FIELD_ROWS: dict[str, str] = {
    "校內負責單位": "department",
    "負責老師": "teacher_in_charge",
    "帶隊老師/教練": "escort_teachers",
    "須查看電子通告的老師": "notice_teachers",
    "活動名稱": "activity_name",
    "主辦/合辦機構": "organizer",
    "活動目的": "purpose",
    "活動日期": "activity_dates",
    "活動地點": "venue",
    "集合地點": "meeting_point",
    "解散地點": "dismissal_point",
    "集合時間": "meeting_times",
    "解散時間": "dismissal_times",
    "交通工具": "transport",
    "參加人數": "participant_count",
    "費用": "fee",
    "服飾": "dress",
    "回覆": "reply",
    "日期": "notice_dates",
    "其他事項/備註": "notes",
}

TRANSPORT_TEXT = {
    "self": "自行前往",
    "coach": "旅遊車",
}

FEE_TEXT = {
    "free": "全免",
    "subsidized": "費用由校方津貼",
}

DRESS_TEXT = {
    "summer_uniform": "整齊夏季校服",
    "winter_uniform": "整齊冬季校服",
    "sports": "整齊運動服套裝",
    "casual": "端莊便服(參照便服指引)",
    "team": "球隊隊衣",
}


def _normalize_label(text: str) -> str:
    t = (text or "").strip()
    t = t.lstrip("*")
    t = re.sub(r"\s+", "", t)
    return t


def _set_cell_text(cell: Any, text: str) -> None:
    if cell.paragraphs:
        cell.paragraphs[0].text = text
        for p in cell.paragraphs[1:]:
            p.text = ""
    else:
        cell.text = text


def _find_row(table: Any, label_key: str) -> int | None:
    target = _normalize_label(label_key)
    for i, row in enumerate(table.rows):
        if _normalize_label(row.cells[0].text) == target:
            return i
    return None


def _format_bilingual(prefix_zh: str, zh: str, en: str | None) -> str:
    zh = (zh or "").strip()
    en = (en or "").strip()
    if zh and en:
        return f"({prefix_zh}) {zh}\n(英) {en}"
    if zh:
        return f"({prefix_zh}) {zh}"
    if en:
        return f"({prefix_zh})\n(英) {en}"
    return f"({prefix_zh})"


def _format_transport(spec: dict[str, Any]) -> str:
    if spec.get("transport_text"):
        return str(spec["transport_text"])
    mode = spec.get("transport", "self")
    if mode == "other":
        other = (spec.get("transport_other") or "").strip()
        return f"其他﹕{other}" if other else "其他﹕____________"
    return TRANSPORT_TEXT.get(mode, TRANSPORT_TEXT["self"])


def _format_fee(spec: dict[str, Any]) -> str:
    if spec.get("fee_text"):
        return str(spec["fee_text"])
    mode = spec.get("fee", "free")
    if mode == "per_student":
        amount = (spec.get("fee_amount") or "").strip()
        return f"每位${amount}" if amount else "每位$________"
    return FEE_TEXT.get(mode, FEE_TEXT["free"])


def _format_dress(spec: dict[str, Any]) -> str:
    if spec.get("dress_text"):
        return str(spec["dress_text"])
    mode = spec.get("dress", "winter_uniform")
    if mode == "other":
        other = (spec.get("dress_other") or "").strip()
        return other or "其他﹕_______________"
    return DRESS_TEXT.get(mode, DRESS_TEXT["winter_uniform"])


def _format_reply(spec: dict[str, Any]) -> str:
    if spec.get("reply_text"):
        return str(spec["reply_text"])
    mode = spec.get("reply", "mandatory")
    base = " 必須參與           □ 可選擇參與\n 家長及學生聯絡電話(外出活動適用)"
    if mode == "optional":
        return base.replace("□ 可選擇參與", "☑ 可選擇參與")
    return base.replace(" 必須參與", "☑ 必須參與")


def _format_notice_dates(spec: dict[str, Any]) -> str:
    if spec.get("notice_dates_text"):
        return str(spec["notice_dates_text"])
    issue = (spec.get("notice_issue") or "").strip()
    slip = (spec.get("reply_slip") or "").strip()
    if issue or slip:
        return f"出通告﹕{issue}     交回條﹕{slip}"
    return "出通告﹕     月      日     交回條﹕    月      日"


def _format_participant_count(count: Any) -> str:
    if count is None or count == "":
        return "                                   (學生名單見P.2)"
    return f"{count} (學生名單見P.2)"


def _value_for_field(key: str, spec: dict[str, Any]) -> str | None:
    if key == "activity_name":
        return _format_bilingual(
            "中",
            spec.get("activity_name_zh", ""),
            spec.get("activity_name_en"),
        )
    if key == "organizer":
        return _format_bilingual(
            "中",
            spec.get("organizer_zh", ""),
            spec.get("organizer_en"),
        )
    if key == "transport":
        return _format_transport(spec)
    if key == "fee":
        return _format_fee(spec)
    if key == "dress":
        return _format_dress(spec)
    if key == "reply":
        return _format_reply(spec)
    if key == "notice_dates":
        return _format_notice_dates(spec)
    if key == "participant_count":
        return _format_participant_count(spec.get("participant_count"))
    if key in spec and spec[key] is not None:
        return str(spec[key])
    return None


def _has_field_value(key: str, spec: dict[str, Any]) -> bool:
    if key == "activity_name":
        return bool(spec.get("activity_name_zh") or spec.get("activity_name_en"))
    if key == "organizer":
        return bool(spec.get("organizer_zh") or spec.get("organizer_en"))
    return spec.get(key) is not None and spec.get(key) != ""


def _fill_main_table(table: Any, spec: dict[str, Any]) -> list[str]:
    missing: list[str] = []
    for label, key in FIELD_ROWS.items():
        row_idx = _find_row(table, label)
        if row_idx is None:
            if _has_field_value(key, spec) and key not in OPTIONAL_ROWS:
                missing.append(f"row not found: {label}")
            continue
        value = _value_for_field(key, spec)
        if value is None:
            continue
        _set_cell_text(table.rows[row_idx].cells[1], value)
    return missing


def _fill_students(table: Any, students: list[dict[str, Any]]) -> None:
    for idx, student in enumerate(students[:40]):
        row_idx = (idx % 20) + 1
        if row_idx >= len(table.rows):
            break
        if idx < 20:
            cols = (1, 2, 3)
        else:
            cols = (5, 6, 7)
        row = table.rows[row_idx]
        _set_cell_text(row.cells[cols[0]], str(student.get("class", "") or ""))
        _set_cell_text(row.cells[cols[1]], str(student.get("number", "") or ""))
        _set_cell_text(row.cells[cols[2]], str(student.get("name", "") or ""))


def _fill_paragraphs(doc: Document, spec: dict[str, Any]) -> None:
    if spec.get("submission_date") and len(doc.paragraphs) > 1:
        sub = spec["submission_date"]
        doc.paragraphs[1].text = (
            f"交表日期：{sub}\t                通告編號﹕                  (校務處專用)"
        )
    if spec.get("form_updated_date") and len(doc.paragraphs) > 3:
        upd = spec["form_updated_date"]
        doc.paragraphs[3].text = (
            f"更新日期﹕{upd}                                         → P.T.O (學生名單)"
        )


def fill_activity_form(
    *,
    template: Path,
    spec: dict[str, Any],
    output: Path,
) -> Path:
    if not template.is_file():
        raise FileNotFoundError(f"Template not found: {template}")
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, output)

    doc = Document(str(output))
    if not doc.tables:
        raise RuntimeError("Template has no tables")

    _fill_paragraphs(doc, spec)
    missing = _fill_main_table(doc.tables[0], spec)
    if len(doc.tables) > 2:
        _fill_students(doc.tables[2], spec.get("students") or [])

    doc.save(str(output))
    if missing:
        print("Warnings:", "; ".join(missing))
    return output


def _default_output_name(spec: dict[str, Any]) -> str:
    if spec.get("output_basename"):
        base = spec["output_basename"]
        return base if base.endswith(".docx") else f"{base}.docx"
    today = date.today()
    return f"校內外活動申請表_{today.year}{today.month:02d}.docx"


def main() -> int:
    ap = argparse.ArgumentParser(
        description="Copy and fill 校內/外活動通告申請表 from JSON spec."
    )
    ap.add_argument("spec", type=Path, help="JSON spec file")
    ap.add_argument(
        "--template",
        type=Path,
        default=DEFAULT_TEMPLATE,
        help="Official blank form (.docx)",
    )
    ap.add_argument(
        "--out",
        type=Path,
        help="Output .docx path (default: activity_folder / 校內外活動申請表_YYYYMM.docx)",
    )
    ap.add_argument(
        "--activity-folder",
        type=Path,
        help="Activity folder on S: (used when --out omitted)",
    )
    args = ap.parse_args()

    spec = json.loads(args.spec.read_text(encoding="utf-8"))
    if args.out:
        out = args.out
    elif args.activity_folder:
        out = args.activity_folder / _default_output_name(spec)
    elif spec.get("activity_folder"):
        out = Path(spec["activity_folder"]) / _default_output_name(spec)
    else:
        raise SystemExit("Provide --out or --activity-folder (or activity_folder in spec)")

    result = fill_activity_form(template=args.template, spec=spec, output=out)
    print(result)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
