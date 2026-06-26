"""Extract a JSON spec from a filled 校內/外活動申請表 (for cloning / review)."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document

from fill_activity_form import FIELD_ROWS, _normalize_label


def _cell(table, label: str) -> str:
    for row in table.rows:
        if _normalize_label(row.cells[0].text) == _normalize_label(label):
            return (row.cells[1].text or "").strip()
    return ""


def _parse_bilingual(text: str) -> tuple[str, str]:
    zh = ""
    en = ""
    m = re.search(r"\(中\)\s*(.+?)(?:\n\(英\)\s*(.+))?$", text, re.DOTALL)
    if m:
        zh = (m.group(1) or "").strip()
        en = (m.group(2) or "").strip()
    return zh, en


def _extract_students(table) -> list[dict[str, str]]:
    students: list[dict[str, str]] = []
    for row in table.rows[1:]:
        cells = row.cells
        for cols in ((1, 2, 3), (5, 6, 7)):
            if cols[2] >= len(cells):
                continue
            rec = {
                "class": (cells[cols[0]].text or "").strip(),
                "number": (cells[cols[1]].text or "").strip(),
                "name": (cells[cols[2]].text or "").strip(),
            }
            if any(rec.values()):
                students.append(rec)
    return students


def extract_spec(docx_path: Path) -> dict:
    doc = Document(str(docx_path))
    main = doc.tables[0]
    spec: dict = {"source_docx": str(docx_path)}

    for label, key in FIELD_ROWS.items():
        val = _cell(main, label)
        if not val:
            continue
        if key == "activity_name":
            zh, en = _parse_bilingual(val)
            spec["activity_name_zh"] = zh
            if en:
                spec["activity_name_en"] = en
        elif key == "organizer":
            zh, en = _parse_bilingual(val)
            spec["organizer_zh"] = zh
            if en:
                spec["organizer_en"] = en
        elif key == "participant_count":
            m = re.match(r"(\d+)", val)
            if m:
                spec["participant_count"] = int(m.group(1))
        elif key in ("transport", "fee", "dress", "reply", "notice_dates"):
            spec[f"{key}_text"] = val
        else:
            spec[key] = val

    if len(doc.tables) > 2:
        spec["students"] = _extract_students(doc.tables[2])

    if len(doc.paragraphs) > 3:
        m = re.search(r"更新日期﹕(\S+)", doc.paragraphs[3].text)
        if m:
            spec["form_updated_date"] = m.group(1)

    return spec


def main() -> int:
    ap = argparse.ArgumentParser(description="Extract activity form spec from filled DOCX.")
    ap.add_argument("docx", type=Path)
    ap.add_argument("--out", type=Path, required=True)
    args = ap.parse_args()

    spec = extract_spec(args.docx)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(spec, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(args.out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
