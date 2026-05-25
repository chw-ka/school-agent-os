"""Table visual style for F5 ICT Exam02 — matched to 24_25 template DOCX."""

from __future__ import annotations

from copy import deepcopy
from typing import Callable

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table

from docx_inplace import BODY_FONT, CODE_FONT
from f5_ict_template_table import template_table

# Prototype table indices in 24_25_S5_ICT_Exam02.docx
PROTOTYPE_INDEX_BY_KIND: dict[str, int] = {
    "spreadsheet": 3,
    "data": 11,
    "data_code": 11,
    "schema": 18,
    "trace": 6,
    "stack": 31,
    "compare": 4,
    "single": 5,
}

SLOT_TABLE_KIND: dict[str, str] = {
    "b-01": "spreadsheet",
    "b-02": "spreadsheet",
    "b-03": "single",
    "b-04": "trace",
    "b-05": "data_code",
    "b-06": "compare",
    "c-03": "data",
    "c-05": "schema",
    "c-08": "stack",
    "mcq-06": "spreadsheet",
    "mcq-13": "spreadsheet",
    "mcq-15": "data",
}

# From 24_25_S5_ICT_Exam02.docx body tables (tblPr / tcPr patterns).
TBL_STYLE_ID = "af0"
SHADE_HEADER = "D9D9D9"  # spreadsheet col/row headers
SHADE_TRACE = "BFBFBF"  # array trace column headers
TBL_INDENT_DEFAULT = 562  # dxa (~1 cm) — most 乙丙 tables
TBL_INDENT_SCHEMA = 567
TBL_INDENT_STACK = 1004
TBL_WIDTH_SPREADSHEET = 9904  # dxa — table 3


def _tbl_pr(table: Table):
    tbl = table._tbl
    pr = tbl.tblPr
    if pr is None:
        pr = OxmlElement("w:tblPr")
        tbl.insert(0, pr)
    return pr


def _set_tbl_style(pr, style_id: str) -> None:
    el = pr.find(qn("w:tblStyle"))
    if el is None:
        el = OxmlElement("w:tblStyle")
        pr.insert(0, el)
    el.set(qn("w:val"), style_id)


def _set_tbl_indent(pr, dxa: int) -> None:
    el = pr.find(qn("w:tblInd"))
    if el is None:
        el = OxmlElement("w:tblInd")
        pr.append(el)
    el.set(qn("w:w"), str(dxa))
    el.set(qn("w:type"), "dxa")


def _set_tbl_width(pr, dxa: int | None) -> None:
    el = pr.find(qn("w:tblW"))
    if el is None:
        el = OxmlElement("w:tblW")
        pr.append(el)
    if dxa is None:
        el.set(qn("w:type"), "auto")
        el.set(qn("w:w"), "0")
    else:
        el.set(qn("w:type"), "dxa")
        el.set(qn("w:w"), str(dxa))


def _set_tbl_look(pr) -> None:
    """Match template tblLook val=04A0 (first row / first column banding)."""
    el = pr.find(qn("w:tblLook"))
    if el is None:
        el = OxmlElement("w:tblLook")
        pr.append(el)
    el.set(qn("w:val"), "04A0")
    el.set(qn("w:firstRow"), "1")
    el.set(qn("w:lastRow"), "0")
    el.set(qn("w:firstColumn"), "1")
    el.set(qn("w:lastColumn"), "0")
    el.set(qn("w:noHBand"), "0")
    el.set(qn("w:noVBand"), "1")


def _cell_tc_pr(cell):
    return cell._tc.get_or_add_tcPr()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = _cell_tc_pr(cell)
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_valign_center(cell) -> None:
    tc_pr = _cell_tc_pr(cell)
    va = tc_pr.find(qn("w:vAlign"))
    if va is None:
        va = OxmlElement("w:vAlign")
        tc_pr.append(va)
    va.set(qn("w:val"), "center")


def set_cell_paragraph_align(cell, align: WD_ALIGN_PARAGRAPH) -> None:
    for p in cell.paragraphs:
        p.paragraph_format.alignment = align


def set_cell_font(cell, *, name: str | None = None, bold: bool | None = None) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            if name:
                r.font.name = name
            if bold is not None:
                r.bold = bold
        if not p.runs and p.text:
            r = p.add_run(p.text)
            p.text = ""
            if name:
                r.font.name = name
            if bold is not None:
                r.bold = bold


def _style_all_cells(
    table: Table,
    *,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    font: str | None = BODY_FONT,
    valign: bool = True,
) -> None:
    for row in table.rows:
        for cell in row.cells:
            if valign:
                set_cell_valign_center(cell)
            set_cell_paragraph_align(cell, align)
            if font:
                set_cell_font(cell, name=font)


def apply_table_base(
    table: Table,
    *,
    indent_dxa: int = TBL_INDENT_DEFAULT,
    width_dxa: int | None = None,
    style_id: str = TBL_STYLE_ID,
) -> None:
    pr = _tbl_pr(table)
    _set_tbl_style(pr, style_id)
    _set_tbl_indent(pr, indent_dxa)
    _set_tbl_width(pr, width_dxa)
    _set_tbl_look(pr)


def apply_spreadsheet_style(table: Table) -> None:
    """B1-style grid: grey header row + grey first column (table 3 in 24_25)."""
    apply_table_base(table, indent_dxa=TBL_INDENT_DEFAULT, width_dxa=TBL_WIDTH_SPREADSHEET)
    rows, cols = len(table.rows), len(table.columns)
    for r in range(rows):
        for c in range(cols):
            cell = table.cell(r, c)
            set_cell_valign_center(cell)
            set_cell_paragraph_align(cell, WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_font(cell, name=BODY_FONT)
            if r == 0 or c == 0:
                set_cell_shading(cell, SHADE_HEADER)
    if rows > 1 and cols > 1:
        set_cell_shading(table.cell(0, 0), SHADE_HEADER)


def apply_data_table_style(table: Table, *, code_font: bool = False) -> None:
    """ORDER / PID sample tables — centered, indented (table 11, 15)."""
    apply_table_base(table)
    font = CODE_FONT if code_font else BODY_FONT
    _style_all_cells(table, align=WD_ALIGN_PARAGRAPH.CENTER, font=font)


def apply_schema_table_style(table: Table) -> None:
    """FACILITY / RESERVE schema — header row centered bold; body left (table 18–19)."""
    apply_table_base(table, indent_dxa=TBL_INDENT_SCHEMA)
    rows = len(table.rows)
    for r in range(rows):
        for cell in table.rows[r].cells:
            set_cell_valign_center(cell)
            if r == 0:
                set_cell_paragraph_align(cell, WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_font(cell, name=BODY_FONT, bold=True)
            else:
                set_cell_paragraph_align(cell, WD_ALIGN_PARAGRAPH.LEFT)
                set_cell_font(cell, name=BODY_FONT)


def apply_trace_header_style(table: Table) -> None:
    """Array trace — grey header row only (BFBFBF, table 6)."""
    apply_table_base(table)
    cols = len(table.columns)
    for c in range(cols):
        cell = table.cell(0, c)
        set_cell_shading(cell, SHADE_TRACE)
        set_cell_valign_center(cell)
        set_cell_paragraph_align(cell, WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_font(cell, name=BODY_FONT)
    for r in range(1, len(table.rows)):
        for c in range(cols):
            cell = table.cell(r, c)
            set_cell_valign_center(cell)
            set_cell_paragraph_align(cell, WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_font(cell, name=BODY_FONT)


def apply_stack_trace_style(table: Table) -> None:
    """Stack operation trace — first row/col headers centered (table 31)."""
    apply_table_base(table, indent_dxa=TBL_INDENT_STACK)
    rows, cols = len(table.rows), len(table.columns)
    for r in range(rows):
        for c in range(cols):
            cell = table.cell(r, c)
            set_cell_valign_center(cell)
            set_cell_paragraph_align(cell, WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_font(cell, name=BODY_FONT)
            if r == 0 or c == 0:
                set_cell_font(cell, name=BODY_FONT, bold=(r == 0))


def apply_single_box_style(table: Table) -> None:
    """One-cell spec box (e.g. media dimensions)."""
    apply_table_base(table)
    cell = table.cell(0, 0)
    set_cell_valign_center(cell)
    set_cell_paragraph_align(cell, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_font(cell, name=BODY_FONT)


def apply_compare_table_style(table: Table) -> None:
    """Hardware comparison — header row centered (table 4 style 11)."""
    apply_table_base(table, style_id="11")
    rows = len(table.rows)
    for r in range(rows):
        for cell in table.rows[r].cells:
            set_cell_valign_center(cell)
            align = WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_paragraph_align(cell, align)
            set_cell_font(cell, name=BODY_FONT)


TABLE_STYLE_BY_KIND: dict[str, Callable[[Table], None]] = {
    "spreadsheet": apply_spreadsheet_style,
    "data": apply_data_table_style,
    "data_code": lambda t: apply_data_table_style(t, code_font=True),
    "schema": apply_schema_table_style,
    "trace": apply_trace_header_style,
    "stack": apply_stack_trace_style,
    "single": apply_single_box_style,
    "compare": apply_compare_table_style,
}

SLOT_TABLE_KIND: dict[str, str] = {
    "b-01": "spreadsheet",
    "b-02": "spreadsheet",
    "b-03": "single",
    "b-04": "trace",
    "b-05": "data_code",
    "b-06": "compare",
    "c-03": "data",
    "c-05": "schema",
    "c-08": "stack",
    "mcq-06": "spreadsheet",
    "mcq-13": "spreadsheet",
    "mcq-15": "data",
}


def apply_style_from_prototype(table: Table, prototype_index: int) -> None:
    """Copy tblPr (+ table style id) from a template table onto a newly created table."""
    proto = template_table(prototype_index)
    src_pr = proto._tbl.tblPr
    if src_pr is None:
        return
    dst_pr = _tbl_pr(table)
    for child in list(dst_pr):
        dst_pr.remove(child)
    for child in src_pr:
        dst_pr.append(deepcopy(child))


def apply_style_for_slot(table: Table, slot_id: str) -> None:
    kind = SLOT_TABLE_KIND.get(slot_id, "data")
    fn = TABLE_STYLE_BY_KIND.get(kind, apply_data_table_style)
    fn(table)
