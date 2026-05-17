from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from docx_inplace import EnCoverPatch, apply_cmp_cover_en


@dataclass(frozen=True)
class Meta:
    school: str = "Carmel Holy Word Secondary School"
    year: str = "2025 – 2026"
    term: str = "Term 2 Examination"
    subject: str = "Secondary 3 Computer Literacy"
    paper: str = "Written Paper"
    date: str = "__________"
    time: str = "__________"
    time_limit: str = "30 minutes"
    pages: str = "7"
    total_mark: str = "40"


def _replace_paragraph_range_exact(doc: Document, start: int, end_excl: int, lines: list[str]) -> None:
    """Replace text of existing paragraphs only (no insertion)."""
    span = end_excl - start
    if len(lines) > span:
        raise ValueError(f"Replacement has {len(lines)} lines but only {span} paragraphs available.")
    for i in range(span):
        doc.paragraphs[start + i].text = lines[i] if i < len(lines) else ""


def _find_idx(doc: Document, contains: str, *, start: int = 0) -> int:
    for i in range(start, len(doc.paragraphs)):
        if contains in doc.paragraphs[i].text:
            return i
    raise ValueError(f'Cannot find paragraph containing "{contains}"')


def _set_cell(table, r: int, c: int, text: str) -> None:
    table.cell(r, c).text = text


def _apply_cover(doc: Document, meta: Meta) -> None:
    """Patch cover lines in-place; keep template spacing and Instructions unchanged."""
    cover_cell = doc.tables[0].cell(0, 0)
    apply_cmp_cover_en(
        cover_cell,
        EnCoverPatch(
            year_term=f"{meta.year} {meta.term}",
            subject=meta.subject,
            paper=meta.paper,
            date_line=f"\tDate\t: {meta.date}",
            time_line=f"\tTime\t: {meta.time}",
            duration_line=f"\tTime Limit\t: {meta.time_limit}",
            pages_line=f"\tNo. of Pages\t: {meta.pages}",
            total_line=f"\tTotal mark\t: {meta.total_mark}",
        ),
    )


def _build_content():
    # Keep it simple (concept + very short code snippets), aligned to the 7 lessons:
    # libraries/pip, gTTS/playsound, google translate client + env var,
    # SpeechRecognition basics, OpenCV face detection concepts, object tracking ROI/tracker,
    # vibe coding + difficulty curve formula, JSON question format.
    section_a: list[str] = [
        "Section A: Multiple Choice Questions (15 marks)",
        "Choose the most suitable answers.",
        "",
        "1.\tWhich command is used to install a third-party library in Python?",
        "A.\tinstall library",
        "B.\tpip install <library>",
        "C.\tpython add <library>",
        "D.\timport pip",
        "",
        "2.\tWhich library converts text into mp3 audio (TTS) in the notes?",
        "A.\tplaysound",
        "B.\tgTTS",
        "C.\tjson",
        "D.\tos",
        "",
        "3.\tWhich line plays an mp3 file in a TTS program?",
        "A.\tplaysound(\"audio.mp3\")",
        "B.\tgTTS(\"audio.mp3\")",
        "C.\tos.environ(\"audio.mp3\")",
        "D.\ttranslate.Client(\"audio.mp3\")",
        "",
        "4.\tWhich file is used as Google credential in the Cloud Translation example?",
        "A.\tconfig.txt",
        "B.\tpassword.json",
        "C.\tkey.mp3",
        "D.\tcv2.xml",
        "",
        "5.\tWhich module is imported as translate_v2 in the notes?",
        "A.\tfrom google.cloud import translate_v2 as translate",
        "B.\timport translate",
        "C.\timport google.cloud.translate_v2",
        "D.\tfrom translate import google",
        "",
        "6.\tIn SpeechRecognition, which object is used to listen to the microphone input?",
        "A.\tsr.AudioFile",
        "B.\tsr.Recognizer",
        "C.\tsr.Microphone",
        "D.\tsr.Translate",
        "",
        "7.\tWhich function converts speech audio into text using Google service?",
        "A.\trecord()",
        "B.\tlisten()",
        "C.\trecognize_google()",
        "D.\ttranslate()",
        "",
        "8.\tIn OpenCV face detection, what is the purpose of converting to gray scale?",
        "A.\tTo make the image colorful",
        "B.\tTo improve detection speed/efficiency",
        "C.\tTo translate the image",
        "D.\tTo play audio",
        "",
        "9.\tWhich file is a Haar Cascade model file?",
        "A.\thaascade.py",
        "B.\thaascade.jpg",
        "C.\thaascade.xml",
        "D.\thaascade.mp4",
        "",
        "10.\tWhat does cv2.VideoCapture(0) usually open?",
        "A.\tA random image",
        "B.\tThe default camera",
        "C.\tA translation client",
        "D.\tAn mp3 file",
        "",
        "11.\tIn object tracking, what does ROI stand for?",
        "A.\tRange Of Integer",
        "B.\tRegion Of Interest",
        "C.\tRun On Internet",
        "D.\tRotate On Image",
        "",
        "12.\tWhich function lets you draw a rectangle on an OpenCV image?",
        "A.\tcv2.rectangle()",
        "B.\tcv2.translate()",
        "C.\tcv2.speech()",
        "D.\tcv2.random()",
        "",
        "13.\tGiven Speed = 5 × (1.2)^(Level−1), when Level = 2, Speed is:",
        "A.\t5",
        "B.\t6",
        "C.\t7.2",
        "D.\t8.64",
        "",
        "14.\tWhich format is used for storing quiz questions in the lesson?",
        "A.\t.mp3",
        "B.\t.json",
        "C.\t.xml",
        "D.\t.docx",
        "",
        "15.\tIn quiz_data.json, which key stores the correct option?",
        "A.\t\"options\"",
        "B.\t\"question\"",
        "C.\t\"answer\"",
        "D.\t\"score\"",
        "",
        "Section B: Short Questions (15 marks)",
        "All answers should be written in the spaces provided in the answer sheet.",
        "",
        "1.\tText-to-Speech (TTS) and Speech-to-Text (STT) are commonly used in AI applications.\t\t(5 marks)",
        "\t(a)\tState ONE purpose of TTS.\t(1 mark)",
        "\t(b)\tState ONE purpose of STT.\t(1 mark)",
        "\t(c)\tName ONE Python library used for TTS in the notes.\t(1 mark)",
        "\t(d)\tName ONE Python library used for STT in the notes.\t(1 mark)",
        "\t(e)\tState ONE reason why a program may need to use try/except for STT.\t(1 mark)",
        "",
        "2.\tCloud Translation requires credentials and correct parameters.\t\t\t\t\t(5 marks)",
        "\t(a)\tWhat is the purpose of the file password.json?\t(1 mark)",
        "\t(b)\tWrite the environment variable name used to point to the credential file.\t(1 mark)",
        "\t(c)\tState the meaning of source_language and target_language.\t(2 marks)",
        "\t(d)\tGive ONE example of a target_language code used in the notes.\t(1 mark)",
        "",
        "3.\tOpenCV Face Detection and Object Tracking are different tasks.\t\t\t\t\t(5 marks)",
        "\t(a)\tState ONE difference between face detection and object tracking.\t(2 marks)",
        "\t(b)\tWhat does ROI stand for?\t(1 mark)",
        "\t(c)\tState ONE reason why converting an image to gray scale is helpful.\t(1 mark)",
        "\t(d)\tWhich parameter in detectMultiScale can be adjusted to reduce false positives?\t(1 mark)",
        "",
        "Section C: Long Question (10 marks)",
        "1.\tA student wants to build a \"Smart Translator\" program that listens to English speech, translates it to Chinese and plays the translated result.",
        "The student will use SpeechRecognition + Cloud Translation + gTTS.",
        "",
        "(a)\tList the THREE main steps in order (from input to output).\t(3 marks)",
        "(b)\tExplain why the program needs an internet connection.\t(2 marks)",
        "(c)\tThe student gets an error when the program cannot understand the speech.",
        "\tState ONE possible cause and ONE way to handle it in the program.\t(3 marks)",
        "(d)\tState ONE privacy concern if the program is used in a classroom.\t(2 marks)",
        "",
        "End of Paper",
    ]

    # Word banks (single-cell tables 1-3) and answer-sheet tables (8-10)
    bank1 = "gTTS    playsound    SpeechRecognition    google-cloud-translate"
    bank2 = "source_language    target_language    password.json    GOOGLE_APPLICATION_CREDENTIALS"
    bank3 = "OpenCV    Haar Cascade    detectMultiScale    ROI    Tracker"

    # Section B suggested answers (tables 8-10): store key points (teacher reference).
    b1 = {
        "(i)": "TTS: convert text to speech/audio",
        "(ii)": "STT: convert speech/audio to text",
        "(iii)": "gTTS",
        "(iv)": "SpeechRecognition",
        "(v)": "e.g. UnknownValueError / RequestError",
    }
    b2 = {
        "(i)": "Credential file for Google API access",
        "(ii)": "GOOGLE_APPLICATION_CREDENTIALS",
        "(iii)": "source: original language / target: output language",
        "(iv)": "zh (or zh-tw)",
        "(v)": "client.translate(...), result['translatedText']",
    }
    b3 = {
        "(i)": "Detection finds objects each frame; tracking follows a chosen object across frames",
        "(ii)": "Region Of Interest",
        "(iii)": "Gray scale improves speed/accuracy",
        "(iv)": "minNeighbors (or scaleFactor)",
        "(v)": "Reduce false positives by increasing minNeighbors",
    }

    # Long question code table (table 4 cell 1,1): provide a flow skeleton (not fill-blanks).
    long_code = (
        "import os\n"
        "import speech_recognition as sr\n"
        "from google.cloud import translate_v2 as translate\n"
        "from gtts import gTTS\n"
        "from playsound import playsound\n\n"
        "os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = 'password.json'\n"
        "client = translate.Client()\n"
        "recognizer = sr.Recognizer()\n\n"
        "with sr.Microphone() as source:\n"
        "    recognizer.adjust_for_ambient_noise(source)\n"
        "    audio = recognizer.listen(source)\n\n"
        "try:\n"
        "    text = recognizer.recognize_google(audio, language='en')\n"
        "    result = client.translate(text, source_language='en', target_language='zh')\n"
        "    out = result['translatedText']\n"
        "    tts = gTTS(out, lang='zh')\n"
        "    tts.save('out.mp3')\n"
        "    playsound('out.mp3')\n"
        "except Exception as e:\n"
        "    print('Error:', e)\n"
    )

    # MCQ key line format like "BCABB  CCDDB  CBADA"
    mcq_key = "BBABACCBCCBABBBC"[:15]  # placeholder; will compute below
    # Actual answers for section A (15)
    # 1B 2B 3A 4B 5A 6C 7C 8B 9C 10B 11B 12A 13B 14B 15C
    mcq_answers = ["B", "B", "A", "B", "A", "C", "C", "B", "C", "B", "B", "A", "B", "B", "C"]
    mcq_key = "".join(mcq_answers[:5]) + "  " + "".join(mcq_answers[5:10]) + "  " + "".join(mcq_answers[10:15])

    # Suggested answer section C
    c_ans = {
        "a_file1": "STT (listen + recognize) → Translate → TTS (save + play)",
        "a_file2": "Internet needed for Google STT/Translate services",
        "b": "Internet is needed because recognize_google() and Cloud Translation call online services.",
        "c_x": "Cause: noisy environment / unclear speech. Handle: try/except + prompt user to repeat.",
        "c_y": "Privacy: voice content may be sensitive; avoid recording/storing without consent.",
        "c_w": "",
        "c_h": "",
        "d_i": "Example: background noise leads to wrong recognition.",
        "d_ii_line": "",
    }

    return {
        "section_a_to_c_paragraphs": section_a,
        "bank1": bank1,
        "bank2": bank2,
        "bank3": bank3,
        "b1": b1,
        "b2": b2,
        "b3": b3,
        "long_code": long_code,
        "mcq_key": mcq_key,
        "c_ans": c_ans,
    }


def generate(*, template: Path, output: Path, meta: Meta) -> None:
    doc = Document(str(template))
    content = _build_content()

    _apply_cover(doc, meta)

    # Replace the whole question paper text block from start (P0) to "End of Paper"
    start = _find_idx(doc, "Section A: Multiple Choice Questions", start=0)
    end = _find_idx(doc, "End of Paper", start=start)
    # include End of Paper line, and blank following few paragraphs until next header (Answer Sheet)
    block_end = end + 1
    # We must not exceed template paragraph span; replace only within existing span.
    _replace_paragraph_range_exact(doc, start, block_end, content["section_a_to_c_paragraphs"])

    # Word banks tables (1-3)
    doc.tables[1].cell(0, 0).text = content["bank1"]
    doc.tables[2].cell(0, 0).text = content["bank2"]
    doc.tables[3].cell(0, 0).text = content["bank3"]

    # Long question code table (table 4 cell (1,1))
    long_table = doc.tables[4]
    _set_cell(long_table, 1, 1, content["long_code"])
    # Update line numbers in table 4 cell (1,0) to match code line count (roughly)
    lines = content["long_code"].splitlines()
    long_table.cell(1, 0).text = "\n".join(str(i) for i in range(1, len(lines) + 1))

    # Suggested Answer area
    suggested_title = _find_idx(doc, "Suggested Answer", start=0)
    # MCQ key is usually 2 lines below (as seen: Section A: @ 1mark, then key line)
    for i in range(suggested_title, min(suggested_title + 10, len(doc.paragraphs))):
        if doc.paragraphs[i].text.strip().startswith("Section A:"):
            # key expected at i+1
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i + 1].text = content["mcq_key"]
            break

    # Section B answer tables at the end are tables 8-10 (same as 24-25 template)
    t8, t9, t10 = doc.tables[8], doc.tables[9], doc.tables[10]
    _set_cell(t8, 0, 0, "1.\t")
    _set_cell(t8, 0, 1, f"(i)  {content['b1']['(i)']}")
    _set_cell(t8, 0, 2, f"(ii)  {content['b1']['(ii)']}")
    _set_cell(t8, 0, 3, f"(iii)  {content['b1']['(iii)']}")
    _set_cell(t8, 1, 1, f"(iv)  {content['b1']['(iv)']}")
    _set_cell(t8, 1, 2, f"(v)  {content['b1']['(v)']}")

    _set_cell(t9, 0, 0, "2.\t")
    _set_cell(t9, 0, 1, f"(i)  {content['b2']['(i)']}")
    _set_cell(t9, 0, 2, f"(ii)  {content['b2']['(ii)']}")
    _set_cell(t9, 0, 3, f"(iii)  {content['b2']['(iii)']}")
    _set_cell(t9, 1, 2, f"(iv)  {content['b2']['(iv)']}")
    _set_cell(t9, 1, 3, f"(v)  {content['b2']['(v)']}")

    _set_cell(t10, 0, 0, "3.\t")
    _set_cell(t10, 0, 1, f"(i)  {content['b3']['(i)']}")
    _set_cell(t10, 0, 2, f"(ii)  {content['b3']['(ii)']}")
    _set_cell(t10, 0, 3, f"(iii)  {content['b3']['(iii)']}")
    _set_cell(t10, 1, 1, f"(iv)  {content['b3']['(iv)']}")
    _set_cell(t10, 1, 2, f"(v)  {content['b3']['(v)']}")
    _set_cell(t10, 1, 3, f"(v)  {content['b3']['(v)']}")

    # Section C suggested answers (paragraphs after "Section C:")
    c_idx = _find_idx(doc, "Section C:", start=suggested_title)
    # Overwrite a known set of lines like in template (1.(a) ... etc)
    # We locate the first "1.\t(a)" after c_idx.
    a_idx = _find_idx(doc, "1.\t(a)", start=c_idx)
    doc.paragraphs[a_idx].text = f"1.\t(a)\tFile 1: {content['c_ans']['a_file1']}"
    doc.paragraphs[a_idx + 1].text = f"\t\tFile 2: {content['c_ans']['a_file2']}"
    doc.paragraphs[a_idx + 2].text = f"\t(b)\t{content['c_ans']['b']}"
    doc.paragraphs[a_idx + 3].text = f"\t(c)\tx: {content['c_ans']['c_x']}"
    doc.paragraphs[a_idx + 4].text = f"\t\ty: {content['c_ans']['c_y']}"
    doc.paragraphs[a_idx + 5].text = f"\t\tw: {content['c_ans']['c_w']}"
    doc.paragraphs[a_idx + 6].text = f"\t\th: {content['c_ans']['c_h']}"
    doc.paragraphs[a_idx + 7].text = f"\t(d)\t(i)\t{content['c_ans']['d_i']}"
    doc.paragraphs[a_idx + 8].text = f"\t(ii)\tLine: {content['c_ans']['d_ii_line']}"

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def _parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Generate 25-26 S3 CMP Term2 written exam from 24-25 template.")
    p.add_argument("--template", required=True, help="Path to 24-25 Term2 WrittenExam.docx (template).")
    p.add_argument("--output", required=True, help="Output .docx path.")
    p.add_argument("--date", default="__________", help="Cover date text.")
    p.add_argument("--time", default="__________", help="Cover time text.")
    return p.parse_args()


def main() -> int:
    args = _parse_args()
    meta = Meta(date=args.date, time=args.time)
    generate(template=Path(args.template), output=Path(args.output), meta=meta)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

