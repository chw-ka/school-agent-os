#!/usr/bin/env python3
"""
Generate 中二 CMP Term 2 試題簿 from 24_25 S2 template.
保留原有段落／表格骨架，只替換內容；封面改 2025–2026 下學期。
"""
from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from docx_inplace import ZhCoverPatch, apply_cmp_cover_zh


@dataclass(frozen=True)
class Meta:
    school: str = "迦密聖道中學"
    year: str = "2025 – 2026"
    term: str = "下學期考試"
    level: str = "中二級 電腦認知"
    paper: str = "試題簿"
    date: str = "__________"
    time: str = "__________"
    time_limit: str = "30 分鐘"
    pages: str = "9 頁"
    total: str = "50"


def _find_idx(doc: Document, contains: str, *, start: int = 0) -> int:
    for i in range(start, len(doc.paragraphs)):
        if contains in doc.paragraphs[i].text:
            return i
    raise ValueError(f'missing: "{contains}"')


def _replace_block_exact(doc: Document, start: int, end_excl: int, lines: list[str]) -> None:
    span = end_excl - start
    if len(lines) > span:
        raise ValueError(f"Need {span} lines, got {len(lines)}")
    for i in range(span):
        doc.paragraphs[start + i].text = lines[i] if i < len(lines) else ""


def _mcq_spans() -> dict[int, tuple[int, int]]:
    """Stem paragraph index -> (start, end_exclusive). From 24_25_S2_CMP_Term02_Exam.docx."""
    return {
        1: (5, 16),
        2: (16, 19),
        3: (19, 22),
        4: (22, 25),
        5: (25, 28),
        6: (28, 30),
        7: (30, 32),
        8: (32, 35),
        9: (35, 41),
        10: (41, 44),
        11: (44, 47),
        12: (47, 54),
        13: (54, 61),
        14: (61, 67),
        15: (67, 73),
        16: (73, 81),
        17: (81, 94),
        18: (94, 103),
        19: (103, 110),
        20: (110, 113),
    }


def _mcq_blocks() -> dict[int, list[str]]:
    """Each list length must equal span end-start from template."""
    return {
        1: [
            "1.\t下列哪些是生成式人工智能較常見的「服務／應用」例子？",
            "",
            "(i)\t輔助寫作（例如潤色文章）",
            "(ii)\t知識問答（例如解答課業問題）",
            "(iii)\t輔助編程（例如生成程式碼建議）",
            "",
            "\tA.\t只有 (i)",
            "B.\t只有 (i) 和 (ii)",
            "C.\t只有 (ii) 和 (iii)",
            "D.\t(i)、(ii) 和 (iii) 皆是",
            "",
        ],
        2: [
            "2.\t根據筆記，下列哪一項最能描述「幻覺（Hallucination）」？",
            "\nA.\tAI 故意拒絕回答\nB.\tAI 可能編造看似可信但不正確的內容\nC.\tAI 無法顯示中文\nD.\tAI 只能在離線時運作",
            "",
        ],
        3: [
            "3.\t確認 AI 資訊是否可信時，筆記建議的做法不包括下列哪一項？",
            "\nA.\t交叉查證（對照可靠來源）\nB.\t要求 AI 提供來源或連結\nC.\t完全相信 AI 的回覆\nD.\t多模型比對",
            "",
        ],
        4: [
            "4.\t第二章提到文本生成的「四大核心應用」，下列哪一項屬於「語意轉換」？",
            "\nA.\t把長文章濃縮成摘要\nB.\t把口語句子改寫成正式書信語氣\nC.\t撰寫故事開頭\nD.\t把重點製成問答遊戲",
            "",
        ],
        5: [
            "5.\t提示語工程 R-I-C-C-O 當中，「限制（Constraint）」最常做的事是？",
            "\nA.\t指定 AI 扮演的身分\nB.\t設定字數、語氣或禁忌\nC.\t選擇輸出檔案格式\nD.\t輸入使用者密碼",
            "",
        ],
        6: [
            "6.\t下列哪一項屬於「輸出格式（Output）」的要求？\n\n"
            "A.\t指明對象是中二學生\n"
            "B.\t要求用對比表格輸出\n"
            "C.\t告知作文的背景資料\n"
            "D.\t請 AI 扮演老師",
            "",
        ],
        7: [
            "7.\t第三章提到 AI 可以協助閱讀 PDF／Word，下列哪一項不是筆記所列的核心功能？\n\n"
            "A.\t數據提取\n"
            "B.\t邏輯節錄\n"
            "C.\t把 GPU 速度提升十倍\n"
            "D.\t把文字轉化為閃卡（Flashcards）",
            "",
        ],
        8: [
            "8.\t在手機應用程式筆記中，若要使用 Teachable Machine 訓練手勢模型，官方網址是哪一個？",
            "\nA.\thttps://appinventor.mit.edu\nB.\thttps://teachablemachine.withgoogle.com\nC.\thttps://www.python.org\nD.\thttps://news.rthk.hk",
            "",
        ],
        9: [
            "9.\t下列哪一句最能描述「圖像識別」？",
            "\nA.\t把文字翻譯成英文",
            "B.\t把影像中的內容對應到事先定義的類別／標籤",
            "C.\t把程式碼自動除錯",
            "D.\t把音訊轉成文字",
            "",
        ],
        10: [
            "10.\t第四章提到「文生圖」，下列哪一項正確？",
            "\nA.\t上載相片並旋轉像素\nB.\t用文字描述生成全新影像\nC.\t只能生成黑白線稿\nD.\t無法指定風格",
            "",
        ],
        11: [
            "11.\t下列哪一項屬於「圖生圖／影像編輯」方向的應用？",
            "\nA.\t依文字憑空生成一張全新照片（無需原圖）\nB.\t上載相片後要求 AI 依指示修改或二次創作\nC.\t把語音轉成文字\nD.\t把問卷結果製成統計圖",
            "",
        ],
        12: [
            "12.\t第一章表格指出「AI 並不像人類一樣真正理解」，其主要運作方式較接近下列哪一項？",
            "",
            "\tA.\t憑感覺創作藝術",
            "\tB.\t透過計算字詞出現的統計規律預測下一個字",
            "\tC.\t透過攝影鏡頭直接「看到」現實世界",
            "\tD.\t先把資料永久記進硬碟才可回答",
            "",
        ],
        13: [
            "13.\t筆記提及可用哪些做法來核對 AI 的回答？（選出最佳的一項）",
            "",
            "\tA.\t只問一次就不要再檢查",
            "\tB.\t交叉查證、要求來源、換模型比對",
            "\tC.\t把所有答案抄進課本就必定正確",
            "\tD.\t把字型改成標楷體就更準確",
            "",
        ],
        14: [
            "14.\tTeachable Machine 訓練流程中，下列哪一項應在「訓練模型」之前完成？",
            "\nA.\t充分測試模型準確性",
            "B.\t建立分類標籤並加入訓練樣本",
            "C.\t取得可分享連結",
            "D.\t立即分享連結給陌生人公開下載",
            "",
        ],
        15: [
            "15.\t下列哪一項是機器學習用於圖像辨識的主要原因？",
            "A.\t電腦只看見像素便可自動知道類別",
            "B.\t影像包含許多變化，需要大量例子來提升準確性",
            "C.\t機器學習可把相片自動印刷",
            "D.\t機器學習會取代鏡頭功能",
            "",
        ],
        16: [
            "16.\t關於生成式 AI「謬誤與真相」，下列敘述哪一項較符合「真相」？",
            "",
            "A.\tAI 一定具有自我意識與主觀情感",
            "B.\tAI 展現的情緒可能只是語言風格的統計模仿",
            "C.\t輸入數據足夠後就不需要人類回饋",
            "D.\tAI 必然會短期內取代所有人類工作",
            "",
            "",
        ],
        17: [
            "17.\t以下哪一項最不符合 Teachable Machine「影像分類『訓練』」的正確做法？",
            "",
            "A.\t先為每個類別準備足夠且有代表性的影像樣本",
            "B.\t樣本越多越好（在同學能力範圍內盡量多元化）",
            "C.\t訓練前完全不收集樣本，直接按 Train Model",
            "D.\t完成訓練後應反覆測試辨識結果",
            "",
            "請參考下列情境（與手機應用程式筆記『Teachable Machine』第一部分一致），回答第18至20題。",
            "",
            "",
            "",
            "",
            "",
        ],
        18: [
            "18.\t同學要在 Teachable Machine 建立『向左移動／向右移動』兩種手勢控制遊戲角色，下列哪個組合最合理？",
            "",
            "",
            "",
            "A.\t建立兩個類別（標籤），分別代表向左／向右的手勢影像",
            "B.\t只需要下載 App Inventor，不必建立標籤",
            "C.\t只能用文字輸入指令完成訓練",
            "D.\t只能用音訊建立模型",
            "",
        ],
        19: [
            "19.\t筆記強調訓練後要『充分測試』模型準確性，下列哪一項是最主要原因？",
            "",
            "A.\t增加電腦耗電量",
            "B.\t確認真實環境下手勢／光線變化下的辨識表現",
            "C.\t令網站版面更美觀",
            "D.\t令瀏覽器自動更新",
            "",
        ],
        20: [
            "20.\t完成訓練後『記下可分享的連結』的主要用途是甚麼？\n\n"
            "A.\t取代瀏覽器搜尋功能\n"
            "B.\t稍後在 App Inventor 擴充／網頁專案載入模型（例如 TMIC）\n"
            "C.\t自動購買雲端硬碟\n"
            "D.\t下載老師的標準答案",
            "",
            "",
        ],
    }


def _section_b_lines() -> list[str]:
    """57 paragraphs: indices 113–169 inclusive (matches template before ~全卷完~)."""
    return [
        "乙部 – 問答題 (30 分)",
        "題組一（生成式人工智能）：老師要求同學運用生成式 AI 閱讀一份通告 PDF，並整理「中二級」的注意事項；同時要以負責任態度核對內容，避免誤信錯誤資訊。",
        "",
        "\t\t任務重點：(i) 清晰指令　(ii) 對象與背景　(iii) 輸出格式（點列／表格）",
        "\t\t提醒：PDF 排版複雜時，AI 整理可能出錯；務必與老師或原件交叉核對。",
        "\t\t與「幻覺」相關：若發現日期／行程可疑，應如何處理？（請在作答時交代）",
        "",
        "\t(a)\t請完成下表，為上述通告整理任務設計一段提示語：請按 R-I-C-C-O 五項要素填寫（可使用短句）。　（6 分）",
        "",
        "",
        "\t(b)\t請依提示語工程的角度，解釋為何「限制（Constraint）」與「輸出格式（Output）」對通告整理特別重要（須舉一例）。　（8 分）",
        "\t\t要素：限制包括截稿時間、對象、語氣；輸出格式例如時間線表格／分段項目。",
        "請分兩段作答：先說明限制可減少甚麼問題，再說明輸出格式如何方便閱讀。",
        "提示：若指令含糊，AI 可能省略「服飾／物品／注意事項」等重要細節。",
        "亦請指出「交叉查證」可如何避免漏抄／抄錯通告重點。",
        "\t\t作答時請留意：不少錯誤源於合併儲存格／複雜版面造成 AI 誤讀。",
        "\t\t若你需要列出核對清單，請用最少三項（例如：日期、負責老師、集合地點）。",
        "",
        "\t(c)\t請分別說明「數據提取」與「邏輯節錄」在文件分析中的用途（各舉一個學習情境）。　（6 分）",
        "\t\t情境例子：通告日程／教科書長文章／課堂講義；請勿抄錄無關長篇。",
        "數據提取：例如抽出日期、地點、聯絡方法；節錄：例如濃縮成三段重點。",
        "請各用約 2–3 句完成；並指出若結果可疑應如何修正（例如對照原件）。",
        "\t\t進階思考：若要比對不同段落的要求，使用表格輸出會否更清晰？為甚麼？",
        "\t\t提醒：使用 AI 製作溫習工具時，仍須回到課本核對是否「一本正經地錯」。",
        "",
        "\t(d)\t請以「負責任使用生成式 AI」為題，寫出兩點建議（例如：引用來源、避免抄襲、私隱保護）。　（2 分）",
        "",
        "",
        "2.\t題組二（Teachable Machine — 手機應用程式筆記）：老師要求同學先完成「Teachable Machine」網站上的手勢模型訓練（第一部分），並於網站中進行測試與取得模型連結；稍後才會把模型放到 App Inventor 專案中。（根據 Dodge Game 筆記「任務 1」）",
        "",
        "\t\t網站：https://teachablemachine.withgoogle.com",
        "\t\t要求概述：建立至少兩個標籤（例如向左／向右）；收集樣本並訓練；測試準確性；取得分享連結。",
        "\t\tTMIC／App Inventor 版面請見題目後表格框架（僅作作答草稿區）。",
        "\t\t請勿於考試期間登入任何個人帳戶；此題為紙筆設計與概念題。",
        "",
        "\t(a)\t請在下表簡述 Teachable Machine「訓練之前／訓練之後」各要完成的一件關鍵事情。（以短句作答即可）　（1 分）",
        "",
        "",
        "\t(b)\tTeachable Machine 建立可供使用的影像分類模型時，下列步驟的正確順序為何？請將 A–D 排列為 1–4。　（3 分）",
        "",
        "",
        "\t\t(i)\t寫出應有的順序（請填 1–4）：",
        "",
        "\t\t(ii)\t請將下列步驟按正確順序排列（請填寫 1–4）：",
        "A.\t匯出／分享模型（例如取得連結方便稍後使用）",
        "B.\t加入影像樣本（網鏡拍攝／上載）",
        "C.\t按下「Train Model」進行訓練",
        "D.\t建立分類標籤（至少兩個類別）",
        "",
        "\t(c)\t為甚麼要在網站完成「充分測試」後，才把模型連結應用到 App Inventor？請從「準確度／使用者體驗」說明。（4 分）",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
    ]


def _apply_cover(doc: Document, meta: Meta) -> None:
    """Patch cover lines in-place; keep template spacing and 考生須知 unchanged."""
    cover_cell = doc.tables[0].cell(0, 0)
    apply_cmp_cover_zh(
        cover_cell,
        ZhCoverPatch(
            year_term=f"{meta.year} {meta.term}",
            paper=meta.paper,
            date_line=f"\t日期:\t{meta.date}",
            time_line=f"\t時間:\t{meta.time}",
            duration_line=f"\t時限:\t{meta.time_limit}",
            pages_line=f"\t頁數:\t{meta.pages}",
            total_line=f"\t總分:\t{meta.total}",
        ),
    )


def _apply_tables(doc: Document) -> None:
    """Rewrite Section B tables for Gen AI + Teachable Machine."""
    # Table 5: keep as MCQ reference area — optional reuse for rubric; leave one instructional row
    doc.tables[5].cell(1, 1).text = (
        "（本分題若印有情境／辨識題，請在此草稿區整理思路；無須填滿。）\n"
        "Teachable Machine：標籤 → 樣本 → Train → 測試 → 分享連結。"
    )

    # Table 6: R-I-C-C-O property sheet (matches old component table shape)
    rows = [
        ("要素", "請填寫你的提示語設計（可短句）"),
        ("R 角色（Role）", "(i)　______________________________"),
        ("I 指令（Instruction）", "(ii)　______________________________"),
        ("C 背景（Context）", "(iii)　_____________________________"),
        ("C 限制（Constraint）", "(iv)　____________________________"),
        ("O 輸出格式（Output）", "(v)　_____________________________"),
        ("（整合檢查）", "(vi)　請用一句話概括你最擔心 AI 會抄錯的重點（例如日期）："),
        ("（備用）", ""),
    ]
    t6 = doc.tables[6]
    for r, (left, right) in enumerate(rows):
        if r >= len(t6.rows):
            break
        t6.cell(r, 0).text = left
        t6.cell(r, 1).text = right

    # Table 7: answer slots for (i)–(vi)
    t7 = doc.tables[7]
    t7.cell(0, 0).text = "1. (a)"
    t7.cell(0, 1).text = "(i)"
    t7.cell(0, 2).text = "(ii)"
    t7.cell(0, 3).text = "(iii)"
    t7.cell(1, 1).text = "(iv)"
    t7.cell(1, 2).text = "(v)"
    t7.cell(1, 3).text = "(vi)"

    # Table 8: scratch workspace — label rows for (b)(c)
    t8 = doc.tables[8]
    t8.cell(0, 0).text = "(b)"
    t8.cell(3, 0).text = "(c)"

    # Table 9: TM steps — labels stay A-D; student fills order in paragraph area
    doc.tables[9].cell(0, 0).text = "(b)\t順序"

    # Table 10: TM reasoning short slots
    doc.tables[10].cell(0, 0).text = "(c)"
    doc.tables[10].cell(0, 1).text = "(i) 準確度"
    doc.tables[10].cell(0, 2).text = "(ii) 使用者體驗"


def generate(template: Path, output: Path, meta: Meta) -> None:
    doc = Document(str(template))
    spans = _mcq_spans()
    blocks = _mcq_blocks()

    _apply_cover(doc, meta)

    for q in range(1, 21):
        start, end = spans[q]
        lines = blocks[q]
        _replace_block_exact(doc, start, end, lines)

    b_start = _find_idx(doc, "乙部 – 問答題", start=0)
    end_b = _find_idx(doc, "~全卷完~", start=b_start)
    sec_b = _section_b_lines()
    _replace_block_exact(doc, b_start, end_b, sec_b)

    _apply_tables(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main() -> int:
    ap = argparse.ArgumentParser(description="Generate S2 CMP Term2 exam from 24-25 template.")
    ap.add_argument(
        "--template",
        default="/Users/warren_chan/Projects/school-agent-os/Subjects/PastPaper/CMP+ICT/2024-2025/2nd Term/F2 CMP/24_25_S2_CMP_Term02_Exam.docx",
        help="Source template docx",
    )
    ap.add_argument(
        "--output",
        default="/Users/warren_chan/Projects/school-agent-os/Subjects/PastPaper/CMP+ICT/2025-2026/S2 CMP/25_26_S2_CMP_Term02_Exam.docx",
        help="Output docx path",
    )
    ap.add_argument("--date", default="__________")
    ap.add_argument("--time", default="__________")
    ns = ap.parse_args()

    meta = Meta(date=ns.date, time=ns.time)
    generate(Path(ns.template), Path(ns.output), meta)
    print(f"Written: {ns.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
