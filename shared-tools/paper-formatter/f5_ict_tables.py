"""F5/S5 ICT written-exam table helpers (24_25 template layout)."""

from __future__ import annotations

from typing import Iterable

from docx import Document
from docx.table import Table

from docx_inplace import (
    clear_all_tables_except,
    clear_table_cells,
    delete_tables_except,
    set_paragraph_text_distribute,
)

# Cover (always keep) + body tables used by 25-26 blueprint content.
# All other template tables (MCQ diagrams, appendix, spare trace grids) are removed.
F5_ICT_REQUIRED_TABLES: frozenset[int] = frozenset(
    {
        0,  # cover
        3,  # B1 spreadsheet
        4,  # B2 validation
        5,  # B3 multimedia spec
        6,  # B4 linear-search trace
        11,  # B6 ORDER sample
        14,  # C2 MOVIE insert
        15,  # C3 Order2024/2025
        18,  # C5 FACILITY
        19,  # C5 RESERVE
        20,  # C6 SQL trace step 1
        21,  # C6 SQL trace step 2
        22,  # C6 SQL trace step 3
        23,  # C6 SQL trace step 4
        24,  # C6 SQL trace step 5
        25,  # C7 transaction reference
        31,  # C8 stack trace
    }
)


def clear_table(table: Table) -> None:
    """Blank every cell while preserving table dimensions and styles."""
    clear_table_cells(table)


def clear_tables(doc: Document, indices: Iterable[int]) -> None:
    for i in indices:
        if 0 <= i < len(doc.tables):
            clear_table(doc.tables[i])


def clear_entire_table(table: Table) -> None:
    """Blank every cell in a table (all rows/columns)."""
    clear_table_cells(table)


def remove_unused_f5_ict_tables(doc: Document) -> None:
    """Delete template tables not required by the current paper (keep cover + body slots)."""
    delete_tables_except(doc, F5_ICT_REQUIRED_TABLES)


def clear_all_body_tables_before_write(doc: Document) -> None:
    """Blank non-cover tables before populate (after unused tables removed)."""
    clear_all_tables_except(doc, keep_indices=(0,))


def set_cell_text(table: Table, r: int, c: int, text: str) -> None:
    cell = table.cell(r, c)
    if cell.paragraphs:
        set_paragraph_text_distribute(cell.paragraphs[0], text)
    else:
        cell.text = text


def apply_f5_ict_table_content(doc: Document) -> None:
    """Populate body tables for the 25-26 DSE blueprint variant."""
    # B1 — charity sale spreadsheet (table 3)
    t3 = doc.tables[3]
    clear_entire_table(t3)
    headers = ["A", "B", "C", "D", "E", "F", "G", "H", "I"]
    for c, h in enumerate(headers):
        set_cell_text(t3, 0, c, h)
    rows = [
        ["1", "日期", "商品", "單價", "數量", "會員", "總價", "統計", "備註"],
        ["2", "2026-04-10", "明信片", "15", "8", "Y", "", "", ""],
        ["3", "2026-04-11", "布袋", "45", "3", "N", "", "", ""],
        ["4", "2026-04-12", "徽章", "20", "6", "Y", "", "", ""],
    ]
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            set_cell_text(t3, r, c, val)

    # B2 — validation methods (table 4)
    t4 = doc.tables[4]
    clear_entire_table(t4)
    set_cell_text(t4, 0, 0, "資料欄")
    set_cell_text(t4, 0, 1, "適當的數據有效性檢驗")
    set_cell_text(t4, 0, 2, "例子")
    samples = [
        ("學號", "一個字母 + 五個數字", "A12345"),
        ("班別", "清單：甲班／乙班／丙班", "乙班"),
        ("年齡", "整數 15–18", "16"),
    ]
    for r, (a, b, c) in enumerate(samples, start=1):
        set_cell_text(t4, r, 0, a)
        set_cell_text(t4, r, 1, b)
        set_cell_text(t4, r, 2, c)

    # B3 — photo spec note (table 5)
    t5 = doc.tables[5]
    clear_entire_table(t5)
    set_cell_text(
        t5,
        0,
        0,
        "相片規格：\n1600 × 1200 像素\n24 bit 真彩色（未壓縮 BMP）\n共 80 張",
    )

    # B4 — linear search trace (table 6 only)
    t6 = doc.tables[6]
    clear_entire_table(t6)
    for c, h in enumerate(["步驟", "i", "found"]):
        if c < len(t6.rows[0].cells):
            set_cell_text(t6, 0, c, h)

    # B6 — ORDER sample (table 11) — 4 columns in template
    t11 = doc.tables[11]
    clear_entire_table(t11)
    hdr = ["OID", "CNAME", "PHONE", "EMAIL"]
    for c, h in enumerate(hdr):
        set_cell_text(t11, 0, c, h)
    data = [
        ("O001", "Wong Ka Ming", "61234567", "wkm@mail.com"),
        ("O002", "Lau Hoi Yan", "62345678", "lhy@mail.com"),
        ("O003", "Cheung Tsz Kin", "63456789", "ctk@mail.com"),
    ]
    for r, row in enumerate(data, start=1):
        for c, val in enumerate(row):
            set_cell_text(t11, r, c, val)

    # C2 — MOVIE insert sample (table 14)
    t14 = doc.tables[14]
    clear_entire_table(t14)
    set_cell_text(
        t14,
        0,
        0,
        '影碟編號：M10001\n片名："Coding Fun"\n租金：35',
    )

    # C3 — Order2024 / Order2025 (table 15; template has extra columns — clear all)
    t15 = doc.tables[15]
    clear_entire_table(t15)
    for c, h in enumerate(["PID", "SID", "AMT"]):
        set_cell_text(t15, 0, c, h)
    rows15 = [
        ("P01", "S01", "150"),
        ("P02", "S02", "90"),
        ("P03", "S03", "0"),
        ("P04", "S04", "60"),
    ]
    for r, row in enumerate(rows15, start=1):
        for c, val in enumerate(row):
            set_cell_text(t15, r, c, val)

    # C5 — FACILITY / RESERVE (tables 18–19)
    t18 = doc.tables[18]
    clear_entire_table(t18)
    for c, h in enumerate(["欄名", "數據類型", "描述", "例子"]):
        set_cell_text(t18, 0, c, h)
    set_cell_text(t18, 1, 0, "FID")
    set_cell_text(t18, 1, 1, "CHAR(4)")
    set_cell_text(t18, 1, 2, "設施編號")
    set_cell_text(t18, 1, 3, "F101")
    set_cell_text(t18, 2, 0, "FNAME")
    set_cell_text(t18, 2, 1, "VARCHAR(30)")
    set_cell_text(t18, 2, 2, "設施名稱")
    set_cell_text(t18, 2, 3, "舞蹈室")
    set_cell_text(t18, 3, 0, "CAPACITY")
    set_cell_text(t18, 3, 1, "INTEGER")
    set_cell_text(t18, 3, 2, "容納人數")
    set_cell_text(t18, 3, 3, "20")

    t19 = doc.tables[19]
    clear_entire_table(t19)
    for c, h in enumerate(["欄名", "數據類型", "描述", "例子"]):
        set_cell_text(t19, 0, c, h)
    set_cell_text(t19, 1, 0, "RID")
    set_cell_text(t19, 1, 1, "CHAR(7)")
    set_cell_text(t19, 1, 2, "預約編號")
    set_cell_text(t19, 1, 3, "R260501")
    set_cell_text(t19, 2, 0, "MEMID")
    set_cell_text(t19, 2, 1, "CHAR(6)")
    set_cell_text(t19, 2, 2, "會員編號")
    set_cell_text(t19, 2, 3, "M10001")
    set_cell_text(t19, 3, 0, "RDATE")
    set_cell_text(t19, 3, 1, "DATE")
    set_cell_text(t19, 3, 2, "預約日期")
    set_cell_text(t19, 3, 3, "2026-07-12")

    # C6 — query trace grids (tables 20–24)
    for ti in range(20, 25):
        tb = doc.tables[ti]
        clear_entire_table(tb)
        set_cell_text(tb, 0, 0, "步驟")
        if len(tb.rows[0].cells) > 1:
            set_cell_text(tb, 0, 1, "結果列")

    # C7 — transaction reference (table 25)
    t25 = doc.tables[25]
    clear_entire_table(t25)
    set_cell_text(t25, 0, 0, "SQL 子句")
    set_cell_text(t25, 0, 1, "用途")
    ops = [
        ("BEGIN TRANSACTION", "開始交易"),
        ("COMMIT / ROLLBACK", "確認或還原變更"),
    ]
    for r, (a, b) in enumerate(ops, start=1):
        set_cell_text(t25, r, 0, a)
        set_cell_text(t25, r, 1, b)

    # C8 — stack trace (table 31)
    t31 = doc.tables[31]
    clear_entire_table(t31)
    set_cell_text(t31, 0, 0, "步驟")
    set_cell_text(t31, 0, 1, "操作")
    set_cell_text(t31, 0, 2, "頂端")
    set_cell_text(t31, 0, 3, "Stack（底→頂）")
    stack_ops = [
        ("0", "初始", "—", "∅"),
        ("1", "PUSH 3", "3", "3"),
        ("2", "PUSH 7", "7", "3, 7"),
        ("3", "POP", "3", "3"),
    ]
    for r, row in enumerate(stack_ops, start=1):
        for c, val in enumerate(row):
            if c < len(t31.rows[r].cells):
                set_cell_text(t31, r, c, val)
