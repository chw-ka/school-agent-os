#!/usr/bin/env python3
"""Generate easier 25-26 S3 CMP Term 2 written exam from S2 Term 2 template (MCQ/Matching/T-F/Fill)."""

from __future__ import annotations

import argparse
import random
import re
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from docx_inplace import ZhCoverPatch, apply_cmp_cover_zh, set_paragraph_text_distribute

from mcq_answer_keys import build_random_mcq_key
from matching_layout import build_shuffled_matching
from s3_cmp_term2_exam_spec import (
    build_s3_cmp_term2_exam_spec,
    format_fill_rubric_blocks,
    format_matching_rubric_blocks,
    format_tf_rubric,
)
from tf_fill_layout import build_shuffled_fill, build_shuffled_tf

_RNG = random.Random(20260517)

REPO = Path(__file__).resolve().parents[2]
S2_TEMPLATE = (
    REPO
    / "Subjects/S2-CMP/past-papers/2025-2026/Term 02/WrittenExam/25_26_S2_CMP_Term02_Exam.docx"
)
DEFAULT_OUT = (
    REPO
    / "Subjects/S3-CMP/past-papers/2025-2026/Term 02/WrittenExam/25_26_S3_CMP_Term02_Exam.docx"
)

# Paragraph spans for MCQ blocks (from 25_26_S2_CMP_Term02_Exam.docx).
MCQ_SPANS: dict[int, tuple[int, int]] = {
    1: (5, 16),
    2: (16, 19),
    3: (19, 22),
    4: (22, 25),
    5: (25, 28),
    6: (28, 30),
    7: (30, 32),
    8: (32, 35),
    9: (35, 41),
    10: (41, 44),
    11: (44, 47),
    12: (47, 54),
    13: (54, 61),
    14: (61, 67),
    15: (67, 74),
    16: (74, 82),
    17: (82, 90),
    18: (90, 97),
    19: (97, 104),
    20: (104, 107),
}


@dataclass(frozen=True)
class Meta:
    school: str = "迦密聖道中學"
    year: str = "2025 – 2026"
    term: str = "下學期考試"
    level: str = "中三級 電腦認知"
    paper: str = "試題簿"
    date: str = "__________"
    time: str = "__________"
    time_limit: str = "30 分鐘"
    pages: str = "9 頁"
    total: str = "50"


def _find_idx(doc: Document, contains: str, *, start: int = 0) -> int:
    for i in range(start, len(doc.paragraphs)):
        if contains in doc.paragraphs[i].text:
            return i
    raise ValueError(f'missing: "{contains}"')


def _find_mcq_answer_line(doc: Document, *, start: int = 150) -> int:
    for i in range(start, len(doc.paragraphs)):
        compact = re.sub(r"\s+", "", doc.paragraphs[i].text.upper())
        if len(compact) >= 20 and all(c in "ABCD" for c in compact[:20]):
            return i
    raise ValueError("missing MCQ answer key line in answer sheet")


def _replace_block(doc: Document, start: int, end_excl: int, lines: list[str]) -> None:
    span = end_excl - start
    if len(lines) > span:
        raise ValueError(f"Need ≤{span} lines at {start}:{end_excl}, got {len(lines)}")
    for i in range(span):
        set_paragraph_text_distribute(
            doc.paragraphs[start + i], lines[i] if i < len(lines) else ""
        )


def _mcq_opts_compact(options: list[str]) -> str:
    """Tab-indented A–D; each option on its own line."""
    return "\n".join(f"\t{L}.\t{t}" for L, t in zip("ABCD", options))


def _mcq_stem_and_opts(q: int, stem: str, options: list[str]) -> str:
    return f"{q}.\t{stem}\n{_mcq_opts_compact(options)}"


def _mcq_opts_paragraph(options: list[str]) -> str:
    """Options-only paragraph: leading newline keeps tab before A. in Word."""
    return "\n" + _mcq_opts_compact(options)


def _mcq_compact(q: int, stem: str, options: list[str]) -> list[str]:
    return [_mcq_stem_and_opts(q, stem, options), "", ""]


def _mcq_micro(q: int, stem: str, options: list[str]) -> list[str]:
    return [_mcq_stem_and_opts(q, stem, options), ""]


def _mcq_spread(
    q: int,
    stem: str,
    options: list[str],
    *,
    trailing_blank: bool = True,
    extra_blanks: int = 0,
) -> list[str]:
    lines = [f"{q}.\t{stem}", _mcq_opts_paragraph(options)]
    if trailing_blank:
        lines.append("")
    lines.extend([""] * extra_blanks)
    return lines


def _mcq_blocks() -> dict[int, list[str]]:
    """甲部 MCQ：情境題；函數名稱等留乙–戊。組合選項題（如 Q1）不打亂 A–D，其餘題打亂選項次序。"""
    return {
        1: [
            "1.\t老師要求用生成式 AI 整理一份 PDF 通告重點，下列哪項較屬合理應用？",
            "",
            "\t\t(1)\t把長文摘成短摘要",
            "\t\t(2)\t就內容提出問答",
            "\t\t(3)\t協助檢查是否遺漏日期／地點",
            "",
            "\tA.\t只有 (1)",
            "\tB.\t只有 (1) 和 (2)",
            "\tC.\t只有 (2) 和 (3)",
            "\tD.\t(1)、(2) 和 (3) 皆是",
            "",
        ],
        2: _mcq_compact(
            2,
            "執行 import cv2 時出現 'ModuleNotFoundError'，較合理處理是？",
            ["刪除 Python", "用 pip 安裝套件", "改用 Word", "不用處理"],
        ),
        3: _mcq_compact(
            3,
            "AI 捏造看似可信但不存在的內容，較適合稱為？",
            ["快取", "降噪", "幻覺（Hallucination）", "迭代"],
        ),
        4: _mcq_compact(
            4,
            "用 AI 協助寫報告，下列哪項最符合「負責任使用」？",
            ["只貼 AI 答案", "完全不用工具", "假設 AI 永遠正確", "對照官方資料再寫"],
        ),
        5: _mcq_compact(
            5,
            "語音轉文字程式因網絡中斷崩潰，較應加入？",
            ["try…except", "for 迴圈", "import json", "while True"],
        ),
        6: _mcq_micro(
            6,
            "要在影片追蹤足球，下列步驟次序最合理？",
            [
                "init → selectROI → update",
                "selectROI → init → update",
                "只 update",
                "只 imread",
            ],
        ),
        7: _mcq_micro(
            7,
            "人臉偵測誤判框太多，較合理調整是？",
            ["降低 minNeighbors", "刪模型檔", "提高 minNeighbors", "播 mp3"],
        ),
        8: _mcq_compact(
            8,
            "使用生成式 AI 閱讀長 PDF 時，下列哪項較**不**合理？",
            ["提取日期", "整理要點", "翻譯關鍵字", "刪除同學檔案"],
        ),
        9: _mcq_spread(
            9,
            "'cv2.VideoCapture(0)' 的 0 在課堂實驗通常指？",
            ["預設鏡頭", "第 0 張圖", "第 0 秒", "靜音"],
        ),
        10: _mcq_compact(
            10,
            "Vibe Coding 下，學生最合適扮演？",
            ["抄襲 AI", "總監（構思、測試、修正）", "只畫圖", "只閱卷"],
        ),
        11: _mcq_compact(
            11,
            "AI 寫的遊戲太快，學生最應？",
            ["直接交卷", "刪除程式", "測試後請 AI 調整（迭代）", "改紙筆玩"],
        ),
        12: _mcq_spread(
            12,
            "拍照做人臉偵測前先把影像轉灰階，主要原因是？",
            ["令相片變彩色", "用來播放 mp3", "用來翻譯句子", "減少運算量、加快處理"],
        ),
        13: _mcq_spread(
            13,
            "即時鏡頭追蹤要較流暢，較宜選 KCF 而非 CSRT，因為？",
            ["KCF 一般較快", "KCF 一定較準確", "CSRT 不能追蹤", "兩者完全相同"],
        ),
        14: _mcq_spread(
            14,
            "Google 翻譯 API 顯示找不到憑證，較可能是？",
            ["未裝 tkinter", "未設定 password.json 路徑", "未灰階", "未圈選 ROI"],
        ),
        15: _mcq_spread(
            15,
            "要刪除專案內舊的 sound.mp3，較應使用？",
            ["gTTS", "detectMultiScale", "os.remove", "recognize_google"],
        ),
        16: _mcq_spread(
            16,
            "測驗程式把題目放在 quiz_data.json，主要是為了？",
            ["不用寫任何程式", "一定不會出錯", "只能老師使用", "程式與題目資料分開，方便更新題庫"],
        ),
        17: _mcq_spread(
            17,
            "製作有按鈕、選項的測驗視窗介面，較應使用？",
            ["tkinter", "gTTS", "Haar Cascade", "playsound"],
        ),
        18: _mcq_spread(
            18,
            "追蹤準確度要求高、速度可慢，較應選哪種 tracker？",
            ["MOSSE", "CSRT", "KCF", "不用 tracker"],
        ),
        19: _mcq_spread(
            19,
            "把 AI 生成內容直接交功課且完全不註明，最違反？",
            ["檔名", "亮度", "清潔", "學術誠信"],
        ),
        20: _mcq_compact(
            20,
            "課堂鏡頭畫面全黑，較應先檢查？",
            [
                "是否已用 pip 安裝 gTTS",
                "JSON 格式",
                "是否已轉灰階",
                "鏡頭／VideoCapture 是否成功開啟",
            ],
        ),
    }


def _section_b_paragraphs() -> dict[str, list[str]]:
    return {
        "header": [
            "乙部 – 配對題 (10 分)",
            "函數庫配對：請將正確功能字母填入括號內。\t(5分)",
        ],
        "bank1": [
            "\t\t功能庫：",
            "A. 播放 mp3 音訊",
            "B. 把文字轉成語音（TTS）",
            "C. 把語音轉成文字（STT）",
            "D. 在影像中偵測人臉",
            "E. 刪除資料夾內的檔案",
        ],
        "header2": [
            "程式步驟配對：根據描述選出最合適的步驟字母。\t(5分)",
        ],
        "bank2": [
            "步驟：",
            "A. adjust_for_ambient_noise()",
            "B. audio.save(\"audio.mp3\")",
            "C. recognize_google()",
            "D. cv2.cvtColor(..., cv2.COLOR_BGR2GRAY)",
            "E. cv2.selectROI()",
        ],
    }


def _short_answer() -> str:
    return (
        "根據筆記，Vibe Coding 為什麼需要「迭代（Iteration）」？"
        "請舉一個例子說明。（5 分）"
    )


def _apply_cover(doc: Document, meta: Meta) -> None:
    cover_cell = doc.tables[0].cell(0, 0)
    apply_cmp_cover_zh(
        cover_cell,
        ZhCoverPatch(
            school=meta.school,
            year_term=f"{meta.year} {meta.term}",
            level=meta.level,
            paper=meta.paper,
            date_line=f"\t日期:\t{meta.date}",
            time_line=f"\t時間:\t{meta.time}",
            duration_line=f"\t時限:\t{meta.time_limit}",
            pages_line=f"\t頁數:\t{meta.pages}",
            total_line=f"\t總分:\t{meta.total}",
        ),
    )


def _apply_matching_tables(
    doc: Document,
    rows1: list[tuple[str, str]],
    rows2: list[tuple[str, str]],
) -> None:
    for t, rows in ((doc.tables[1], rows1), (doc.tables[2], rows2)):
        for r, (a, b) in enumerate(rows):
            t.cell(r, 0).text = a
            t.cell(r, 1).text = b


def _apply_word_banks(doc: Document, bank_a: list[str], bank_b: list[str]) -> None:
    doc.tables[3].cell(0, 0).text = "\t" + "\t\t".join(bank_a)
    doc.tables[4].cell(0, 0).text = "\t" + "\t\t".join(bank_b)


def _apply_answer_sheet(
    doc: Document,
    meta: Meta,
    *,
    mcq_answers_display: str,
    matching_answers: list[str],
    tf_answers: str,
    fill_answers: list[list[str]],
) -> None:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "中二級 電腦認知":
            set_paragraph_text_distribute(p, meta.level)
        if "中二級 電腦認知科" in p.text:
            set_paragraph_text_distribute(
                p, f"{meta.year} {meta.level}科{meta.term} (評分準則)"
            )
    ans_mcq = _find_mcq_answer_line(doc, start=150)
    set_paragraph_text_distribute(doc.paragraphs[ans_mcq], mcq_answers_display)

    match_ans = format_matching_rubric_blocks(matching_answers)
    start = _find_idx(doc, "乙部： 配對題", start=150)
    # find line after header "乙部： 配對題"
    idx = start + 1
    for j, line in enumerate(match_ans):
        if idx + j < len(doc.paragraphs):
            set_paragraph_text_distribute(doc.paragraphs[idx + j], line)

    tf_ans = format_tf_rubric(tf_answers)
    tf_start = _find_idx(doc, "丙部： 是非題", start=150)
    for j, line in enumerate(tf_ans):
        set_paragraph_text_distribute(doc.paragraphs[tf_start + 1 + j], line)

    fill_ans = format_fill_rubric_blocks(fill_answers)
    fill_start = _find_idx(doc, "丁部：填充題", start=150)
    for j, line in enumerate(fill_ans):
        set_paragraph_text_distribute(doc.paragraphs[fill_start + 1 + j], line)

    sa_start = _find_idx(doc, "戊部：短答題", start=150)
    set_paragraph_text_distribute(
        doc.paragraphs[sa_start + 1],
        "1.\t\tAI 很少一次就完美；學生要測試後再提出修改（例如調慢速度、改顏色）。"
        "例子：接蘋果遊戲太慢／太快要叫 AI 改。",
    )


def generate(template: Path, output: Path, meta: Meta) -> tuple[str, list[str], str, list[list[str]]]:
    doc = Document(str(template))

    _apply_cover(doc, meta)

    raw_blocks = _mcq_blocks()
    blocks, mcq_key = build_random_mcq_key(raw_blocks, rng=_RNG)
    mcq_display = " ".join(mcq_key[i : i + 5] for i in range(0, 20, 5))

    match_rows1, match_rows2, matching_keys = build_shuffled_matching(_RNG)
    tf_lines, tf_key = build_shuffled_tf(_RNG)
    fill_lines, fill_answers, bank_a, bank_b = build_shuffled_fill(_RNG)

    for q, (start, end) in MCQ_SPANS.items():
        span = end - start
        block = list(blocks[q])
        if len(block) > span:
            raise ValueError(f"MCQ#{q} needs ≤{span} lines, got {len(block)}")
        if len(block) < span:
            block.extend([""] * (span - len(block)))
        _replace_block(doc, start, end, block)

    b = _section_b_paragraphs()
    set_paragraph_text_distribute(doc.paragraphs[107], "")
    _replace_block(doc, 108, 110, b["header"])
    _replace_block(doc, 110, 116, b["bank1"])
    _replace_block(doc, 117, 119, b["header2"])
    _replace_block(doc, 119, 125, b["bank2"])

    _replace_block(doc, 130, 135, tf_lines)
    _replace_block(doc, 137, 153, fill_lines)
    set_paragraph_text_distribute(doc.paragraphs[155], _short_answer())

    idx_fill = _find_idx(doc, "丁部", start=0)
    if "(5分)" in doc.paragraphs[idx_fill].text:
        set_paragraph_text_distribute(doc.paragraphs[idx_fill], "丁部 – 填充題 (10 分)")

    _apply_matching_tables(doc, match_rows1, match_rows2)
    _apply_word_banks(doc, bank_a, bank_b)
    _apply_answer_sheet(
        doc,
        meta,
        mcq_answers_display=mcq_display,
        matching_answers=matching_keys,
        tf_answers=tf_key,
        fill_answers=fill_answers,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))

    spec_path = output.with_suffix(".spec.json")
    _save_spec_and_apply_footer(
        output,
        spec_path,
        mcq_key=mcq_key,
        matching_answers=matching_keys,
        tf_lines=tf_lines,
        tf_answers=tf_key,
        fill_answers=fill_answers,
        fill_word_banks=[bank_a, bank_b],
    )
    return mcq_key, matching_keys, tf_key, fill_answers


def _save_spec_and_apply_footer(
    docx_path: Path,
    spec_path: Path,
    *,
    mcq_key: str,
    matching_answers: list[str],
    tf_lines: list[str],
    tf_answers: str,
    fill_answers: list[list[str]],
    fill_word_banks: list[list[str]],
) -> None:
    import sys

    for _pkg in ("question-quality-check", "paper-quality-check"):
        sys.path.insert(0, str(Path(__file__).resolve().parents[1] / _pkg))
    from exam_spec import save_spec
    from footer import apply_footer_meta

    spec = build_s3_cmp_term2_exam_spec(
        mcq_answers=mcq_key,
        matching_answers=matching_answers,
        tf_lines=tf_lines,
        tf_answers=tf_answers,
        fill_answers=fill_answers,
        fill_word_banks=fill_word_banks,
    )
    save_spec(spec_path, spec)
    footer = spec.get("meta", {}).get("footer")
    if footer:
        apply_footer_meta(docx_path, footer)


def _run_quality_check(docx_path: Path, spec_path: Path) -> int:
    import os
    import subprocess
    import sys

    root = Path(__file__).resolve().parents[1]
    env = {**os.environ, "PYTHONIOENCODING": "utf-8"}
    py = sys.executable
    q_cmd = [
        py,
        str(root / "question-quality-check" / "check_docx.py"),
        "--candidate",
        str(docx_path.resolve()),
        "--candidate-spec",
        str(spec_path.resolve()),
        "--subject",
        "S3 CMP",
        "--years",
        "3",
    ]
    q = subprocess.run(q_cmd, cwd=str(root), env=env)
    p_cmd = [
        py,
        str(root / "paper-quality-check" / "check_docx.py"),
        "--candidate",
        str(docx_path.resolve()),
        "--candidate-spec",
        str(spec_path.resolve()),
        "--skip-written",
    ]
    p = subprocess.run(p_cmd, cwd=str(root), env=env)
    return max(q.returncode, p.returncode)


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--template", type=Path, default=S2_TEMPLATE)
    ap.add_argument("--output", type=Path, default=DEFAULT_OUT)
    ap.add_argument("--date", default="__________")
    ap.add_argument("--time", default="__________")
    ap.add_argument("--skip-check", action="store_true", help="Skip quality checks after generate")
    ns = ap.parse_args()
    meta = Meta(date=ns.date, time=ns.time)
    mcq_key, match_keys, tf_key, fill_ans = generate(ns.template, ns.output, meta)
    print(f"MCQ key: {' '.join(mcq_key[i:i+5] for i in range(0,20,5))}")
    print(f"Matching: {match_keys[0]} / {match_keys[1]}")
    print(f"T/F key: {tf_key}")
    print(f"Fill block1: {fill_ans[0]}")
    print(f"Fill block2: {fill_ans[1]}")
    print(f"Written: {ns.output}")
    spec_path = ns.output.with_suffix(".spec.json")
    print(f"Written: {spec_path}")
    if not ns.skip_check:
        code = _run_quality_check(ns.output, spec_path)
        if code != 0:
            print("Quality checks reported issues (see above).", file=sys.stderr)
            return code
        print("Quality checks: all passed.")
    gen_dir = ns.output.parent.parent / "_generation"
    gen_dir.mkdir(parents=True, exist_ok=True)
    gen_spec = gen_dir / spec_path.name
    gen_spec.write_text(spec_path.read_text(encoding="utf-8"), encoding="utf-8")
    print(f"Copied spec: {gen_spec}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
