from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from docx_inplace import (
    ZhCoverPatch,
    apply_cmp_cover_zh,
    replace_in_paragraph_runs,
    set_paragraph_text_distribute,
)


@dataclass(frozen=True)
class ExamMeta:
    school_name: str
    year_text: str
    term_text: str  # e.g. 上學期考試 / 下學期考試
    subject_text: str  # e.g. 中三級 電腦認知
    paper_text: str  # e.g. 試題簿
    date_text: str  # e.g. 2026年1月15日 or __________
    time_text: str  # e.g. 上午10:45 – 上午11:15 or __________
    duration_text: str  # e.g. 30分鐘
    pages_text: str  # e.g. 5頁
    total_marks_text: str  # e.g. 50


def _replace_in_all_paragraphs(doc: Document, needle: str, replacement: str) -> None:
    for p in doc.paragraphs:
        if needle in p.text:
            p.text = p.text.replace(needle, replacement)


def _find_paragraph_index(doc: Document, contains: str, *, start: int = 0) -> int:
    for i in range(start, len(doc.paragraphs)):
        if contains in doc.paragraphs[i].text:
            return i
    raise ValueError(f'Cannot find paragraph containing: "{contains}"')


def _set_paragraph_block_exact(doc: Document, start_idx: int, end_idx_exclusive: int, lines: list[str]) -> None:
    """
    Overwrite doc.paragraphs[start_idx:end_idx_exclusive] with provided lines.
    - Never inserts/appends paragraphs (to avoid layout shifting).
    - If fewer lines than available paragraphs, remaining paragraphs are blanked.
    - If more lines than available paragraphs, raise to prevent structure drift.
    """
    span = max(0, end_idx_exclusive - start_idx)
    if len(lines) > span:
        raise ValueError(
            f"Replacement has {len(lines)} lines but only {span} paragraphs available "
            f"(range {start_idx}:{end_idx_exclusive})."
        )
    for i in range(span):
        _set_paragraph_text_distribute(doc.paragraphs[start_idx + i], lines[i] if i < len(lines) else "")


def _set_paragraph_text_distribute(paragraph, text: str) -> None:
    set_paragraph_text_distribute(paragraph, text)


def _replace_in_paragraph_runs(paragraph, needle: str, replacement: str) -> None:
    replace_in_paragraph_runs(paragraph, needle, replacement)


def _set_table_cell_text(table, r: int, c: int, text: str) -> None:
    cell = table.cell(r, c)
    # Preserve the cell's paragraph/run structure by writing into existing paragraphs.
    # If the caller passes newlines, map them onto existing paragraphs where possible.
    lines = text.split("\n")
    for i, p in enumerate(cell.paragraphs):
        _set_paragraph_text_distribute(p, lines[i] if i < len(lines) else "")
    # Never append new paragraphs (would shift XML/layout). Extra lines are dropped.


def _build_term2_content():
    # Keep EXACT paragraph structure of the template (no insertions).
    # Section A blocks must match the template's per-question paragraph spans:
    # Q01 len=7, Q02-08 len=3, Q09 len=10, Q10-11 len=3, Q12 len=12, Q13 len=8, Q14 len=6, Q15 len=8.

    mcq_blocks: dict[int, list[str]] = {
        1: [
            "1.\t下列哪一項最能描述「外建／第三方函數庫」？",
            "",
            "\tA.\t跟 Python 一起安裝，無需額外處理",
            "\tB.\t需要用 pip 安裝後才能 import 使用",
            "\tC.\t只能在網頁上使用，不能在電腦上用",
            "\tD.\t只能用於 AI，不可用於一般程式",
            "",
        ],
        2: [
            "2.\t在 TTS 程式中，負責把文字轉成語音（mp3）的函數庫是？",
            "\nA.\tgTTS\nB.\tplaysound\nC.\tos\nD.\tjson",
            "",
        ],
        3: [
            "3.\t在課堂範例中，為了避免播放失敗，playsound 建議安裝哪個版本？",
            "\nA.\tplaysound==2.0\nB.\tplaysound==1.2.2\nC.\tplaysound==0.1\nD.\t不用安裝 playsound",
            "",
        ],
        4: [
            "4.\tCloud Translation 範例中，以下哪一行用來設定 Google 憑證檔？",
            "\nA.\tos.environ[\"GOOGLE_APPLICATION_CREDENTIALS\"] = \"password.json\"\nB.\tclient = translate.Client()\nC.\tresult = client.translate(text, target_language=\"zh\")\nD.\tpip install google-cloud-translate==2.0.1",
            "",
        ],
        5: [
            "5.\tSpeechRecognition 中，用來做環境降噪的步驟是？",
            "\nA.\trecognizer.listen(source)\nB.\trecognizer.adjust_for_ambient_noise(source)\nC.\trecognizer.record(source)\nD.\trecognizer.recognize_google(audio)",
            "",
        ],
        6: [
            "6.\t如果要辨識「廣東話」，recognize_google 的 language 參數應使用？",
            "\nA.\ten\nB.\tzh\nC.\tyue\nD.\tjp",
            "",
        ],
        7: [
            "7.\t在 OpenCV 人臉偵測流程中，把彩色圖片轉成灰階通常使用？",
            "\nA.\tcv2.imread(img, cv2.COLOR_BGR2GRAY)\nB.\tcv2.cvtColor(img, cv2.COLOR_BGR2GRAY)\nC.\tcv2.gray(img)\nD.\tcv2.convert(img, \"gray\")",
            "",
        ],
        8: [
            "8.\tHaar Cascade 的人臉模型檔案通常是什麼格式？",
            "\nA.\t.py\nB.\t.json\nC.\t.xml\nD.\t.mp3",
            "",
        ],
        9: [
            "9.\t以下哪一段程式最符合「先圈選 ROI，再開始追蹤物件」？",
            "\nA.\ttracker = cv2.TrackerCSRT_create()",
            "\troi = cv2.selectROI(frame)",
            "\ttracker.init(frame, roi)",
            "\ttracker.update(frame)",
            "B.\troi = cv2.selectROI(frame)",
            "\ttracker = cv2.TrackerCSRT_create()",
            "\ttracker.init(frame, roi)",
            "\ttracker.update(frame)",
            "",
        ],
        10: [
            "10.\tcv2.VideoCapture(0) 中的 0 代表？",
            "\nA.\t第 0 張圖片\nB.\t預設鏡頭裝置\nC.\t影片第 0 秒\nD.\t靜音模式",
            "",
        ],
        11: [
            "11.\t在 detectMultiScale 中，提高 minNeighbors 通常會令結果？",
            "\nA.\t更容易誤判（更多假陽性）\nB.\t更穩定、誤判較少（更嚴格）\nC.\t只能偵測到更小的臉\nD.\t只影響顏色，不影響偵測",
            "",
        ],
        12: [
            "12.\t以下哪一個 `quiz_data.json` 結構最合理？",
            "",
            "\tA.\t{\"question\": \"...\", \"options\": [\"A\",\"B\"], \"answer\": 0}",
            "\tB.\t{\"q\": \"...\", \"opt\": \"A,B\", \"ans\": \"A\"}",
            "\tC.\t{\"question\": [\"...\"], \"answer\": \"A\"}",
            "\tD.\t{\"options\": [\"A\",\"B\"], \"answer\": 0}",
            "",
            "",
            "",
            "",
            "",
            "",
        ],
        13: [
            "13.\t物件追蹤中，以下哪個追蹤器通常「準確度很好，但速度較慢」？",
            "",
            "A.\tMOSSE",
            "B.\tKCF",
            "C.\tCSRT",
            "D.\tBOOSTING",
            "",
            "",
        ],
        14: [
            "14.\t若初始速度為 5，Level = 3，Speed = 5 × (1.2)^(Level−1)，Speed 約為？",
            "\nA.\t5.0",
            "B.\t6.0",
            "C.\t7.2",
            "D.\t8.64",
            "",
        ],
        15: [
            "15.\tSpeechRecognition 中，哪個方法用於把語音轉成文字？",
            "",
            "A.\tlisten()",
            "B.\trecord()",
            "C.\trecognize_google()",
            "D.\ttranslate()",
            "",
            "",
        ],
    }

    # Matching table (5 rows): left term + right description letters A-E (in descriptions)
    matching_rows = [
        ("gTTS", "A. 播放已儲存的 mp3 音訊"),
        ("playsound()", "B. 把文字轉換成語音並輸出 mp3"),
        ("SpeechRecognition", "C. 把語音轉成文字（STT）"),
        ("detectMultiScale()", "D. 在圖片中找出人臉並回傳座標"),
        ("selectROI()", "E. 手動圈選要追蹤的目標區域"),
    ]

    # Section C/D block must fit template span exactly.
    output_lines: list[str] = [
        "丙部 – 短答題（15 分）",
        "請用簡短文字回答下列問題。",
        "",
        "1.\t甚麼是「外建／第三方函數庫」？請舉出一個例子。（3 分）",
        "",
        "2.\t在 Cloud Translation 範例中，`password.json` 有甚麼用途？（3 分）",
        "",
        "3.\t簡述 SpeechRecognition 進行語音辨識時，為何要加入 `try…except`。（3 分）",
        "",
        "4.\tOpenCV 的 `detectMultiScale()` 會回傳甚麼資料？（3 分）",
        "",
        "5.\t簡述「物件偵測」與「物件追蹤」的分別。（3 分）",
        "",
        "丁部 – 結構題（15 分）",
        "請按題目要求作答。",
        "",
        "1.\t《翻譯播音員》系統設計\t(5分)",
        "",
        "",
        "2.\tOpenCV 人臉偵測：降低誤判\t(5分)",
        "",
        "",
        "3.\tVibe Coding：難度曲線\t(5分)",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        "~全卷完~",
    ]

    # Tables 2-4: reuse as structured question answer areas (keep same table layout).
    fill1_code = (
        "題目：某同學要製作《翻譯播音員》（廣東話 → 英文 → 播放英文語音）。\n"
        "(a) 請寫出系統流程的三個主要步驟。（3 分）\n"
        "(b) 請指出至少兩個需要用到的函數庫。（2 分）\n"
    )
    fill1_out = "（作答區）"

    fill2_code = (
        "題目：同學使用 OpenCV 進行人臉偵測，但出現誤判（把非人臉當成人臉）。\n"
        "(a) 請提出一個可能原因。（2 分）\n"
        "(b) 請指出應調整哪一個參數以減少誤判，並簡述調整方向。（3 分）\n"
    )
    fill2_out = "（作答區）"

    fill3_code = (
        "題目：遊戲設計的難度曲線公式為  Speed = 5 × (1.2)^(Level−1)。\n"
        "(a) Level = 1、2、3 時，Speed 分別是多少（保留 2 位小數）？（3 分）\n"
        "(b) 簡述為何使用「指數增長」能令玩家更有壓力。（2 分）\n"
    )
    fill3_out = "（作答區）"

    # Answer key
    # Answer key formatted like template line groups: "ABCDE ABCDE ABCDE" but with A-D letters.
    mcq_letters = list("BABBBCBCBBBCDAC")
    mcq_answers = "".join(mcq_letters[:5]) + " " + "".join(mcq_letters[5:10]) + " " + "".join(mcq_letters[10:15])
    matching_answers = ["1. B", "2. A", "3. C", "4. D", "5. E"]
    fill_answers = {
        "fill1": {
            "(i)": "流程：STT → 翻譯 → TTS 播放（或：輸入語音→文字→翻譯→語音輸出）",
            "(ii)": "函數庫：speech_recognition、google-cloud-translate_v2、gTTS、playsound",
        },
        "fill2": {
            "(i)": "原因：光線/背景/角度/模型限制導致誤判",
            "(ii)": "調整：minNeighbors（提高可減少誤判）",
        },
        "fill3": {
            "(i)": "Level1=5.00；Level2=6.00；Level3=7.20",
            "(ii)": "原因：增長更快，令玩家更快感到壓力/挑戰",
        },
    }

    return {
        "mcq_blocks": mcq_blocks,
        "matching_rows": matching_rows,
        "output_block": output_lines,
        "fill_tables": [(fill1_code, fill1_out), (fill2_code, fill2_out), (fill3_code, fill3_out)],
        "mcq_answers": mcq_answers,
        "matching_answers": matching_answers,
        "fill_answers": fill_answers,
    }


def build_term2_exam_docx(*, template_path: Path, output_path: Path, meta: ExamMeta) -> None:
    doc = Document(str(template_path))

    cover_cell = doc.tables[0].cell(0, 0)
    apply_cmp_cover_zh(
        cover_cell,
        ZhCoverPatch(
            school=meta.school_name,
            year_term=f"{meta.year_text} {meta.term_text}",
            term_needle="上學期考試",
            term_replacement=meta.term_text,
            level=meta.subject_text,
            paper=meta.paper_text,
            date_line=f"\t日期:\t{meta.date_text}",
            time_line=f"\t時間:\t{meta.time_text}",
            duration_line=f"\t時限:\t{meta.duration_text}",
            pages_line=f"\t頁數:\t{meta.pages_text}",
            total_line=f"\t總分:\t{meta.total_marks_text}",
        ),
    )

    content = _build_term2_content()

    # Replace MCQ questions in-place (do not change paragraph count).
    # Template structure: keep headings and blank lines; only overwrite Q1-15 blocks.
    import re

    a_start = _find_paragraph_index(doc, "甲部 – 多項選擇題", start=0)
    b_heading = _find_paragraph_index(doc, "乙部 – 配對題", start=a_start + 1)
    stem_re = re.compile(r"^(\d+)\.\t")
    stem_idxs: list[int] = []
    for i in range(a_start, b_heading):
        m = stem_re.match(doc.paragraphs[i].text.strip())
        if m:
            stem_idxs.append(i)
    if len(stem_idxs) != 15:
        raise ValueError(f"Expected 15 MCQ stems in template, found {len(stem_idxs)}.")
    for qi, si in enumerate(stem_idxs, start=1):
        end = stem_idxs[qi] if qi < len(stem_idxs) else b_heading
        block_lines = content["mcq_blocks"].get(qi)
        if block_lines is None:
            raise ValueError(f"Missing MCQ block content for Q{qi}.")
        _set_paragraph_block_exact(doc, si, end, block_lines)

    # Matching table is TABLE 1 in template.
    match_table = doc.tables[1]
    for r, (left, desc) in enumerate(content["matching_rows"]):
        _set_table_cell_text(match_table, r, 0, left)
        _set_table_cell_text(match_table, r, 2, desc)

    # Replace section C/D paragraphs block: from first "丙部 – 程式輸出題" to "~全卷完~"
    c_start = _find_paragraph_index(doc, "丙部 – 程式輸出題", start=0)
    end_marker = _find_paragraph_index(doc, "~全卷完~", start=c_start)
    _set_paragraph_block_exact(doc, c_start, end_marker + 1, content["output_block"])

    # Fill-in tables are TABLE 2, 3, 4 (indexing from 0): they each have 2 rows x 2 cols
    for ti, (code, out) in enumerate(content["fill_tables"], start=2):
        t = doc.tables[ti]
        _set_table_cell_text(t, 1, 0, code)
        _set_table_cell_text(t, 1, 1, out)

    # Answer sheet / headers: replace 上學期考試 -> 下學期考試 across all paragraphs (and keep year text).
    for p in doc.paragraphs:
        if "上學期考試" in p.text:
            _replace_in_paragraph_runs(p, "上學期考試", meta.term_text)

    # Replace the printed date line on answer sheet (two occurrences in paragraphs).
    # If template has fixed 15/01/2026, overwrite to meta.date_text if it's numeric; otherwise blank it.
    # We keep the same "日期：" label.
    for p in doc.paragraphs:
        if "日期：" in p.text and "時間：" in p.text:
            # Example: 班別： (\t)\t日期：15/01/2026
            parts = p.text.split("日期：")
            if len(parts) == 2:
                _set_paragraph_text_distribute(
                    p,
                    parts[0] + "日期：" + (meta.date_text if meta.date_text else "__________"),
                )

    # MCQ answer key paragraph (the line with 15 letters)
    ans_idx = _find_paragraph_index(doc, "DDBCB", start=0)
    _set_paragraph_text_distribute(doc.paragraphs[ans_idx], content["mcq_answers"])

    # Matching answer row table (TABLE 11) is 1 row x 5 cols
    match_ans_table = doc.tables[11]
    for i, txt in enumerate(content["matching_answers"]):
        _set_table_cell_text(match_ans_table, 0, i, txt)

    # Fill-in marking scheme tables (TABLE 14-16) are 3 rows x 3 cols
    t14, t15, t16 = doc.tables[14], doc.tables[15], doc.tables[16]

    _set_table_cell_text(t14, 0, 1, f"(i)\t{content['fill_answers']['fill1']['(i)']}")
    _set_table_cell_text(t14, 1, 1, f"(ii)\t{content['fill_answers']['fill1']['(ii)']}")
    _set_table_cell_text(t14, 2, 1, "(iii)\t(不適用)")

    _set_table_cell_text(t15, 0, 1, f"(i)\t{content['fill_answers']['fill2']['(i)']}")
    _set_table_cell_text(t15, 1, 1, f"(ii)\t{content['fill_answers']['fill2']['(ii)']}")
    _set_table_cell_text(t15, 2, 1, "(iii)\t(不適用)")

    _set_table_cell_text(t16, 0, 1, f"(i)\t{content['fill_answers']['fill3']['(i)']}")
    _set_table_cell_text(t16, 1, 1, f"(ii)\t{content['fill_answers']['fill3']['(ii)']}")
    _set_table_cell_text(t16, 2, 1, "(iii)\t(不適用)")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Generate S3 CMP Term 2 exam docx from Term 1 template.")
    p.add_argument("--template", required=True, help="Path to Term 1 exam template .docx")
    p.add_argument("--output", required=True, help="Output .docx path")
    p.add_argument("--date", default="__________", help="Date text on cover and answer sheet")
    p.add_argument("--time", default="__________", help="Time text on cover")
    return p.parse_args()


def main() -> int:
    args = _parse_args()
    meta = ExamMeta(
        school_name="迦密聖道中學",
        year_text="2025 – 2026",
        term_text="下學期考試",
        subject_text="中三級 電腦認知",
        paper_text="試題簿",
        date_text=str(args.date),
        time_text=str(args.time),
        duration_text="30分鐘",
        pages_text="5頁",
        total_marks_text="50",
    )
    build_term2_exam_docx(
        template_path=Path(args.template),
        output_path=Path(args.output),
        meta=meta,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

