"""Fixed paragraph layout for 24_25 F5 ICT Exam02 template (body after cover)."""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

from docx_inplace import set_paragraph_text_distribute
from f5_ict_tables import clear_all_body_tables_before_write

if TYPE_CHECKING:
    from docx.document import Document

# Must match spec_mcq_render.MCQ_SPANS (6,6,6,6,6,10,8,10,6,6,6,6,10,6,10,10,6,6,6,6,10,10,9,13,12,7,6,14,6,14)
_MCQ_SPANS: tuple[int, ...] = (
    6, 6, 6, 6, 6, 10, 8, 10, 6, 6, 6, 6, 10, 6, 10, 10, 6, 6, 6, 6,
    10, 10, 9, 13, 13, 11, 8, 14, 6, 14,
)

# Half-open paragraph ranges on 24_25_S5_ICT_Exam02.docx (verified once from template).
F5_SECTION_HEADERS: frozenset[int] = frozenset({1, 311, 423})
F5_PART_B_RANGE = (313, 422)  # inclusive end
F5_PART_C_RANGE = (424, 624)  # inclusive end

# MCQ slots 1–30: (start, end_exclusive) — end matches template option block + padding.
F5_ICT_MCQ_BLOCKS: tuple[tuple[int, int], ...] = (
    (4, 10),
    (12, 18),
    (20, 26),
    (28, 34),
    (36, 42),
    (44, 54),
    (57, 65),
    (67, 77),
    (79, 85),
    (87, 93),
    (95, 101),
    (103, 109),
    (112, 122),
    (124, 130),
    (132, 142),
    (144, 154),
    (156, 162),
    (164, 170),
    (173, 179),
    (181, 187),
    (189, 199),
    (201, 211),
    (213, 222),
    (225, 238),
    (240, 253),
    (253, 264),
    (264, 272),
    (272, 286),
    (288, 294),
    (296, 310),
)

assert len(F5_ICT_MCQ_BLOCKS) == len(_MCQ_SPANS)
for (s, e), span in zip(F5_ICT_MCQ_BLOCKS, _MCQ_SPANS, strict=True):
    assert e - s == span, (s, e, span)


def prepare_f5_ict_template_body(doc: Document) -> None:
    """
    Reset deliverable body: drop spare template tables, blank cells, clear paragraphs.

    Keeps cover (table 0) and section headers (甲／乙／丙). Content is written after this.
    """
    clear_all_body_tables_before_write(doc)

    # Clear MCQ area (2–310), 乙部 content, 丙部 content; keep section header lines.
    for i in range(2, 311):
        if i not in F5_SECTION_HEADERS:
            set_paragraph_text_distribute(doc.paragraphs[i], "")
    for i in range(F5_PART_B_RANGE[0], F5_PART_B_RANGE[1] + 1):
        set_paragraph_text_distribute(doc.paragraphs[i], "")
    for i in range(F5_PART_C_RANGE[0], F5_PART_C_RANGE[1] + 1):
        set_paragraph_text_distribute(doc.paragraphs[i], "")


def _last_nonempty_paragraph(doc: Document, start: int, end: int) -> int:
    for i in range(end - 1, start - 1, -1):
        if doc.paragraphs[i].text.strip():
            return i
    return start - 1


def _first_nonempty_paragraph(doc: Document, start: int, end: int) -> int:
    for i in range(start, end):
        if doc.paragraphs[i].text.strip():
            return i
    return end


def _collapse_paragraph(paragraph, profile: dict | None = None) -> None:
    """Minimize vertical space for template padding / excess gap lines."""
    if profile is None:
        from template_profile import load_f5_ict_profile

        profile = load_f5_ict_profile()
    from paper_format.renderer.paragraph_write import write_mcq_line

    write_mcq_line(paragraph, "", profile, collapsed=True)
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(1)


def _normal_blank_paragraph(paragraph, profile: dict | None = None) -> None:
    if profile is None:
        from template_profile import load_f5_ict_profile

        profile = load_f5_ict_profile()
    from paper_format.renderer.paragraph_write import write_mcq_line

    write_mcq_line(paragraph, "", profile)


def normalize_mcq_vertical_gaps(
    doc: Document,
    blocks: tuple[tuple[int, int], ...] = F5_ICT_MCQ_BLOCKS,
    *,
    between_gap_lines: int = 2,
) -> None:
    """
    Between MCQ items: keep exactly ``between_gap_lines`` visible blank paragraphs.

    Template rows are wider than content; unused rows inside/between blocks are
    collapsed to ~1pt line height so Q6→Q7 does not show 7 empty lines.
    """
    if between_gap_lines < 1:
        raise ValueError("between_gap_lines must be >= 1")

    from template_profile import load_f5_ict_profile

    profile = load_f5_ict_profile()

    for bi in range(len(blocks) - 1):
        start_i, end_i = blocks[bi]
        start_j, end_j = blocks[bi + 1]
        last_i = _last_nonempty_paragraph(doc, start_i, end_i)
        first_j = _first_nonempty_paragraph(doc, start_j, end_j)
        gap_indices = list(range(last_i + 1, first_j))
        if not gap_indices:
            continue
        if len(gap_indices) <= between_gap_lines:
            for i in gap_indices:
                _normal_blank_paragraph(doc.paragraphs[i], profile)
            continue
        for i in gap_indices[:-between_gap_lines]:
            _collapse_paragraph(doc.paragraphs[i], profile)
        for i in gap_indices[-between_gap_lines:]:
            _normal_blank_paragraph(doc.paragraphs[i], profile)

    start, end = blocks[-1]
    last = _last_nonempty_paragraph(doc, start, end)
    for i in range(last + 1, end):
        _collapse_paragraph(doc.paragraphs[i], profile)
