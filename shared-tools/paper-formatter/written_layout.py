"""
Reusable layout helpers for DSE / F5 ICT short (乙部) and long (丙部) written questions.

Preserves template paragraph styles by writing through existing runs (see docx_inplace).
"""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx_inplace import set_paragraph_text_rich

if TYPE_CHECKING:
    from docx.document import Document


# Tab-only lines used as answer / diagram space in past-paper templates.
ANSWER_BLANK = "\t\t\t\t\t"
ANSWER_BLANK_LONG = "\t\t\t\t\t\t\t"
DIAGRAM_BLANK = "\t\t"


def marks(m: int, *, spaced: bool = True) -> str:
    """Suffix like ``(5 分)`` or ``(5分)``."""
    return f"({m} 分)" if spaced else f"({m}分)"


def subpart(
    label: str,
    text: str,
    points: int | None = None,
    *,
    depth: int = 1,
    spaced_marks: bool = True,
) -> str:
    """
    Format a sub-question line: ``\\t(a)\\t...\\t(N 分)``.

    ``depth`` controls leading tabs (1 → ``\\t(a)``, 2 → ``\\t\\t(i)``).
    """
    lead = "\t" * depth
    suffix = f"\t{marks(points, spaced=spaced_marks)}" if points is not None else ""
    return f"{lead}({label})\t{text}{suffix}"


def topic_header(number: str, title: str, points: int) -> str:
    """Numbered section topic for List Paragraph slots, e.g. ``2.\\t數據控制\\t(10 分)``."""
    return f"{number}\t{title}\t{marks(points)}"


def stem(text: str, points: int | None = None, *, spaced_marks: bool = False, depth: int = 1) -> str:
    """Single-line stem with leading tab(s) (ERD / instruction lines)."""
    suffix = f"\t{marks(points, spaced=spaced_marks)}" if points is not None else ""
    return f"{'\t' * depth}{text}{suffix}"


def code_line(text: str, *, depth: int = 2) -> str:
    """Pseudocode / algorithm line."""
    return f"{'\t' * depth}{text}"


def sql_line(text: str, *, depth: int = 3) -> str:
    """Indented SQL or schema line."""
    return f"{'\t' * depth}{text}"


def blank_lines(count: int) -> list[str]:
    return [""] * count


def replace_span(
    doc: Document,
    start: int,
    end_inclusive: int,
    lines: list[str],
    *,
    profile: dict | None = None,
) -> None:
    """
    Overwrite ``doc.paragraphs[start:end_inclusive]`` using template role profiles.

    Raises if line count does not match the paragraph span exactly.
    """
    span = end_inclusive - start + 1
    if len(lines) != span:
        raise RuntimeError(
            f"replace_span {start}-{end_inclusive} need {span} lines, got {len(lines)}"
        )
    if profile is None:
        from template_profile import load_f5_ict_profile

        profile = load_f5_ict_profile()
    from paper_format.renderer.paragraph_write import write_written_line

    for k, text in enumerate(lines):
        write_written_line(doc.paragraphs[start + k], text, profile)


def set_paragraph_block(doc: Document, start: int, end_exclusive: int, lines: list[str]) -> None:
    """Overwrite a half-open paragraph range; extra slots are blanked."""
    span = max(0, end_exclusive - start)
    if len(lines) > span:
        raise ValueError(
            f"Replacement has {len(lines)} lines but only {span} paragraphs "
            f"(range {start}:{end_exclusive})."
        )
    for i in range(span):
        set_paragraph_text_rich(
            doc.paragraphs[start + i],
            lines[i] if i < len(lines) else "",
        )
