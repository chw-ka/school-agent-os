"""Insert 乙／丙 body tables after render — slot-driven, aligned with spec text."""

from __future__ import annotations

import re
from typing import Any

from docx import Document

from docx_inplace import set_paragraph_text_rich
from f5_ict_table_style import PROTOTYPE_INDEX_BY_KIND, SLOT_TABLE_KIND, apply_style_for_slot
from f5_ict_tables import set_cell_text
from f5_ict_template_table import insert_table_clone_after
from written_slot_ranges import WRITTEN_SLOT_PARAGRAPHS

def format_table_label(name: str, *, written: bool = False) -> str:
    """Table caption before grid — tab-indented like MCQ (1)(2)(3) / 乙部 body."""
    if written:
        return f"\t{name}\t"
    return f"\t\t{name}\t"


# Slots that always get a table when rendered (school Exam02 layout).
_ALWAYS_TABLE: frozenset[str] = frozenset(
    {"b-01", "b-02", "b-03", "b-05", "b-06", "c-05", "c-06"}
)

_TABLE_HINTS = (
    "試算表",
    "下表",
    "資料表",
    "追蹤",
    "堆疊",
    "Stack",
    "FACILITY",
    "RESERVE",
    "BOOKING",
    "ORDER",
)


def _fill_grid(table, grid: list[list[str]]) -> None:
    for r in range(len(table.rows)):
        row = grid[r] if r < len(grid) else []
        cells = table.rows[r].cells
        for c in range(len(cells)):
            val = row[c] if c < len(row) else ""
            cell = cells[c]
            if cell.paragraphs:
                set_paragraph_text_rich(cell.paragraphs[0], val)
            else:
                cell.text = val


def _para(doc: Document, para_index: int):
    """``para_index`` matches ``doc.paragraphs`` / ``replace_span`` (0-based)."""
    if para_index < 0 or para_index >= len(doc.paragraphs):
        raise IndexError(f"paragraph {para_index} out of range (0–{len(doc.paragraphs) - 1})")
    return doc.paragraphs[para_index]


def _insert_grid_after_paragraph(para, grid: list[list[str]], *, slot_id: str):
    if not grid:
        return
    rows, cols = len(grid), max(len(r) for r in grid)
    kind = SLOT_TABLE_KIND.get(slot_id, "data")
    if kind == "spreadsheet":
        from docx_inplace import insert_table_after_paragraph

        table = insert_table_after_paragraph(para, rows, cols)
        from f5_ict_table_style import apply_spreadsheet_style

        apply_spreadsheet_style(table)
    elif slot_id in ("b-03", "b-06"):
        from docx_inplace import insert_table_after_paragraph

        table = insert_table_after_paragraph(para, rows, cols)
        apply_style_for_slot(table, slot_id)
    else:
        proto_idx = PROTOTYPE_INDEX_BY_KIND.get(kind)
        if proto_idx is not None:
            table = insert_table_clone_after(para, proto_idx, rows=rows, cols=cols)
        else:
            from docx_inplace import insert_table_after_paragraph

            table = insert_table_after_paragraph(para, rows, cols)
            apply_style_for_slot(table, slot_id)
    _fill_grid(table, grid)
    return table


def _insert_grid(doc: Document, after_para: int, grid: list[list[str]], *, slot_id: str):
    _insert_grid_after_paragraph(_para(doc, after_para), grid, slot_id=slot_id)


def _paragraph_after_element(doc: Document, element) -> Any | None:
    nxt = element.getnext()
    while nxt is not None:
        if nxt.tag.endswith("}p"):
            for p in doc.paragraphs:
                if p._element is nxt:
                    return p
        nxt = nxt.getnext()
    return None


def insert_named_tables_after_stem(
    doc: Document,
    stem_para_index: int,
    blocks: list[tuple[str, list[list[str]]]],
    *,
    slot_id: str,
    written: bool = False,
) -> int:
    """Insert TABLE_NAME paragraph then table for each block (stem → name → table → …)."""
    from docx_inplace import set_paragraph_text_distribute
    from exam_generator import _insert_paragraph_after

    anchor_p = doc.paragraphs[stem_para_index]
    count = 0
    for name, grid in blocks:
        if not grid:
            continue
        label_p = _insert_paragraph_after(anchor_p, "")
        set_paragraph_text_distribute(label_p, format_table_label(name, written=written))
        table = _insert_grid_after_paragraph(label_p, grid, slot_id=slot_id)
        if table is not None:
            nxt = _paragraph_after_element(doc, table._tbl)
            anchor_p = nxt if nxt is not None else label_p
        else:
            anchor_p = label_p
        count += 1
    return count


def _slot_range(slot_id: str) -> tuple[int, int]:
    start, end, _ = WRITTEN_SLOT_PARAGRAPHS[slot_id]
    return start, end


def _first_nonempty_in_slot(doc: Document, slot_id: str) -> int:
    start, end = _slot_range(slot_id)
    for i in range(start, end + 1):
        if doc.paragraphs[i].text.strip():
            return i
    return start


def _anchor_after_scenario(doc: Document, slot_id: str) -> int:
    """First stem line of slot — insert table immediately after (before (a))."""
    return _first_nonempty_in_slot(doc, slot_id)


def _anchor_before_subpart(doc: Document, slot_id: str, label: str) -> int:
    start, end = _slot_range(slot_id)
    pat = re.compile(rf"^\s*\({re.escape(label)}\)")
    for i in range(start, end + 1):
        if pat.match(doc.paragraphs[i].text.strip()):
            return max(start, i - 1)
    return _anchor_after_scenario(doc, slot_id)


def _needs_table(slot_id: str, pick: dict[str, Any]) -> bool:
    if slot_id in _ALWAYS_TABLE:
        return True
    text = pick.get("text") or ""
    return any(h in text for h in _TABLE_HINTS)


def _parse_columns(text: str) -> list[tuple[str, str]]:
    cols: list[tuple[str, str]] = []
    for m in re.finditer(r"欄\s*([A-Z])\s*[=＝]\s*([^、。\n]+)", text):
        cols.append((m.group(1), m.group(2).strip()))
    for m in re.finditer(r"欄\s*([A-Z])\s*為\s*([^、。\n]+)", text):
        cols.append((m.group(1), m.group(2).strip()))
    return cols


def _parse_table_name(text: str) -> str:
    m = re.search(r"資料表\s+(\w+)", text)
    if m:
        return m.group(1)
    m = re.search(r"表\s+(\w+)\s*\(", text)
    if m:
        return m.group(1)
    return "BOOKING"


def build_excel_compact_grid(
    columns: list[tuple[str, str]],
    data_rows: list[list[str]],
) -> list[list[str]]:
    """Excel-style grid with exactly len(columns) data columns (no empty padding)."""
    n = len(columns)
    if n == 0:
        return []
    letters = [c[0] for c in columns]
    labels = [c[1] for c in columns]
    grid: list[list[str]] = [[""] + letters, ["1"] + labels]
    for i, row in enumerate(data_rows, start=2):
        grid.append([str(i)] + (list(row) + [""] * n)[:n])
    return grid


def build_excel_like_grid(
    columns: list[tuple[str, str]],
    data_rows: list[list[str]],
    *,
    min_cols: int = 6,
    pad_to: int | None = None,
) -> list[list[str]]:
    """
    Excel-style corner + column letters + field labels + numbered data rows.

    Matches 24_25 table 3 layout:
      r0: [empty] | A | B | C | ...
      r1: [1]     | field labels
      r2: [2]     | data ...
    """
    letters = [c[0] for c in columns]
    labels = [c[1] for c in columns]
    n = pad_to or max(len(letters), min_cols)
    while len(letters) < n:
        ch = chr(ord("A") + len(letters))
        letters.append(ch)
        labels.append("")
    letters = letters[:n]
    labels = labels[:n]

    grid: list[list[str]] = [[""] + letters, ["1"] + labels[:n]]
    for i, row in enumerate(data_rows, start=2):
        padded = (list(row) + [""] * n)[:n]
        grid.append([str(i)] + padded)
    return grid


def _parse_sheet_name(text: str) -> str:
    m = re.search(r"試算表[「\"](\w+)[」\"]", text)
    if m:
        return m.group(1)
    m = re.search(r"「(\w+)」", text)
    return m.group(1) if m else "Order"


def _build_b01_order_grid(_pick: dict[str, Any]) -> list[list[str]]:
    cols = [
        ("A", "日期"),
        ("B", "商品"),
        ("C", "單價"),
        ("D", "數量"),
        ("E", "會員"),
        ("F", "總價"),
    ]
    return build_excel_compact_grid(
        cols,
        [
            ["2026-04-10", "明信片", "15", "8", "Y", ""],
            ["2026-04-11", "布袋", "45", "3", "N", ""],
            ["2026-04-12", "徽章", "20", "6", "Y", ""],
        ],
    )


def _build_b01_price_grid(_pick: dict[str, Any]) -> list[list[str]]:
    return build_excel_compact_grid(
        [("H", "商品"), ("I", "參考單價")],
        [("明信片", "15"), ("布袋", "45"), ("徽章", "20")],
    )


def _build_b01_spreadsheet(pick: dict[str, Any]) -> list[list[str]]:
    """Legacy single-grid API — prefer named Order + PriceRef via table_grids_for_pick."""
    return _build_b01_order_grid(pick)


def _build_b02_mini_sheet(_pick: dict[str, Any]) -> list[list[str]]:
    """DSE-style contrast: A = 文書級；B = 多工／影像級（規格差異明顯）。"""
    return [
        ["硬件規格", "電腦 A", "電腦 B"],
        ["CPU", "Intel Core i3（3.3 GHz）", "Intel Core i7（4.2 GHz）"],
        ["RAM", "8 GB", "16 GB"],
        ["主儲存", "256 GB 硬碟（HDD）", "512 GB 固態硬碟（SSD）"],
        ["顯示卡", "內置顯示", "NVIDIA 獨立顯卡（4 GB）"],
        ["螢幕", "21 吋", "27 吋"],
    ]


def _build_b03_media_spec(pick: dict[str, Any]) -> list[list[str]]:
    text = pick.get("text") or ""
    m = re.search(r"(\d+)\s*[×x]\s*(\d+)\s*像素.*?(\d+)\s*bit", text, re.I | re.S)
    n_m = re.search(r"(\d+)\s*張", text)
    n_img = n_m.group(1) if n_m else "40"
    if m:
        w, h, bits = m.group(1), m.group(2), m.group(3)
        body = (
            f"相片規格（未壓縮 BMP）\n"
            f"解像度：{w} × {h} 像素\n"
            f"色彩：{bits} bit 真彩色\n"
            f"數量：{n_img} 張"
        )
    else:
        body = "相片規格（未壓縮 BMP）\n1600 × 1200 像素\n24 bit 真彩色\n40 張"
    return [[body]]


def _build_b04_trace(_pick: dict[str, Any]) -> list[list[str]]:
    return [["步驟", "i", "found"], ["1", "", ""], ["2", "", ""], ["3", "", ""]]


def _build_b05_sql_sample(pick: dict[str, Any]) -> list[list[str]]:
    return [
        ["TID", "Item", "Qty", "ADate"],
        ["T001", "明信片", "8", "2026-04-10"],
        ["T002", "布袋", "3", "2026-04-11"],
        ["T003", "徽章", "6", "2026-04-12"],
    ]


def _build_b06_hardware(_pick: dict[str, Any]) -> list[list[str]]:
  # 學生於 (a) 填寫 CPU 例子；其餘欄提供可比較的參考規格
    return [
        ["部件", "候選機 A", "候選機 B"],
        ["CPU", "（由考生填寫）", "（由考生填寫）"],
        ["RAM", "8 GB", "16 GB"],
        ["儲存（系統碟）", "256 GB SSD", "512 GB SSD"],
        ["顯示器", '15.6" FHD', '16" FHD'],
    ]


def _build_c03_orders(_pick: dict[str, Any]) -> list[list[str]]:
    return [
        ["PID", "SID", "AMT"],
        ["P01", "S01", "150"],
        ["P02", "S02", "90"],
        ["P03", "S03", "0"],
        ["P04", "S04", "60"],
    ]


def _build_c05_member(_pick: dict[str, Any]) -> list[list[str]]:
    return [
        ["MID", "MName"],
        ["M10001", "陳大文"],
        ["M10002", "李美玲"],
        ["M10003", "王志強"],
    ]


def _build_c05_facility(_pick: dict[str, Any]) -> list[list[str]]:
    return [
        ["FID", "FName"],
        ["F01", "舞蹈室"],
        ["F02", "活動室"],
        ["F03", "會議室"],
    ]


def _build_c05_reserve(_pick: dict[str, Any]) -> list[list[str]]:
    return [
        ["RID", "MEMID", "FID", "RDATE"],
        ["R260501", "M10001", "F01", "2026-07-12"],
        ["R260502", "M10001", "F02", "2026-07-18"],
        ["R260503", "M10002", "F01", "2026-07-12"],
        ["R260504", "M10003", "F03", "2026-08-01"],
    ]


def _build_c06_grid(_pick: dict[str, Any]) -> list[list[str]]:
    return [
        ["", "col1", "col2", "col3", "col4", "col5"],
        ["row1", "0", "0", "1", "0", "0"],
        ["row2", "0", "0", "1", "0", "0"],
        ["row3", "0", "0", "1", "0", "0"],
        ["row4", "0", "0", "0", "0", "0"],
        ["row5", "0", "0", "0", "0", "0"],
    ]


def _build_c08_stack(_pick: dict[str, Any]) -> list[list[str]]:
    return [
        ["步驟", "操作", "頂端", "Stack（底→頂）"],
        ["0", "初始", "—", "∅"],
        ["1", "PUSH 3", "3", "3"],
        ["2", "PUSH 7", "7", "3, 7"],
        ["3", "POP", "3", "3"],
    ]


def table_grids_for_pick(slot_id: str, pick: dict[str, Any]) -> list[tuple[str, list[list[str]]]]:
    """Named table grids that DOCX render would insert (no Document required)."""
    if not _needs_table(slot_id, pick):
        return []
    out: list[tuple[str, list[list[str]]]] = []
    if slot_id == "b-01":
        sheet = _parse_sheet_name(pick.get("text") or "")
        out.append((sheet, _build_b01_order_grid(pick)))
        out.append(("PriceRef", _build_b01_price_grid(pick)))
    elif slot_id == "b-02":
        out.append(("部件比較", _build_b02_mini_sheet(pick)))
    elif slot_id == "b-03":
        out.append(("相片規格", _build_b03_media_spec(pick)))
    elif slot_id == "b-04":
        out.append(("追蹤表", _build_b04_trace(pick)))
    elif slot_id == "b-05":
        out.append(("TRANSACTION", _build_b05_sql_sample(pick)))
    elif slot_id == "b-06":
        out.append(("硬件規格", _build_b06_hardware(pick)))
    elif slot_id == "c-03":
        out.append(("ORDER", _build_c03_orders(pick)))
    elif slot_id == "c-05":
        out.extend(
            [
                ("MEMBER", _build_c05_member(pick)),
                ("FACILITY", _build_c05_facility(pick)),
                ("RESERVE", _build_c05_reserve(pick)),
            ]
        )
    elif slot_id == "c-06":
        out.append(("Grid", _build_c06_grid(pick)))
    elif slot_id == "c-08":
        out.append(("Stack", _build_c08_stack(pick)))
    return out


def _anchor_for_slot(doc: Document, slot_id: str, pick: dict[str, Any]) -> int:
    if slot_id == "b-04":
        return _anchor_before_subpart(doc, slot_id, "a")
    return _anchor_after_scenario(doc, slot_id)


def _resolve_block_jobs(
    doc: Document, picks: dict[str, dict]
) -> list[tuple[int, str, list[tuple[str, list[list[str]]]]]]:
    """(stem_anchor, slot_id, [(table_name, grid), …]) sorted by anchor descending."""
    blocks: list[tuple[int, str, list[tuple[str, list[list[str]]]]]] = []
    for slot_id, pick in picks.items():
        if not _needs_table(slot_id, pick):
            continue
        grids = table_grids_for_pick(slot_id, pick)
        if not grids:
            continue
        blocks.append((_anchor_for_slot(doc, slot_id, pick), slot_id, grids))
    blocks.sort(key=lambda x: x[0], reverse=True)
    return blocks


def apply_written_tables_from_picks(doc: Document, picks: dict[str, dict]) -> int:
    """Insert tables for written slots; returns count of tables inserted."""
    if not picks:
        return 0
    count = 0
    for anchor, slot_id, named_grids in _resolve_block_jobs(doc, picks):
        use_named = len(named_grids) > 1 or (
            len(named_grids) == 1 and bool(named_grids[0][0])
        )
        if use_named:
            count += insert_named_tables_after_stem(
                doc, anchor, named_grids, slot_id=slot_id, written=True
            )
        elif named_grids:
            _insert_grid(doc, anchor, named_grids[0][1], slot_id=slot_id)
            count += 1
    return count
