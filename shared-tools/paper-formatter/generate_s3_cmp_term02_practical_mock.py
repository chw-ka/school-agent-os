#!/usr/bin/env python3
"""Generate 25-26 S3 CMP Term 2 Practical Mock (2× last-year exam tasks + vibe coding)."""

from __future__ import annotations

import argparse
import json
import shutil
from pathlib import Path

from docx import Document

from docx_inplace import set_paragraph_text_distribute

REPO = Path(__file__).resolve().parents[2]
TEMPLATE = (
    REPO
    / "Subjects/S3-CMP/past-papers/2024-2025/Term 02/PracticalAssessment"
    / "2425_S3_CMP_Term02_PracticalAssessment.docx"
)
EXAM_2425 = (
    REPO
    / "Subjects/S3-CMP/past-papers/2024-2025/Term 02/PracticalExam/2425_S3_CMP_Term02_PracticalExam"
)
HAAR_SRC = (
    REPO
    / "Subjects/S3-CMP/past-papers/2024-2025/Term 02/PracticalAssessment"
    / "haarcascade_frontalface_default.xml"
)
DEFAULT_OUT = REPO / "Subjects/S3-CMP/past-papers/2025-2026/Term 02/PracticalMock"


def _set(doc: Document, idx: int, text: str) -> None:
    set_paragraph_text_distribute(doc.paragraphs[idx], text)


def _blank_range(doc: Document, start: int, end_exclusive: int) -> None:
    for i in range(start, end_exclusive):
        _set(doc, i, "")


def _build_task1_lines() -> list[str]:
    return [
        "任務 1\t(15 分)【禁止使用 Gemini 或其他 AI】",
        "完成 Python 程式《文字轉語音》（改編自 24-25 實習試任務三）：",
        "(a)\t下載 task1_2526.py，用 VS Code 開啟。",
        "(b)\t匯入 gTTS、playsound、os。",
        "(c)\t請使用者輸入想轉成語音的句子（英文）。",
        "(d)\tgTTS 產生語音、儲存 output.mp3、播放，然後用 os.remove 刪除該檔。",
        "(e)\t提交 task1_2526.py。",
        "(f)\t例子（粗體為輸入）：Enter the text you want to convert to speech: Hello",
        "圖 1",
    ]


def _build_task2_lines() -> list[str]:
    return [
        "任務 2\t(15 分)【禁止使用 Gemini 或其他 AI】",
        "完成 Python 程式《影片人面偵測》（改編自 24-25 實習試任務四）：",
        "(a)\t下載 task2_2526.py、video.mp4、haarcascade_frontalface_default.xml。",
        "(b)\t匯入 cv2；載入人臉模型；開啟 video.mp4。",
        "(c)\twhile 迴圈讀取每一格；若讀取失敗則 break。",
        "(d)\t轉灰階、detectMultiScale；每張臉畫藍色方框及綠色圓點（中心）。",
        "(e)\timshow 顯示；waitKey(1)。",
        "(f)\t提交 task2_2526.py。",
        "圖 2",
    ]


def _build_task3_lines() -> list[str]:
    return [
        "任務 3\t(20 分)【必須使用 Gemini】",
        "Vibe Coding《遊戲頒獎典禮》：讀取 game_scores.json，製作有趣的 party_show.txt。",
        "(a)\t下載 game_scores.json、game_scores_sample.json、party_show_sample.txt、",
        "vibe_spec.md、task3_2526.py、gemini_prompts.txt。",
        "(b)\t先用 sample 檔向 Gemini 試寫程式；再處理正式 game_scores.json。",
        "(c)\t在 gemini_prompts.txt 記錄至少 2 則你輸入的完整提示（初稿 + 修改）。",
        "(d)\t執行 task3_2526.py 產生 party_show.txt（內容須有趣、可讀）。",
        "(e)\t提交 task3_2526.py、party_show.txt、gemini_prompts.txt。",
        "(f)\t評分：提示清晰度、能否迭代、輸出是否符規格及有趣程度。",
        "圖 3",
    ]


def _apply_cover_zh(doc: Document) -> None:
    _set(doc, 0, "迦密聖道中學")
    _set(doc, 1, "2025-2026 下學期")
    _set(doc, 2, "中三級 電腦認知科")
    _set(doc, 3, "模擬實習試")
    _set(doc, 4, "時間: 45 分鐘")


def _apply_instructions(doc: Document) -> None:
    _set(doc, 8, "請在方格內填寫姓名、班級和學號。")
    _set(doc, 9, "請依照以下說明完成考試任務。")
    _set(doc, 10, "如果你未能上傳檔案，將不予評分。")
    _set(
        doc,
        12,
        "重要：任務 1、2 必須獨立完成，禁止使用 Gemini、ChatGPT 或其他 AI 代寫或解釋程式。"
        "違者該題零分。任務 3 必須使用 Gemini 進行 Vibe Coding，並提交提示紀錄。",
    )
    _set(
        doc,
        13,
        "登入 Teams「S3 下學期實習模擬試：任務一／任務二／任務三」，下載所有檔案到同一資料夾。",
    )
    _set(
        doc,
        14,
        "提交：task1_2526.py、task2_2526.py、task3_2526.py、party_show.txt、gemini_prompts.txt。",
    )
    _set(doc, 15, "（學校會於考試時監察；請誠實應考。）")


TASK1_START, TASK1_END = 18, 31
TASK2_START, TASK2_END = 31, 48
TASK3_START, TASK3_END = 48, 65


def _apply_task_block(doc: Document, start: int, end: int, lines: list[str]) -> None:
    span = end - start
    if len(lines) > span:
        raise ValueError(f"Task needs {len(lines)} lines but template only has {span} ({start}:{end})")
    for i, line in enumerate(lines):
        _set(doc, start + i, line)
    _blank_range(doc, start + len(lines), end)


def _apply_tasks(doc: Document) -> None:
    _apply_task_block(doc, TASK1_START, TASK1_END, _build_task1_lines())
    _apply_task_block(doc, TASK2_START, TASK2_END, _build_task2_lines())
    _apply_task_block(doc, TASK3_START, TASK3_END, _build_task3_lines())
    _blank_range(doc, 65, 81)
    _set(doc, 81, "全卷完")


def _apply_name_table(doc: Document) -> None:
    labels = ["姓名", "班別", "學號", "日期", "分數"]
    for i, label in enumerate(labels):
        doc.tables[0].cell(i, 0).paragraphs[0].text = label


def _set_cell(table, r: int, c: int, text: str) -> None:
    table.cell(r, c).paragraphs[0].text = text


def _apply_appendix_zh(doc: Document) -> None:
    for i, p in enumerate(doc.paragraphs):
        if "Appendix" in p.text or "附錄" in p.text:
            _set(doc, i, "附錄：函數庫及 API 參考表")
            break

    t1 = doc.tables[1]
    _set_cell(t1, 0, 0, "函數庫")
    _set_cell(t1, 0, 1, "pip 安裝指令")
    _set_cell(t1, 0, 2, "匯入語句範例")
    rows1 = [
        ("gTTS", "gTTS", "from gtts import gTTS"),
        ("Playsound", "playsound==1.2.2", "from playsound import playsound"),
        ("os", "（內建）", "import os"),
        ("OpenCV", "opencv-contrib-python", "import cv2"),
    ]
    for i, (a, b, c) in enumerate(rows1, start=1):
        if i < len(t1.rows):
            _set_cell(t1, i, 0, a)
            _set_cell(t1, i, 1, b)
            _set_cell(t1, i, 2, c)
    for i in range(len(rows1) + 1, len(t1.rows)):
        for j in range(len(t1.columns)):
            _set_cell(t1, i, j, "")

    t2 = doc.tables[2]
    _set_cell(t2, 0, 0, "函數庫")
    _set_cell(t2, 0, 1, "說明")
    _set_cell(t2, 0, 2, "函數範例")
    rows2 = [
        ("gTTS", "建立 gTTS 物件", 'audio = gTTS(text="I Love You", lang="en")'),
        ("gTTS", "儲存語音為檔案", 'audio.save("audio.mp3")'),
        ("Playsound", "播放音訊檔", 'playsound("audio.mp3")'),
        ("os", "刪除資料夾內的檔案", 'os.remove("output.mp3")'),
        ("OpenCV", "讀取圖片檔", 'image = cv2.imread("face.jpg")'),
        ("OpenCV", "開啟影片或鏡頭", 'cap = cv2.VideoCapture("video.mp4")'),
        ("OpenCV", "讀取影片中的一格畫面", "isTrue, frame = cap.read()"),
        ("OpenCV", "轉為灰階", "cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)"),
        (
            "OpenCV",
            "在人臉周圍畫方框",
            "cv2.rectangle(frame, (x, y), (x + w, y + h), (255, 0, 0), 2)",
        ),
        ("OpenCV", "在視窗顯示畫面", "cv2.imshow('Video', frame)"),
        ("OpenCV", "等待按鍵（可設延遲）", "cv2.waitKey(1)"),
    ]
    for i, (a, b, c) in enumerate(rows2, start=1):
        if i < len(t2.rows):
            _set_cell(t2, i, 0, a)
            _set_cell(t2, i, 1, b)
            _set_cell(t2, i, 2, c)
    for i in range(len(rows2) + 1, len(t2.rows)):
        for j in range(len(t2.columns)):
            _set_cell(t2, i, j, "")


def build_docx(out_path: Path) -> None:
    doc = Document(TEMPLATE)
    _apply_name_table(doc)
    _apply_cover_zh(doc)
    _apply_instructions(doc)
    _apply_tasks(doc)
    _apply_appendix_zh(doc)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


# --- Task 1: 24-25 exam task 3 (TTS) ---
TASK1_STARTER = '''# (b) 匯入函數庫
from
from
import

text = input("Enter the text you want to convert to speech: ")

# (c) 文字轉語音、播放、刪除檔案
audio =
audio.

os.
'''

TASK1_ANSWER = '''from gtts import gTTS
from playsound import playsound
import os

text = input("Enter the text you want to convert to speech: ")

audio = gTTS(text=text, lang="en")
audio.save("output.mp3")
playsound("output.mp3")
os.remove("output.mp3")
'''

# --- Task 2: 24-25 exam task 4 (CV video) ---
TASK2_STARTER = '''# (b) 匯入函數庫
import

# (c) 載入人臉模型
face_cascade = cv.

# (d) 開啟影片
video = cv.

while True:
    # (e) 讀取每一格
    isTrue, frame =
    if
        break

    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)

    faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=5, minSize=(30, 30))

    for (x, y, w, h) in faces:
        cv2.rectangle(frame, (x, y), (x + w, y + h), (255, 0, 0), 2)
        center_x, center_y = x + w // 2, y + h // 2
        cv2.circle(frame, (center_x, center_y), 10, (0, 255, 0), 2)

    # (f) 顯示影片
    cv2.

    # (g) 等待按鍵
    cv2.
'''

TASK2_ANSWER = '''import cv2

face_cascade = cv2.CascadeClassifier("haarcascade_frontalface_default.xml")

video = cv2.VideoCapture("video.mp4")

while True:
    isTrue, frame = video.read()
    if not isTrue:
        break

    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)

    faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=5, minSize=(30, 30))

    for (x, y, w, h) in faces:
        cv2.rectangle(frame, (x, y), (x + w, y + h), (255, 0, 0), 2)
        center_x, center_y = x + w // 2, y + h // 2
        cv2.circle(frame, (center_x, center_y), 10, (0, 255, 0), 2)

    cv2.imshow("Video", frame)

    cv2.waitKey(1)
'''

# --- Task 3: Vibe coding ---
GAME_SCORES_SAMPLE = [
    {"name": "小明", "score": 95, "team": "紅隊"},
    {"name": "小華", "score": 72, "team": "藍隊"},
]

GAME_SCORES = [
    {"name": "陳大文", "score": 88, "team": "紅隊"},
    {"name": "李小美", "score": 95, "team": "藍隊"},
    {"name": "張志強", "score": 76, "team": "紅隊"},
    {"name": "黃詠琳", "score": 91, "team": "藍隊"},
]

PARTY_SHOW_SAMPLE = """🎉🎉🎉 遊戲頒獎典禮 🎉🎉🎉

🏆 最高分玩家：小明（95 分）— 紅隊
🥈 亞軍：小華（72 分）— 藍隊

📊 紅隊總分：95　｜　藍隊總分：72
🎊 恭喜紅隊勝出！

（由 Vibe Coding 自動生成）
"""

VIBE_SPEC = """# 任務三 Vibe Coding 規格 —《遊戲頒獎典禮》

## 必須使用 Gemini
本任務測試你的 **提示（Prompt）** 能力：你是總監，Gemini 是程式員。

## 輸入
讀取 `game_scores.json`（陣列），每個物件：
- `name`（字串）
- `score`（整數）
- `team`（字串，例如「紅隊」「藍隊」）

## 輸出
寫入 `party_show.txt`（純文字，utf-8），內容要 **有趣、可讀**，並包含：
1. 標題（例如含 🎉）
2. 最高分玩家姓名、分數、隊伍
3. 亞軍（第二高分）姓名、分數、隊伍
4. 兩隊總分比較及勝出隊伍
5. 結尾一句鼓勵說話

## 建議流程
1. 用 `game_scores_sample.json` 向 Gemini 試寫 `task3_2526.py`
2. 對照 `party_show_sample.txt` 修改提示
3. 改讀 `game_scores.json` 產生正式 `party_show.txt`
4. 把提示貼到 `gemini_prompts.txt`

## 提交
- task3_2526.py
- party_show.txt
- gemini_prompts.txt（至少 2 則完整提示）
"""

TASK3_STUB = '''"""
任務三（Vibe Coding）：讀取 game_scores.json，寫入 party_show.txt
請用 Gemini 協助完成；規格見 vibe_spec.md
"""

import json

# TODO: 用 Gemini 完成此程式


def main() -> None:
    with open("game_scores.json", encoding="utf-8") as f:
        players = json.load(f)
    # TODO: 產生有趣文字並寫入 party_show.txt
    raise NotImplementedError("Use Gemini to implement")


if __name__ == "__main__":
    main()
'''

TASK3_ANSWER = '''import json


def build_show(players: list[dict]) -> str:
    sorted_players = sorted(players, key=lambda p: p["score"], reverse=True)
    champion = sorted_players[0]
    runner_up = sorted_players[1]
    team_scores: dict[str, int] = {}
    for p in players:
        team_scores[p["team"]] = team_scores.get(p["team"], 0) + p["score"]
    winner_team = max(team_scores, key=team_scores.get)
    lines = [
        "🎉🎉🎉 遊戲頒獎典禮 🎉🎉🎉",
        "",
        f"🏆 最高分玩家：{champion['name']}（{champion['score']} 分）— {champion['team']}",
        f"🥈 亞軍：{runner_up['name']}（{runner_up['score']} 分）— {runner_up['team']}",
        "",
    ]
    for team, total in team_scores.items():
        lines.append(f"📊 {team}總分：{total}")
    lines.append(f"🎊 恭喜{winner_team}勝出！")
    lines.append("")
    lines.append("（由 Vibe Coding 自動生成）")
    return "\\n".join(lines)


def main() -> None:
    with open("game_scores.json", encoding="utf-8") as f:
        players = json.load(f)
    text = build_show(players)
    with open("party_show.txt", "w", encoding="utf-8") as f:
        f.write(text)
    print("Wrote party_show.txt")


if __name__ == "__main__":
    main()
'''

GEMINI_PROMPTS_TEMPLATE = """# Gemini 提示紀錄（任務三 — 必交）
# 請貼上你實際輸入的完整提示（至少 2 則）

## 提示 1（初稿）
（在此貼上）


## 提示 2（修改／優化後）
（在此貼上）

"""


def build_package(out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    build_docx(out_dir / "2526_S3_CMP_Term02_PracticalMock.docx")

    (out_dir / "task1_2526.py").write_text(TASK1_STARTER, encoding="utf-8")
    (out_dir / "task2_2526.py").write_text(TASK2_STARTER, encoding="utf-8")
    (out_dir / "task3_2526.py").write_text(TASK3_STUB, encoding="utf-8")
    (out_dir / "task1_2526_answer.py").write_text(TASK1_ANSWER, encoding="utf-8")
    (out_dir / "task2_2526_answer.py").write_text(TASK2_ANSWER, encoding="utf-8")
    (out_dir / "task3_2526_answer.py").write_text(TASK3_ANSWER, encoding="utf-8")

    (out_dir / "game_scores.json").write_text(
        json.dumps(GAME_SCORES, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (out_dir / "game_scores_sample.json").write_text(
        json.dumps(GAME_SCORES_SAMPLE, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (out_dir / "party_show_sample.txt").write_text(PARTY_SHOW_SAMPLE, encoding="utf-8")
    (out_dir / "vibe_spec.md").write_text(VIBE_SPEC, encoding="utf-8")
    (out_dir / "gemini_prompts.txt").write_text(GEMINI_PROMPTS_TEMPLATE, encoding="utf-8")

    shutil.copy2(EXAM_2425 / "video.mp4", out_dir / "video.mp4")
    shutil.copy2(HAAR_SRC, out_dir / "haarcascade_frontalface_default.xml")

    for stale in ("class_photo.jpg", "announcement.txt", "check_task3.py"):
        p = out_dir / stale
        if p.exists():
            p.unlink()

def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--out-dir", type=Path, default=DEFAULT_OUT)
    args = ap.parse_args()
    build_package(args.out_dir)
    print(f"Wrote package to {args.out_dir}")


if __name__ == "__main__":
    main()
