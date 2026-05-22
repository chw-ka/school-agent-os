#!/usr/bin/env python3
"""Generate 25-26 S2 CMP Term 2 written exam (甲 MCQ / 乙 matching / 丙 T-F / 丁 fill / 戊 SA)."""

from __future__ import annotations

import argparse
import os
import random
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from docx_inplace import ZhCoverPatch, apply_cmp_cover_zh, set_paragraph_text_distribute
from mcq_answer_keys import build_random_mcq_key
from s2_cmp_matching_layout import build_shuffled_matching
from s2_cmp_term2_exam_spec import (
    build_s2_cmp_term2_exam_spec,
    format_fill_rubric_blocks,
    format_matching_rubric_blocks,
    format_tf_rubric,
)
from s2_cmp_tf_fill_layout import build_shuffled_fill, build_shuffled_tf

_RNG = random.Random(20260523)

REPO = Path(__file__).resolve().parents[2]
DEFAULT_TEMPLATE = (
    REPO
    / "Subjects/S2-CMP/past-papers/2025-2026/Term 02/WrittenExam/25_26_S2_CMP_Term02_Exam.docx"
)
DEFAULT_OUT = DEFAULT_TEMPLATE

S2_MCQ_CORRECT_INDEX: tuple[int, ...] = (
    3,
    1,
    2,
    1,
    1,
    1,
    2,
    1,
    1,
    1,
    1,
    1,
    1,
    1,
    1,
    1,
    2,
    0,
    1,
    1,
)

# Paragraph spans in python-docx order (see s2_cmp_term2_exam_from_2425_template._mcq_spans).
MCQ_SPANS: dict[int, tuple[int, int]] = {
    1: (5, 16),
    2: (16, 19),
    3: (19, 22),
    4: (22, 25),
    5: (25, 28),
    6: (28, 30),
    7: (30, 32),
    8: (32, 35),
    9: (35, 41),
    10: (41, 44),
    11: (44, 47),
    12: (47, 54),
    13: (54, 61),
    14: (61, 67),
    15: (67, 74),
    16: (74, 82),
    17: (82, 90),
    18: (90, 97),
    19: (97, 104),
    20: (104, 107),
}

SHORT_ANSWER_Q = (
    "老師要求你用生成式 AI 整理一份 PDF 戶外考察通告的重點。"
    "請分別說明在 R-I-C-C-O 提示語中，「R（角色）」和「C（背景）」可以如何設定；"
    "並解釋為何還需要設定「C（限制）」。(5分)"
)

SHORT_ANSWER_RUBRIC = (
    "1.\t\tR：例如要求 AI 扮演「行政校務助理」；"
    "C（背景）：提供通告內容或說明這是戶外考察物資清單。"
    "C（限制）：例如只提取與學生有關的項目、不可編造日期、限制字數或輸出格式，"
    "以減少 AI 產生無關或錯誤資訊。"
)


@dataclass(frozen=True)
class Meta:
    school: str = "迦密聖道中學"
    year: str = "2025 – 2026"
    term: str = "下學期考試"
    level: str = "中二級 電腦認知"
    paper: str = "試題簿"
    date: str = "__________"
    time: str = "__________"
    time_limit: str = "30 分鐘"
    pages: str = "9 頁"
    total: str = "50"


def _find_idx(doc: Document, contains: str, *, start: int = 0) -> int:
    for i in range(start, len(doc.paragraphs)):
        if contains in doc.paragraphs[i].text:
            return i
    raise ValueError(f'missing: "{contains}"')


def _find_mcq_answer_line(doc: Document, *, start: int = 150) -> int:
    for i in range(start, len(doc.paragraphs)):
        compact = re.sub(r"\s+", "", doc.paragraphs[i].text.upper())
        if len(compact) >= 20 and all(c in "ABCD" for c in compact[:20]):
            return i
    raise ValueError("missing MCQ answer key line in answer sheet")


def _replace_block(doc: Document, start: int, end_excl: int, lines: list[str]) -> None:
    span = end_excl - start
    if len(lines) > span:
        raise ValueError(f"Need ≤{span} lines at {start}:{end_excl}, got {len(lines)}")
    for i in range(span):
        set_paragraph_text_distribute(
            doc.paragraphs[start + i], lines[i] if i < len(lines) else ""
        )


def _fix_option_tabs(block: list[str]) -> list[str]:
    out: list[str] = []
    for line in block:
        fixed = line
        for letter in "ABCD":
            fixed = fixed.replace(f"\n{letter}.\t", f"\n\t{letter}.\t")
            fixed = fixed.replace(f"\n{letter}.", f"\n\t{letter}.\t")
        out.append(fixed)
    return out


def _mcq_blocks() -> dict[int, list[str]]:
    from s2_cmp_term2_exam_from_2425_template import _mcq_blocks as raw

    blocks = raw()
    blocks[17] = [
        "17.\t以下哪一項最不符合 Teachable Machine「影像分類『訓練』」的正確做法？",
        "",
        "A.\t先為每個類別準備足夠且有代表性的影像樣本",
        "B.\t樣本越多越好（在同學能力範圍內盡量多元化）",
        "C.\t訓練前完全不收集樣本，直接按 Train Model",
        "D.\t完成訓練後應反覆測試辨識結果",
        "",
        "請根據『Teachable Machine』的筆記內容，回答第18至20題。",
    ]
    blocks[18] = [
        "18.\t同學要在 Teachable Machine 建立『向左移動／向右移動』兩種手勢控制遊戲角色，下列哪個組合最合理？",
        "",
        "A.\t建立兩個類別（標籤），分別代表向左／向右的手勢影像",
        "B.\t只需要下載 App Inventor，不必建立標籤",
        "C.\t只能用文字輸入指令完成訓練",
        "D.\t只能用音訊建立模型",
        "",
    ]
    blocks[19] = [
        "19.\t筆記強調訓練後要『充分測試』模型準確性，下列哪一項是最主要原因？",
        "",
        "A.\t增加電腦耗電量",
        "B.\t確認真實環境下手勢／光線變化下的辨識表現",
        "C.\t令網站版面更美觀",
        "D.\t令瀏覽器自動更新",
        "",
    ]
    blocks[20] = [
        "20.\t完成訓練後『記下可分享的連結』的主要用途是甚麼？\n\n"
        "A.\t取代瀏覽器搜尋功能\n"
        "B.\t稍後在 App Inventor 擴充／網頁專案載入模型（例如 TMIC）\n"
        "C.\t自動購買雲端硬碟\n"
        "D.\t下載老師的標準答案",
        "",
    ]
    return blocks


def _section_b_paragraphs() -> dict[str, list[str]]:
    return {
        "header": [
            "乙部 – 配對題 (10 分)",
            "提示語工程 (RICCO) 配對：請將正確的組件英文字母填入括號內。\t(5分)",
        ],
        "bank1": [
            "",
            "\t\t組件庫：",
            "A. 請使用「兩欄對比表格」呈現結果。",
            "B. 假設你是一位專業的「行政校務助理」。",
            "C. 絕對不要顯示中二級以外的活動，且不可編造任何日期。",
            "D. 學校即將舉辦戶外考察，同學們需要一份清晰的清單來準備物資。",
            "E. 請閱讀附件的 PDF 通告，並提取當中關於學生的注意事項。",
        ],
        "header2": [
            "功能與流程配對：根據描述選出最合適的術語字母。\t(5分)",
            "",
            "術語：",
            "A. 摘要",
            "B. 語意轉換",
            "C. 文生圖",
            "D. 圖生圖",
            "E. 積木式編程",
        ],
    }


def _apply_cover(doc: Document, meta: Meta) -> None:
    cover_cell = doc.tables[0].cell(0, 0)
    apply_cmp_cover_zh(
        cover_cell,
        ZhCoverPatch(
            school=meta.school,
            year_term=f"{meta.year} {meta.term}",
            level=meta.level,
            paper=meta.paper,
            date_line=f"\t日期:\t{meta.date}",
            time_line=f"\t時間:\t{meta.time}",
            duration_line=f"\t時限:\t{meta.time_limit}",
            pages_line=f"\t頁數:\t{meta.pages}",
            total_line=f"\t總分:\t{meta.total}",
        ),
    )


def _apply_matching_tables(
    doc: Document,
    rows1: list[tuple[str, str]],
    rows2: list[tuple[str, str]],
) -> None:
    for t, rows in ((doc.tables[1], rows1), (doc.tables[2], rows2)):
        for r, (a, b) in enumerate(rows):
            t.cell(r, 0).text = a
            t.cell(r, 1).text = b


def _apply_word_banks(doc: Document, bank_a: list[str], bank_b: list[str]) -> None:
    doc.tables[3].cell(0, 0).text = "\t" + "\t\t".join(bank_a)
    doc.tables[4].cell(0, 0).text = "\t" + "\t\t".join(bank_b)


def _apply_answer_sheet(
    doc: Document,
    meta: Meta,
    *,
    mcq_answers_display: str,
    matching_answers: list[str],
    tf_answers: str,
    fill_answers: list[list[str]],
) -> None:
    for p in doc.paragraphs:
        if p.text.strip() == "中二級 電腦認知":
            set_paragraph_text_distribute(p, meta.level)
        if "中二級 電腦認知科" in p.text:
            set_paragraph_text_distribute(
                p, f"{meta.year} {meta.level}科{meta.term} (評分準則)"
            )
    ans_mcq = _find_mcq_answer_line(doc, start=150)
    set_paragraph_text_distribute(doc.paragraphs[ans_mcq], mcq_answers_display)

    match_ans = format_matching_rubric_blocks(matching_answers)
    start = _find_idx(doc, "乙部： 配對題", start=150)
    for j, line in enumerate(match_ans):
        set_paragraph_text_distribute(doc.paragraphs[start + 1 + j], line)

    tf_ans = format_tf_rubric(tf_answers)
    tf_start = _find_idx(doc, "丙部： 是非題", start=150)
    for j, line in enumerate(tf_ans):
        set_paragraph_text_distribute(doc.paragraphs[tf_start + 1 + j], line)

    fill_ans = format_fill_rubric_blocks(fill_answers)
    fill_start = _find_idx(doc, "丁部：填充題", start=150)
    for j, line in enumerate(fill_ans):
        set_paragraph_text_distribute(doc.paragraphs[fill_start + 1 + j], line)

    for marker in ("戊部：問答題", "戊部：短答題"):
        try:
            sa_start = _find_idx(doc, marker, start=150)
            break
        except ValueError:
            continue
    else:
        raise ValueError("missing 戊部 answer rubric")
    set_paragraph_text_distribute(doc.paragraphs[sa_start + 1], SHORT_ANSWER_RUBRIC)


def generate(template: Path, output: Path, meta: Meta) -> tuple[str, list[str], str, list[list[str]]]:
    doc = Document(str(template))

    _apply_cover(doc, meta)

    raw_blocks = _mcq_blocks()
    blocks, mcq_key = build_random_mcq_key(
        raw_blocks, correct_indices=S2_MCQ_CORRECT_INDEX, rng=_RNG
    )
    mcq_display = " ".join(mcq_key[i : i + 5] for i in range(0, 20, 5))

    match_rows1, match_rows2, matching_keys = build_shuffled_matching(_RNG)
    tf_lines, tf_key = build_shuffled_tf(_RNG)
    fill_lines, fill_answers, bank_a, bank_b = build_shuffled_fill(_RNG)

    for q, (start, end) in MCQ_SPANS.items():
        span = end - start
        block = _fix_option_tabs(list(blocks[q]))
        if len(block) > span:
            raise ValueError(f"MCQ#{q} needs ≤{span} lines, got {len(block)}")
        if len(block) < span:
            block.extend([""] * (span - len(block)))
        _replace_block(doc, start, end, block)

    b = _section_b_paragraphs()
    _replace_block(doc, 107, 109, b["header"])
    _replace_block(doc, 109, 116, b["bank1"])
    _replace_block(doc, 117, 125, b["header2"])

    _replace_block(doc, 130, 135, tf_lines)
    _replace_block(doc, 137, 153, fill_lines)
    set_paragraph_text_distribute(doc.paragraphs[127], "丙部 – 是非題 (5分)")
    set_paragraph_text_distribute(doc.paragraphs[137], "丁部 – 填充題 (10 分)")
    set_paragraph_text_distribute(doc.paragraphs[154], "戊部 – 問答題 (5分)")
    set_paragraph_text_distribute(doc.paragraphs[155], SHORT_ANSWER_Q)

    _apply_matching_tables(doc, match_rows1, match_rows2)
    _apply_word_banks(doc, bank_a, bank_b)
    _apply_answer_sheet(
        doc,
        meta,
        mcq_answers_display=mcq_display,
        matching_answers=matching_keys,
        tf_answers=tf_key,
        fill_answers=fill_answers,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))

    spec_path = output.with_suffix(".spec.json")
    _save_spec_and_apply_footer(
        output,
        spec_path,
        mcq_key=mcq_key,
        matching_answers=matching_keys,
        tf_lines=tf_lines,
        tf_answers=tf_key,
        fill_answers=fill_answers,
        fill_word_banks=[bank_a, bank_b],
    )
    return mcq_key, matching_keys, tf_key, fill_answers


def _save_spec_and_apply_footer(
    docx_path: Path,
    spec_path: Path,
    *,
    mcq_key: str,
    matching_answers: list[str],
    tf_lines: list[str],
    tf_answers: str,
    fill_answers: list[list[str]],
    fill_word_banks: list[list[str]],
) -> None:
    for pkg in ("question-quality-check", "paper-quality-check"):
        p = Path(__file__).resolve().parents[1] / pkg
        if str(p) not in sys.path:
            sys.path.insert(0, str(p))
    from exam_spec import save_spec
    from footer import apply_footer_meta

    spec = build_s2_cmp_term2_exam_spec(
        mcq_answers=mcq_key,
        matching_answers=matching_answers,
        tf_lines=tf_lines,
        tf_answers=tf_answers,
        fill_answers=fill_answers,
        fill_word_banks=fill_word_banks,
        short_answer=SHORT_ANSWER_Q,
    )
    save_spec(spec_path, spec)
    footer = spec.get("meta", {}).get("footer")
    if footer:
        apply_footer_meta(docx_path, footer)


def _run_quality_check(docx_path: Path, spec_path: Path) -> int:
    root = Path(__file__).resolve().parents[1]
    env = {**os.environ, "PYTHONIOENCODING": "utf-8"}
    py = sys.executable
    q = subprocess.run(
        [
            py,
            str(root / "question-quality-check" / "check_docx.py"),
            "--candidate",
            str(docx_path.resolve()),
            "--candidate-spec",
            str(spec_path.resolve()),
            "--subject",
            "S2 CMP",
            "--years",
            "3",
        ],
        cwd=str(root),
        env=env,
    )
    p = subprocess.run(
        [
            py,
            str(root / "paper-quality-check" / "check_docx.py"),
            "--candidate",
            str(docx_path.resolve()),
            "--candidate-spec",
            str(spec_path.resolve()),
            "--skip-written",
        ],
        cwd=str(root),
        env=env,
    )
    return max(q.returncode, p.returncode)


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    ap.add_argument("--output", type=Path, default=DEFAULT_OUT)
    ap.add_argument("--date", default="__________")
    ap.add_argument("--time", default="__________")
    ap.add_argument("--skip-check", action="store_true")
    ns = ap.parse_args()

    gen_dir = ns.output.parent.parent / "_generation"
    gen_dir.mkdir(parents=True, exist_ok=True)
    if ns.template.resolve() == ns.output.resolve() and ns.template.exists():
        backup = gen_dir / f"{ns.template.stem}.backup.docx"
        shutil.copy2(ns.template, backup)
        print(f"Backup: {backup}")

    meta = Meta(date=ns.date, time=ns.time)
    mcq_key, match_keys, tf_key, fill_ans = generate(ns.template, ns.output, meta)
    from collections import Counter

    print(f"MCQ key: {' '.join(mcq_key[i:i+5] for i in range(0, 20, 5))}")
    print(f"MCQ balance: {dict(Counter(mcq_key))}")
    print(f"Matching: {match_keys[0]} / {match_keys[1]}")
    print(f"T/F key: {tf_key}")
    print(f"Fill block1: {fill_ans[0]}")
    print(f"Fill block2: {fill_ans[1]}")
    print(f"Written: {ns.output}")
    spec_path = ns.output.with_suffix(".spec.json")
    print(f"Written: {spec_path}")

    gen_spec = gen_dir / spec_path.name
    gen_spec.write_text(spec_path.read_text(encoding="utf-8"), encoding="utf-8")
    print(f"Copied spec: {gen_spec}")

    if not ns.skip_check:
        code = _run_quality_check(ns.output, spec_path)
        if code != 0:
            print("Quality checks reported issues (see above).", file=sys.stderr)
            return code
        print("Quality checks: all passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
