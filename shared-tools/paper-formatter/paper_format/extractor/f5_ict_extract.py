"""Build F5 ICT template profile JSON from reference DOCX."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document

from paper_format.extractor.paragraph_extract import extract_paragraph_profile
from paper_format.f5_ict_roles import F5_ICT_ROLE_SOURCE_PARAS
from paper_format.schema.paragraph_profile import ParagraphProfile


def build_role_library(docx_path: Path) -> dict[str, dict[str, Any]]:
    doc = Document(str(docx_path.expanduser().resolve()))
    roles: dict[str, dict[str, Any]] = {}
    for role, para_idx in F5_ICT_ROLE_SOURCE_PARAS.items():
        if para_idx >= len(doc.paragraphs):
            continue
        prof = extract_paragraph_profile(doc.paragraphs[para_idx], role=role)
        prof.notes = f"seed from para {para_idx}"
        roles[role] = prof.to_dict()
    return roles


def build_slot_layout(docx_path: Path) -> dict[str, Any]:
    """Paragraph index maps (legacy) + table indices — still needed for slot addressing."""
    import sys

    root = Path(__file__).resolve().parents[1]
    if str(root) not in sys.path:
        sys.path.insert(0, str(root))
    from f5_ict_layout import F5_ICT_MCQ_BLOCKS, F5_SECTION_HEADERS
    from f5_ict_tables import F5_ICT_REQUIRED_TABLES
    from written_slot_ranges import PART_B_RANGE, PART_C_RANGE, WRITTEN_SLOT_PARAGRAPHS

    return {
        "mcq_blocks": [list(b) for b in F5_ICT_MCQ_BLOCKS],
        "section_headers": sorted(F5_SECTION_HEADERS),
        "part_b_range": list(PART_B_RANGE),
        "part_c_range": list(PART_C_RANGE),
        "written_slots": {
            k: list(v) for k, v in WRITTEN_SLOT_PARAGRAPHS.items()
        },
        "required_tables": sorted(F5_ICT_REQUIRED_TABLES),
        "mcq_inter_gap_lines": 2,
    }


def build_mcq_patterns() -> dict[str, Any]:
    """Logical MCQ slot patterns (role sequences), not fixed template spans."""
    return {
        "mcq_single": {
            "description": "單選：題幹 → 空行 → A–D",
            "roles": ["mcq.stem", "mcq.blank_before_options", "mcq.option"] * 1
            + ["mcq.option"] * 3,
            "min_lines": 6,
        },
        "mcq_combo": {
            "description": "組合題：(1)(2)(3) 在題幹後、選項前",
            "roles": [
                "mcq.stem",
                "mcq.blank_before_options",
                "mcq.combo_sub",
                "mcq.combo_sub",
                "mcq.combo_sub",
                "mcq.blank_before_options",
                "mcq.option",
                "mcq.option",
                "mcq.option",
                "mcq.option",
            ],
            "min_lines": 10,
        },
        "mcq_with_code": {
            "description": "題幹 + 偽代碼（左縮排 Courier）+ 追問 + 選項",
            "note": "code lines repeat mcq.code_left; renderer expands by content",
        },
    }


def extract_f5_ict_profile(docx_path: Path, *, source_label: str) -> dict[str, Any]:
    docx_path = docx_path.expanduser().resolve()
    repo = Path(__file__).resolve().parents[4]
    try:
        source_rel = str(docx_path.relative_to(repo))
    except ValueError:
        source_rel = str(docx_path)
    return {
        "version": 1,
        "kind": "f5_ict_exam02_template_profile",
        "source_docx": source_rel,
        "source_label": source_label,
        "roles": build_role_library(docx_path),
        "layout": build_slot_layout(docx_path),
        "patterns": build_mcq_patterns(),
        "tables": {
            "note": "Table cell formatting uses same role library; populate via apply_f5_ict_table_content",
        },
    }


def save_profile(profile: dict[str, Any], out_path: Path) -> None:
    out_path = out_path.expanduser().resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")


def load_profile(path: Path) -> dict[str, Any]:
    return json.loads(path.expanduser().resolve().read_text(encoding="utf-8"))


def get_role_profile(profile: dict[str, Any], role: str) -> ParagraphProfile | None:
    raw = (profile.get("roles") or {}).get(role)
    if not raw:
        return None
    return ParagraphProfile.from_dict(raw)
