"""Clone body tables from 24_25 F5 ICT Exam02 template (preserve Word cell/table styles)."""

from __future__ import annotations

from copy import deepcopy
from functools import lru_cache
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

_REPO = Path(__file__).resolve().parents[2]
F5_ICT_TEMPLATE_DOCX = (
    _REPO
    / "Subjects/S5-ICT/past-papers/2024-2025/Term 02/WrittenExam/24_25_S5_ICT_Exam02.docx"
)

# Prototype indices in 24_25 template (verified 2026-05).
TBL_SPREADSHEET = 3
TBL_ORDER_SAMPLE = 11
TBL_SCHEMA = 18
TBL_STACK_TRACE = 31
TBL_TRACE_ARRAY = 6


@lru_cache(maxsize=1)
def _template_document() -> Document:
    return Document(str(F5_ICT_TEMPLATE_DOCX))


def template_table(index: int) -> Table:
    doc = _template_document()
    if index < 0 or index >= len(doc.tables):
        raise IndexError(f"template table {index} not in {F5_ICT_TEMPLATE_DOCX.name}")
    return doc.tables[index]


def _table_rows(tbl_el) -> list:
    return list(tbl_el.findall(qn("w:tr")))


def _row_cells(tr) -> list:
    return list(tr.findall(qn("w:tc")))


def resize_table_uniform(table: Table, rows: int, cols: int) -> None:
    """Trim or grow a uniform grid table while cloning row/cell XML from last row."""
    if rows < 1 or cols < 1:
        raise ValueError("rows and cols must be >= 1")
    tbl = table._tbl
    trs = _table_rows(tbl)
    while len(trs) > rows:
        tbl.remove(trs[-1])
        trs = _table_rows(tbl)
    while len(trs) < rows:
        proto = trs[-1]
        tbl.append(deepcopy(proto))
        trs = _table_rows(tbl)
    for tr in _table_rows(tbl):
        tcs = _row_cells(tr)
        while len(tcs) > cols:
            tr.remove(tcs[-1])
            tcs = _row_cells(tr)
        while len(tcs) < cols:
            tr.append(deepcopy(tcs[-1]))
            tcs = _row_cells(tr)


def insert_table_clone_after(
    paragraph,
    prototype_index: int,
    *,
    rows: int,
    cols: int,
) -> Table:
    """Insert a deep copy of a template table after `paragraph`, resized to rows×cols."""
    proto = template_table(prototype_index)
    new_tbl = deepcopy(proto._tbl)
    paragraph._p.addnext(new_tbl)
    table = Table(new_tbl, paragraph.part)
    resize_table_uniform(table, rows, cols)
    return table
