#!/usr/bin/env python3
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document


def _norm(s: str) -> str:
    return " ".join(s.replace("\t", " ").replace("\r", "").split())


def main(argv: list[str]) -> int:
    ap = argparse.ArgumentParser(description="Content sanity-check for S3 CMP Term02 exam docx.")
    ap.add_argument("--docx", required=True, help="Path to generated Term02 exam .docx")
    args = ap.parse_args(argv)

    p = Path(args.docx).expanduser().resolve()
    if not p.exists():
        print(f"ERROR: not found: {p}", file=sys.stderr)
        return 2

    d = Document(str(p))

    must_contain = [
        # Cover / term
        "下學期考試",
        # Section headings (question paper)
        "甲部 – 多項選擇題",
        "乙部 – 配對題",
        "丙部 – 短答題",
        "丁部 – 結構題",
        # A few key questions to ensure content isn't truncated
        "1.\t下列哪一項最能描述「外建／第三方函數庫」？",
        "12.\t以下哪一個 `quiz_data.json` 結構最合理？",
        # Answer key line should exist (grouped with spaces like template)
        "BABBB CBCBB BCDAC",
    ]

    # Flatten all visible text: paragraphs + tables.
    texts: list[str] = []
    texts += [para.text for para in d.paragraphs if para.text]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text:
                    texts.append(cell.text)

    hay = _norm("\n".join(texts))

    missing = [m for m in must_contain if _norm(m) not in hay]
    if missing:
        print("FAIL: missing expected content snippets:")
        for m in missing:
            print(f"- {m}")
        return 1

    print("OK: key content snippets found (not truncated).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))

