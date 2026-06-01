"""Validate 乙部 (short/structured) and 丙部 (long) written-question layout in rendered DOCX."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from docx import Document

# Section headers in F5/S6 ICT past papers.
SECTION_B_MARKERS = ("乙部", "Section B")
SECTION_C_MARKERS = ("丙部", "Section C")

# Sub-part labels at start of line (after optional tabs).
_SUBPART_RE = re.compile(r"^\t*\(([a-z]|[ivx]+)\)\t", re.IGNORECASE)
_MARKS_RE = re.compile(r"\((\d+)\s*分\)\s*$")
_TOPIC_NUM_RE = re.compile(r"^\d+\.\t")
# Flag common markdown artifacts that should not appear in DOCX.
# Note: underscores are also used as answer blanks (e.g. "________"), so do NOT flag "__".
_MARKDOWN_RE = re.compile(r"\*\*|`")

# Minimum tab-only blank lines expected after marked sub-parts in 乙部.
_MIN_ANSWER_BLANK_TABS = 4


@dataclass
class WrittenIssue:
    severity: str  # error | warning
    code: str
    message: str
    paragraph: Optional[int] = None

    def to_dict(self) -> dict[str, Any]:
        d: dict[str, Any] = {
            "severity": self.severity,
            "code": self.code,
            "message": self.message,
        }
        if self.paragraph is not None:
            d["paragraph"] = self.paragraph
        return d


@dataclass
class WrittenSectionCheckResult:
    ok: bool
    issues: list[WrittenIssue] = field(default_factory=list)
    sections_found: list[str] = field(default_factory=list)
    part_b_paragraphs: int = 0
    part_c_paragraphs: int = 0

    def to_dict(self) -> dict[str, Any]:
        return {
            "ok": self.ok,
            "sections_found": self.sections_found,
            "part_b_paragraphs": self.part_b_paragraphs,
            "part_c_paragraphs": self.part_c_paragraphs,
            "issues": [i.to_dict() for i in self.issues],
        }


def _find_section_starts(doc: Document) -> dict[str, int]:
    found: dict[str, int] = {}
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        if any(m in t for m in SECTION_B_MARKERS) and "section_b" not in found:
            found["section_b"] = i
        if any(m in t for m in SECTION_C_MARKERS) and "section_c" not in found:
            found["section_c"] = i
    return found


def _section_range(doc: Document, starts: dict[str, int], key: str) -> tuple[int, int] | None:
    if key not in starts:
        return None
    start = starts[key]
    order = ["section_b", "section_c", "answer"]
    idx = order.index(key)
    end = len(doc.paragraphs)
    for later in order[idx + 1 :]:
        if later in starts:
            end = starts[later]
            break
    # Heuristic: answer sheet often starts with 「標準答案」 or is mostly empty tail.
    for i in range(start + 1, min(len(doc.paragraphs), start + 400)):
        t = doc.paragraphs[i].text.strip()
        if t.startswith("標準答案") or t.startswith("Standard Answer"):
            end = min(end, i)
            break
    return start, end


def _is_tab_blank(text: str) -> bool:
    return bool(text) and not text.strip() and "\t" in text


def _has_leading_tab(text: str) -> bool:
    return text.startswith("\t")


def _check_part(
    doc: Document,
    start: int,
    end: int,
    *,
    part_name: str,
    strict_blanks: bool,
) -> list[WrittenIssue]:
    issues: list[WrittenIssue] = []
    marked_subparts: list[int] = []

    for i in range(start, end):
        text = doc.paragraphs[i].text
        if not text.strip():
            continue

        if _MARKDOWN_RE.search(text):
            issues.append(
                WrittenIssue(
                    "error",
                    "markdown_artifact",
                    f"{part_name}: paragraph contains markdown formatting (**, __, or `).",
                    paragraph=i,
                )
            )

        if "\n" in text and part_name == "乙部":
            issues.append(
                WrittenIssue(
                    "warning",
                    "multiline_paragraph",
                    f"{part_name}: paragraph {i} embeds multiple lines; "
                    "split pseudocode/stems across template slots instead.",
                    paragraph=i,
                )
            )

        if _SUBPART_RE.match(text):
            if not _has_leading_tab(text):
                issues.append(
                    WrittenIssue(
                        "error",
                        "subpart_indent",
                        f"{part_name}: sub-part at paragraph {i} should start with a tab "
                        f"before '(…)', got: {text[:60]!r}…",
                        paragraph=i,
                    )
                )
            if _MARKS_RE.search(text):
                marked_subparts.append(i)

        # Standalone (i)/(ii) lines without leading double-tab in 丙部 SQL blocks.
        if part_name == "丙部" and re.match(r"^\([ivx]+\)\t", text, re.I):
            issues.append(
                WrittenIssue(
                    "error",
                    "nested_subpart_indent",
                    f"丙部: nested sub-part at paragraph {i} should use '\\t\\t(i)\\t…'.",
                    paragraph=i,
                )
            )

        if _TOPIC_NUM_RE.match(text) and not _MARKS_RE.search(text):
            issues.append(
                WrittenIssue(
                    "warning",
                    "topic_missing_marks",
                    f"{part_name}: numbered topic at paragraph {i} has no '(N 分)' suffix.",
                    paragraph=i,
                )
            )

    if strict_blanks:
        for pi in marked_subparts:
            blank_found = False
            for j in range(pi + 1, min(pi + 4, end)):
                nxt = doc.paragraphs[j].text
                if _is_tab_blank(nxt):
                    blank_found = True
                    break
                if nxt.strip() and _SUBPART_RE.match(nxt):
                    break
            if not blank_found:
                issues.append(
                    WrittenIssue(
                        "warning",
                        "missing_answer_blank",
                        f"{part_name}: marked sub-part at paragraph {pi} "
                        "is not followed by a tab-only answer blank line.",
                        paragraph=pi,
                    )
                )

    return issues


def _compare_template_tabs(
    doc: Document,
    template: Document,
    start: int,
    end: int,
) -> list[WrittenIssue]:
    issues: list[WrittenIssue] = []
    t_len = len(template.paragraphs)
    for i in range(start, end):
        if i >= t_len:
            break
        cand = doc.paragraphs[i].text
        ref = template.paragraphs[i].text
        if not cand.strip() or not ref.strip():
            continue
        if _SUBPART_RE.match(ref) or _SUBPART_RE.match(cand):
            ref_tabs = ref.count("\t")
            cand_tabs = cand.count("\t")
            if cand_tabs < ref_tabs - 1:
                issues.append(
                    WrittenIssue(
                        "warning",
                        "tab_depth_mismatch",
                        f"Paragraph {i}: candidate has {cand_tabs} tabs, "
                        f"template reference has {ref_tabs}.",
                        paragraph=i,
                    )
                )
    return issues


def check_written_sections(
    docx_path: Path,
    *,
    template_docx_path: Optional[Path] = None,
    strict_blanks: bool = True,
) -> WrittenSectionCheckResult:
    docx_path = docx_path.expanduser().resolve()
    doc = Document(str(docx_path))
    starts = _find_section_starts(doc)

    issues: list[WrittenIssue] = []
    sections_found: list[str] = []

    if "section_b" not in starts:
        issues.append(
            WrittenIssue("error", "missing_section_b", "Cannot find 乙部 / Section B header.")
        )
    else:
        sections_found.append("乙部")

    if "section_c" not in starts:
        issues.append(
            WrittenIssue("error", "missing_section_c", "Cannot find 丙部 / Section C header.")
        )
    else:
        sections_found.append("丙部")

    part_b_count = 0
    part_c_count = 0

    b_range = _section_range(doc, starts, "section_b")
    if b_range:
        b_start, b_end = b_range
        part_b_count = b_end - b_start
        issues.extend(
            _check_part(doc, b_start, b_end, part_name="乙部", strict_blanks=strict_blanks)
        )

    c_range = _section_range(doc, starts, "section_c")
    if c_range:
        c_start, c_end = c_range
        part_c_count = c_end - c_start
        issues.extend(
            _check_part(
                doc,
                c_start,
                c_end,
                part_name="丙部",
                strict_blanks=False,
            )
        )

    if template_docx_path and template_docx_path.exists():
        template = Document(str(template_docx_path.expanduser().resolve()))
        if b_range:
            issues.extend(_compare_template_tabs(doc, template, b_range[0], b_range[1]))
        if c_range:
            issues.extend(_compare_template_tabs(doc, template, c_range[0], c_range[1]))

    has_errors = any(i.severity == "error" for i in issues)
    return WrittenSectionCheckResult(
        ok=not has_errors,
        issues=issues,
        sections_found=sections_found,
        part_b_paragraphs=part_b_count,
        part_c_paragraphs=part_c_count,
    )


def format_written_report(result: WrittenSectionCheckResult) -> str:
    lines = ["=== Written sections (乙部 / 丙部) ==="]
    if result.sections_found:
        lines.append(f"Sections: {', '.join(result.sections_found)}")
        lines.append(
            f"Paragraph spans: 乙部≈{result.part_b_paragraphs}, 丙部≈{result.part_c_paragraphs}"
        )
    if not result.issues:
        lines.append("No issues.")
    else:
        for issue in result.issues:
            loc = f" [para {issue.paragraph}]" if issue.paragraph is not None else ""
            lines.append(f"  [{issue.severity.upper()}] {issue.code}{loc}: {issue.message}")
    lines.append(f"Result: {'PASS' if result.ok else 'ISSUES FOUND'}")
    return "\n".join(lines)
