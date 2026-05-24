"""Check MCQ blocks in F5 ICT DOCX for unnecessary internal blank lines."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

_OPT = re.compile(r"^\t[A-D]\.\t")
_SUB = re.compile(r"^\t+\(\d+\)\t")


@dataclass
class McqBlankIssue:
    question: int
    paragraph: int
    kind: str
    detail: str

    def to_dict(self) -> dict[str, str]:
        return {
            "question": str(self.question),
            "paragraph": str(self.paragraph),
            "kind": self.kind,
            "detail": self.detail,
        }


@dataclass
class McqBlankCheckResult:
    ok: bool
    issues: list[McqBlankIssue] = field(default_factory=list)
    checked: int = 0

    def to_dict(self) -> dict[str, Any]:
        return {
            "ok": self.ok,
            "checked": self.checked,
            "issues": [i.to_dict() for i in self.issues],
        }


def _first_option_index(lines: list[str]) -> int:
    for i, line in enumerate(lines):
        if _OPT.match(line):
            return i
    return len(lines)


def _is_combo_block(lines: list[str]) -> bool:
    return any(_SUB.match(line) for line in lines)


def check_mcq_block_lines(question: int, lines: list[str], *, para_start: int) -> list[McqBlankIssue]:
    issues: list[McqBlankIssue] = []
    opt_i = _first_option_index(lines)
    if opt_i >= len(lines):
        return issues

    content = lines[:opt_i]
    options = lines[opt_i : opt_i + 4]
    trailing = lines[opt_i + 4 :]
    combo = _is_combo_block(content)

    # Trailing blanks after D are template padding — allowed.
    if any(t.strip() for t in trailing):
        issues.append(
            McqBlankIssue(
                question,
                para_start + opt_i + 4,
                "content_after_options",
                "Non-empty text after option D inside MCQ block",
            )
        )

    # Content region: analyse blank runs
    blank_runs: list[tuple[int, int]] = []
    i = 0
    while i < len(content):
        if content[i].strip() == "":
            start = i
            while i < len(content) and content[i].strip() == "":
                i += 1
            blank_runs.append((start, i))
        else:
            i += 1

    sub_idxs = [j for j, line in enumerate(content) if _SUB.match(line)]

    for start, end in blank_runs:
        run_len = end - start
        para = para_start + start

        if combo:
            # Allowed: one blank before (1) if stem above; one blank before options after (3)
            before_sub = sub_idxs and start < sub_idxs[0]
            after_sub = sub_idxs and start > sub_idxs[-1]
            if before_sub and run_len == 1:
                continue
            if after_sub and run_len == 1:
                continue
            if run_len > 1 or (not before_sub and not after_sub):
                issues.append(
                    McqBlankIssue(
                        question,
                        para,
                        "excess_blank_run",
                        f"Combo MCQ: {run_len} consecutive blank line(s) at content index {start}",
                    )
                )
            elif not before_sub and not after_sub:
                issues.append(
                    McqBlankIssue(
                        question,
                        para,
                        "unexpected_blank",
                        "Blank line between (1)(2)(3) statements",
                    )
                )
        else:
            # Non-combo: at most one blank, immediately before options
            if start != len(content) - 1 or run_len != 1:
                issues.append(
                    McqBlankIssue(
                        question,
                        para,
                        "excess_blank_run",
                        f"{run_len} blank line(s) before options (expected at most one line "
                        f"immediately above A.)",
                    )
                )

    if not combo and not any(content[j].strip() == "" for j in range(len(content))):
        # Single-line stems may have zero or one blank; warn only if span suggests padding
        pass

    if len(options) < 4:
        issues.append(
            McqBlankIssue(
                question,
                para_start + opt_i,
                "missing_options",
                f"Only {len(options)} option lines found (expected 4)",
            )
        )

    return issues


def check_mcq_blank_lines(
    docx_path: Path,
    blocks: tuple[tuple[int, int], ...],
) -> McqBlankCheckResult:
    from docx import Document

    docx_path = docx_path.expanduser().resolve()
    doc = Document(str(docx_path))
    issues: list[McqBlankIssue] = []

    for qi, (start, end) in enumerate(blocks, start=1):
        lines = [doc.paragraphs[i].text for i in range(start, end)]
        issues.extend(check_mcq_block_lines(qi, lines, para_start=start))
        for i in range(start, end):
            p = doc.paragraphs[i]
            if _OPT.match(p.text):
                a = p.paragraph_format.alignment
                if a is not None and a.name == "CENTER":
                    issues.append(
                        McqBlankIssue(
                            qi,
                            i,
                            "option_centered",
                            f"Option line must be left-aligned, not CENTER: {p.text[:40]!r}",
                        )
                    )

    return McqBlankCheckResult(
        ok=not issues,
        issues=issues,
        checked=len(blocks),
    )


def format_mcq_blank_report(result: McqBlankCheckResult) -> str:
    lines = [
        "=== MCQ blank-line check (甲部) ===",
        f"Checked {result.checked} questions",
    ]
    if result.ok:
        lines.append("Status: OK")
        return "\n".join(lines)
    lines.append("Status: ISSUES")
    for issue in result.issues:
        lines.append(
            f"  [Q{issue.question} para {issue.paragraph}] {issue.kind}: {issue.detail}"
        )
    return "\n".join(lines)
