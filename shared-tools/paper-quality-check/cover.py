"""Extract and verify CMP-style exam cover page fields in DOCX."""
from __future__ import annotations

import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Optional

from filename_meta import InferredExamMeta, infer_exam_meta_from_path

try:
    from docx import Document
except ImportError as e:  # pragma: no cover
    raise ImportError("python-docx required: pip install python-docx") from e

# Chinese CMP cover (table 0, cell 0,0) — matches paper-formatter/docx_inplace.py
_ZH_COVER_PARA = {
    "school": 3,
    "year_term": 4,
    "level": 6,
    "paper": 7,
}


@dataclass(frozen=True)
class CoverFields:
    school: str = ""
    year_term: str = ""
    level: str = ""
    paper: str = ""

    def to_dict(self) -> dict[str, str]:
        return asdict(self)


@dataclass
class CoverIssue:
    field: str
    expected: str
    actual: str


@dataclass
class CoverCheckResult:
    candidate: str
    inferred_from: str
    expected: CoverFields
    actual: CoverFields
    ok: bool
    issues: list[CoverIssue] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "ok": self.ok,
            "candidate": self.candidate,
            "inferred_from": self.inferred_from,
            "expected": self.expected.to_dict(),
            "actual": self.actual.to_dict(),
            "issues": [asdict(i) for i in self.issues],
        }


def _norm(s: str) -> str:
    s = re.sub(r"\s+", "", s)
    return s.replace("–", "-").replace("—", "-").replace("－", "-")


def _para_text(cell, idx: int) -> str:
    paras = cell.paragraphs
    if idx >= len(paras):
        return ""
    return paras[idx].text.strip()


def extract_cover_fields(docx_path: Path) -> CoverFields:
    docx_path = docx_path.expanduser().resolve()
    doc = Document(str(docx_path))
    if not doc.tables:
        return CoverFields()
    cell = doc.tables[0].cell(0, 0)
    idx = _ZH_COVER_PARA
    return CoverFields(
        school=_para_text(cell, idx["school"]),
        year_term=_para_text(cell, idx["year_term"]),
        level=_para_text(cell, idx["level"]),
        paper=_para_text(cell, idx["paper"]),
    )


def _expected_from_meta(meta: InferredExamMeta) -> CoverFields:
    return CoverFields(
        school="迦密聖道中學",
        year_term=meta.year_term_cover,
        level=meta.level_line,
        paper=meta.paper_title,
    )


def check_cover(
    docx_path: Path,
    *,
    inferred_meta: Optional[InferredExamMeta] = None,
    filename_for_infer: Optional[Path] = None,
) -> Optional[CoverCheckResult]:
    """
    Verify cover page lines against metadata inferred from filename.
    Returns None when filename cannot be parsed.
    """
    docx_path = docx_path.expanduser().resolve()
    infer_path = filename_for_infer or docx_path
    meta = inferred_meta or infer_exam_meta_from_path(infer_path)
    if meta is None:
        return None

    actual = extract_cover_fields(docx_path)
    expected = _expected_from_meta(meta)
    issues: list[CoverIssue] = []

    checks = [
        ("year_term", expected.year_term, actual.year_term),
        ("level", expected.level, actual.level),
        ("paper", expected.paper, actual.paper),
    ]
    for field_name, exp, act in checks:
        if exp and _norm(exp) not in _norm(act) and _norm(act) not in _norm(exp):
            issues.append(CoverIssue(field=field_name, expected=exp, actual=act))

    return CoverCheckResult(
        candidate=str(docx_path),
        inferred_from=meta.source_filename or Path(infer_path).name,
        expected=expected,
        actual=actual,
        ok=len(issues) == 0,
        issues=issues,
    )


def format_cover_report(result: CoverCheckResult) -> str:
    lines = [
        f"Cover page: {'OK' if result.ok else 'ISSUES FOUND'}",
        f"Inferred from filename: {result.inferred_from}",
        f"  year/term  expected: {result.expected.year_term!r}",
        f"             actual:   {result.actual.year_term!r}",
        f"  level      expected: {result.expected.level!r}",
        f"             actual:   {result.actual.level!r}",
        f"  paper      expected: {result.expected.paper!r}",
        f"             actual:   {result.actual.paper!r}",
    ]
    for issue in result.issues:
        lines.append(f"\n  [{issue.field}] mismatch")
        lines.append(f"    expected: {issue.expected!r}")
        lines.append(f"    actual:   {issue.actual!r}")
    return "\n".join(lines)
