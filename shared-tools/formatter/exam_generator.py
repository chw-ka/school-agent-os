from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PLACEHOLDER_TITLE = "{{TITLE}}"
PLACEHOLDER_QUESTIONS = "{{QUESTIONS}}"


@dataclass(frozen=True)
class ExamQuestion:
    markdown: str
    points: Optional[float] = None


@dataclass(frozen=True)
class ExamSpec:
    title: str
    questions: list[ExamQuestion]


def _set_run_font(run, font_name: str, font_size_pt: int) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    r = run._element.rPr
    rFonts = r.rFonts if r.rFonts is not None else r._add_rFonts()
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _apply_default_style(doc: Document, font_name: str, font_size_pt: int) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _apply_standard_margins_a4(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def _add_markdownish_runs(paragraph, text: str, font_name: str, font_size_pt: int) -> None:
    """
    Minimal Markdown support:
    - **bold**
    - *italic*
    Everything else is treated as plain text.
    """

    i = 0
    while i < len(text):
        bold_m = _BOLD_RE.search(text, i)
        italic_m = _ITALIC_RE.search(text, i)
        candidates = [m for m in [bold_m, italic_m] if m is not None]
        if not candidates:
            run = paragraph.add_run(text[i:])
            _set_run_font(run, font_name, font_size_pt)
            return

        m = min(candidates, key=lambda x: x.start())
        if m.start() > i:
            run = paragraph.add_run(text[i : m.start()])
            _set_run_font(run, font_name, font_size_pt)

        content = m.group(1)
        run = paragraph.add_run(content)
        if m.re is _BOLD_RE:
            run.bold = True
        else:
            run.italic = True
        _set_run_font(run, font_name, font_size_pt)
        i = m.end()


def _iter_nonempty_lines(markdown: str) -> Iterable[str]:
    for raw in markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw.strip()
        if line:
            yield line


def _insert_paragraph_after(paragraph, text: str = "", style: Optional[str] = None):
    # python-docx doesn't expose insert_paragraph_after on Paragraph; emulate via insert_paragraph_before
    # on the following paragraph if it exists, otherwise append at end of document body.
    parent_el = paragraph._element
    next_el = parent_el.getnext()
    body = paragraph.part.element.body

    if next_el is not None:
        new_p = paragraph.insert_paragraph_before(text, style=style)
        body.remove(new_p._element)
        parent_el.addnext(new_p._element)
        return new_p

    doc = paragraph.part.document
    new_p = doc.add_paragraph(text)
    if style:
        new_p.style = style
    return new_p


def _replace_text_in_paragraph(paragraph, needle: str, replacement: str) -> bool:
    if needle not in paragraph.text:
        return False
    paragraph.text = paragraph.text.replace(needle, replacement)
    return True


def _find_first_paragraph(doc: Document, needle: str):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def _load_exam_spec(path: Path) -> ExamSpec:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("Input JSON must be an object.")

    title = data.get("title")
    if not isinstance(title, str) or not title.strip():
        raise ValueError('Input JSON must contain non-empty string field "title".')

    qs_raw = data.get("questions")
    if not isinstance(qs_raw, list) or not qs_raw:
        raise ValueError('Input JSON must contain non-empty array field "questions".')

    questions: list[ExamQuestion] = []
    for idx, q in enumerate(qs_raw, start=1):
        if not isinstance(q, dict):
            raise ValueError(f"Question #{idx} must be an object.")

        md = q.get("markdown")
        if not isinstance(md, str) or not md.strip():
            raise ValueError(f'Question #{idx} must contain non-empty string field "markdown".')

        points = q.get("points")
        if points is not None and not isinstance(points, (int, float)):
            raise ValueError(f'Question #{idx} field "points" must be a number if present.')

        questions.append(ExamQuestion(markdown=md, points=float(points) if points is not None else None))

    return ExamSpec(title=title.strip(), questions=questions)


def _render_questions_at_anchor(
    doc: Document,
    anchor_paragraph,
    spec: ExamSpec,
    *,
    font_name: str,
    font_size_pt: int,
) -> None:
    if PLACEHOLDER_QUESTIONS in anchor_paragraph.text:
        anchor_paragraph.text = anchor_paragraph.text.replace(PLACEHOLDER_QUESTIONS, "").strip()

    current = anchor_paragraph
    for q in spec.questions:
        first_line, *rest_lines = list(_iter_nonempty_lines(q.markdown))

        q_p = _insert_paragraph_after(current, "", style="List Number")
        _add_markdownish_runs(q_p, first_line, font_name, font_size_pt)
        if q.points is not None:
            run = q_p.add_run(f"  ({q.points:g} 分)")
            _set_run_font(run, font_name, font_size_pt)
        current = q_p

        for line in rest_lines:
            sub_p = _insert_paragraph_after(current, "", style="List Bullet 2")
            _add_markdownish_runs(sub_p, line, font_name, font_size_pt)
            current = sub_p

        current = _insert_paragraph_after(current, "", style="Normal")


def generate_exam_docx(
    *,
    input_json_path: Path,
    template_docx_path: Path,
    output_docx_path: Path,
    font_name: str = "新細明體",
    font_size_pt: int = 12,
    enforce_standard: bool = False,
) -> None:
    spec = _load_exam_spec(input_json_path)

    if not template_docx_path.exists():
        raise FileNotFoundError(f"Template not found: {template_docx_path}")

    doc = Document(str(template_docx_path))

    if enforce_standard:
        _apply_default_style(doc, font_name, font_size_pt)
        _apply_standard_margins_a4(doc)

    title_anchor = _find_first_paragraph(doc, PLACEHOLDER_TITLE)
    if title_anchor is not None:
        _replace_text_in_paragraph(title_anchor, PLACEHOLDER_TITLE, spec.title)
    else:
        p0 = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph("")
        p0.text = spec.title

    questions_anchor = _find_first_paragraph(doc, PLACEHOLDER_QUESTIONS)
    if questions_anchor is None:
        questions_anchor = doc.add_paragraph(PLACEHOLDER_QUESTIONS)

    _render_questions_at_anchor(
        doc,
        questions_anchor,
        spec,
        font_name=font_name,
        font_size_pt=font_size_pt,
    )

    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx_path))


def _parse_args(argv: Optional[list[str]] = None) -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Generate an exam .docx from question JSON (Markdown) using a .docx template.",
    )
    p.add_argument("--input", required=True, help="Path to exam JSON (utf-8).")
    p.add_argument("--template", required=True, help="Path to .docx template under templates/.")
    p.add_argument("--output", required=True, help="Output .docx path.")
    p.add_argument("--font-name", default="新細明體", help="Default font name (used for inserted text).")
    p.add_argument("--font-size", type=int, default=12, help="Default font size in pt (used for inserted text).")
    p.add_argument(
        "--enforce-standard",
        action="store_true",
        help="Force Normal style + A4 margins to school standard (may override template).",
    )
    return p.parse_args(argv)


def main(argv: Optional[list[str]] = None) -> int:
    args = _parse_args(argv)
    generate_exam_docx(
        input_json_path=Path(args.input),
        template_docx_path=Path(args.template),
        output_docx_path=Path(args.output),
        font_name=args.font_name,
        font_size_pt=args.font_size,
        enforce_standard=bool(args.enforce_standard),
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
