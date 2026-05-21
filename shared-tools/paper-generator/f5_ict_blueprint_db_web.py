#!/usr/bin/env python3
"""
Generate F5 ICT exam (Database elective; no networking / HTML / web dev) from template layout.
Always runs post-generation similarity check vs template + past 3 years.
"""
from __future__ import annotations

import argparse
import shutil
import sys
from pathlib import Path

from docx import Document

from dse_ict_style import COMBO_OPTS_1_ONLY, COMBO_OPTS_3_ONLY, COMBO_OPTS_ALL, COMBO_OPTS_1_2_OK
from f5_ict_spec import F5_MCQ_CORRECT_INDEX, build_f5_ict_exam_spec
from post_check import repo_root, run_spec_check

_FMT = Path(__file__).resolve().parents[1] / "paper-formatter"
if str(_FMT) not in sys.path:
    sys.path.insert(0, str(_FMT))
from docx_inplace import ZhCoverPatch, apply_cmp_cover_zh_on_en_layout, set_paragraph_text_distribute
from mcq_answer_keys import build_random_mcq_key
from written_layout import (
    ANSWER_BLANK,
    ANSWER_BLANK_LONG,
    DIAGRAM_BLANK,
    code_line,
    replace_span,
    sql_line,
    stem,
    subpart,
    topic_header,
)

_QCHECK = Path(__file__).resolve().parents[1] / "question-quality-check"
_PCHECK = Path(__file__).resolve().parents[1] / "paper-quality-check"
for _d in (_QCHECK, _PCHECK):
    if str(_d) not in sys.path:
        sys.path.insert(0, str(_d))
from exam_spec import save_spec
from footer import apply_footer_meta

# --- generation logic (template-span preserving) ---


def is_option_line(t: str) -> bool:
    s = t.lstrip()
    return len(s) >= 2 and s[0] in "ABCD" and s[1] == "."


def mcq_blocks(doc: Document) -> list[tuple[int, int]]:
    blocks: list[tuple[int, int]] = []
    i = 1
    while i < 311:
        while i < 311 and not doc.paragraphs[i].text.strip():
            i += 1
        if i >= 311:
            break
        if "甲部" in doc.paragraphs[i].text:
            i += 1
            continue
        start = i
        while i < 311:
            t = doc.paragraphs[i].text
            if is_option_line(t):
                break
            i += 1
        while i < 311 and is_option_line(doc.paragraphs[i].text):
            i += 1
        end = i
        blocks.append((start, end))
        while i < 311 and not doc.paragraphs[i].text.strip():
            i += 1
    return blocks


def opt(a: str, b: str, c: str, d: str) -> list[str]:
    return [f"\tA.\t{a}", f"\tB.\t{b}", f"\tC.\t{c}", f"\tD.\t{d}"]


def build_mcq_payload() -> list[list[str]]:
    o = opt
    c1, c3, call, c12 = COMBO_OPTS_1_ONLY, COMBO_OPTS_3_ONLY, COMBO_OPTS_ALL, COMBO_OPTS_1_2_OK
    rows: list[list[str]] = [
        ["下列哪一項是二進制 1100 0101 的十六進制表示？", ""] + o("0xC5", "0xCA", "0xA5", "0x5C"),
        ["以 8 位元二進制補碼表示十進位 -3，下列哪項正確？", ""] + o("1111 1101", "1111 1111", "1000 0011", "0000 0011"),
        ["下列哪句關於 ASCII 與 Unicode（UTF-8）儲存中文，何者正確？", ""]
        + o("兩者均以 1 字節儲存每個中文字", "UTF-8 的中文字通常比 ASCII 編碼佔更多字節", "ASCII 可完整表示所有繁體中文字", "UTF-8 只能儲存英文字母"),
        ["一段未壓縮音訊：取樣頻率 44.1 kHz、取樣精度 16 bit、立體聲（2 聲道），長度 1 秒。下列何者最接近其檔案大小？", ""]
        + o("約 88 KB", "約 176 KB", "約 352 KB", "約 705 KB"),
        ["在試算表中，VLOOKUP／XLOOKUP 的主要用途是？", ""] + o("合併儲存格", "依關鍵值查找並傳回對應欄位", "計算滿足多條件的個數", "將文字轉為日期格式"),
        [
            "老師要統計「同時滿足：班別為甲班、分數>=60」的學生人數。下列哪項（些）最適合？",
            "",
            "\t\t(1)\tCOUNT",
            "\t\t(2)\tCOUNTIF",
            "\t\t(3)\tCOUNTIFS",
            "",
        ]
        + o(*c3),
        [
            "下列哪句關於快取記憶體（Cache）是正確的？",
            "",
            "（提示：快取位於 CPU 與主記憶體之間，用較小但較快的記憶體暫存常用資料以提升效能。）",
            "",
        ]
        + o(
            "在 CPU 與主記憶體之間提供較小但較快的暫存層以提升效能",
            "用作永久儲存作業系統與使用者文件",
            "負責把類比訊號轉為數位訊號",
            "只用作顯示輸出裝置",
        ),
    ]
    rows.extend(
        [
            [
                "比較固態硬碟（SSD）與傳統硬碟（HDD），下列哪項（些）句子是正確的？",
                "",
                "\t\t(1)\tSSD 通常較抗震",
                "\t\t(2)\tHDD 通常讀寫較快",
                "\t\t(3)\t相同容量下 SSD 通常較貴",
                "",
            ]
            + o("只有 (1)", "只有 (1) 和 (3)", "只有 (2) 和 (3)", "(1)、(2) 和 (3)"),
            ["下列哪項最屬於實用程式（Utility）的典型用途？", ""]
            + o("試算表繪製圖表", "磁碟重組或檔案壓縮等系統維護工作", "編輯數碼影片特效", "管理使用者帳戶與檔案權限"),
            [
                "一幅 800×600 像素的點陣圖，每像素以 24 bit 真彩色儲存（未壓縮）。下列何者最接近檔案大小？",
                "",
            ]
            + o("約 1.4 MB", "約 14 MB", "約 140 KB", "約 4.8 MB"),
            [
                "下列哪句關於有損與無損壓縮，何者正確？",
                "",
            ]
            + o(
                "有損壓縮可還原成與原檔完全相同",
                "無損壓縮常用於 JPEG 相片以縮小檔案",
                "有損壓縮可能犧牲部分畫質／音質以換取較小檔案",
                "兩者均不能減少檔案大小",
            ),
            [
                "下列哪句關於「欄位（Field）」是正確的？",
                "",
            ]
            + o(
                "儲存同一類資料的最小單位",
                "由多個欄位組成的一筆記錄",
                "整個資料庫的備份檔",
                "CPU 內的快取行",
            ),
            [
                "細看以下學號輸入規則：只接受「一個英文字母 + 五個數字」。此輸入會涉及下列哪（些）項目？",
                "",
                "\t\t(1)\t數據有效性檢驗",
                "\t\t(2)\t奇偶檢測",
                "\t\t(3)\t順序存取檔案",
                "",
            ]
            + o(*c1),
            [
                "下列哪句關於「資訊」與「數據」是正確的？",
                "",
            ]
            + o(
                "未經處理的原始符號或數值",
                "經整理後對決策有幫助的內容",
                "硬碟上的二進制位元",
                "壓縮演算法中的位元組",
            ),
            [
                "下列哪項（些）關於 RAM 與 ROM 的句子是正確的？",
                "",
                "\t\t(1)\tRAM 通常可讀寫且斷電後內容可能消失",
                "\t\t(2)\tROM 通常用於儲存開機程式",
                "\t\t(3)\tRAM 用作永久儲存使用者文件",
                "",
            ]
            + o("只有 (1)", "只有 (1) 和 (2)", "只有 (3)", "(1)、(2) 和 (3)"),
            [
                "下列哪項（些）輸入裝置與用途的配對是正確的？",
                "",
                "\t\t(1)\t條碼掃描器：讀取商品編號",
                "\t\t(2)\t麥克風：輸入聲音訊號",
                "\t\t(3)\t繪圖板：輸入位置與壓感",
                "",
            ]
            + o("只有 (1)", "只有 (1) 和 (2)", "(1)、(2) 和 (3)", "只有 (3)"),
            [
                "下列哪句關於「一筆記錄」是正確的？",
                "",
            ]
            + o("一個欄位內的單一數值", "由多個欄位值組成的一組相關資料", "整個資料庫的備份檔", "CPU 暫存器內容"),
            [
                "傳輸資料時加入校驗位，使每個字節中 1 的個數維持偶數。下列哪句關於此做法是正確的？",
                "",
            ]
            + o("這屬於奇偶檢測", "這屬於有效性檢驗", "這屬於資料壓縮", "這屬於檔案備份"),
            [
                "下列哪項處理方式最符合「累積一整批訂單後在深夜統一結算」？",
                "",
            ]
            + o("即時處理", "批次處理", "互動式處理", "感應器即時回饋"),
            [
                "比較直接存取與順序存取檔案，下列哪句是正確的？",
                "",
            ]
            + o(
                "順序存取必須從頭開始逐一讀取直至目標記錄",
                "直接存取可在已知索引下較快定位特定記錄",
                "兩者均不能讀取文字檔",
                "直接存取只能用於堆疊結構",
            ),
            [
                "下列哪項（些）關於向量圖與點陣圖的句子是正確的？",
                "",
                "\t\t(1)\t向量圖以數學描述圖形，放大較不易失真",
                "\t\t(2)\t點陣圖由像素組成",
                "\t\t(3)\t相片掃描結果通常以點陣圖儲存",
                "",
            ]
            + o("只有 (1)", "只有 (1) 和 (2)", "(1)、(2) 和 (3)", "只有 (3)"),
            [
                "下列哪項（些）關於作業系統（OS）主要功能的句子是正確的？",
                "",
                "\t\t(1)\t管理記憶體與處理器資源",
                "\t\t(2)\t提供使用者介面",
                "\t\t(3)\t必定只能同時執行一個程式",
                "",
            ]
            + o("只有 (1)", "只有 (1) 和 (2)", "只有 (3)", "(1) 和 (2) 均正確"),
            [
                "細看以下流程：學校以問卷收集學生意見，再人手輸入試算表分析。",
                "",
                "下列哪句最恰當描述此資訊處理流程？",
                "",
                "",
            ]
            + o(
                "問卷輸入屬輸入；試算表分析屬處理",
                "只屬輸出階段",
                "沒有處理階段",
                "不需收集數據",
            ),
            [
                "一段短片：",
                "",
                "\t\t解析度 640×480",
                "\t\t幀率 25 fps",
                "\t\t每幀未壓縮約 1 MB",
                "\t\t長度 2 秒",
                "",
                "\t下列何者最接近未壓縮檔案大小？",
                "",
            ]
            + o("約 25 MB", "約 50 MB", "約 100 MB", "約 200 MB"),
            [
                "下列哪項（些）關於檔案組織與存取的句子是正確的？",
                "",
                "\t\t(1)\t文字檔可按行順序讀取",
                "\t\t(2)\t隨機存取檔案可配合索引定位記錄",
                "\t\t(3)\t所有檔案均只能用順序存取",
                "\t\t（以上為一般課程描述，可能因教材用語略有差異）",
                "\t\t（請以「最合理」為準選擇答案）",
                "",
            ]
            + o("只有 (1)", "只有 (1) 和 (2)", "只有 (3)", "(1)、(2) 和 (3)"),
            [
                "下列哪項最屬於「專用軟件」？",
                "",
                "",
            ]
            + o("醫院掛號系統", "文字處理器", "網頁瀏覽器", "試算表"),
            [
                "一幅點陣圖每像素以 8 bit 儲存顏色。下列哪項正確？",
                "",
            ]
            + o("16", "256", "65536", "16777216"),
            [
                "下列哪項（些）關於數位音訊檔案的句子是正確的？",
                "",
                "\t\t(1)\t取樣頻率愈高，還原音質通常愈接近原聲",
                "\t\t(2)\t取樣精度愈高，每個取樣點可記錄的層次愈多",
                "\t\t(3)\t立體聲比單聲道需要更多聲道資料",
                "",
                "\t\t（請以「最合理」為準選擇答案）",
                "",
                "",
                "",
            ]
            + o("只有 (1)", "只有 (1) 和 (2)", "只有 (3)", "(1)、(2) 和 (3)"),
            ["下列哪項是十進位 10 的二進制表示？", ""]
            + o("1010", "1001", "1100", "0011"),
            [
                "某公司要把銷售數據製成月報表供管理層決策。",
                "",
                "此過程涉及下列哪（些）資訊處理階段？",
                "",
                "\t\t(1)\t輸入銷售數據",
                "\t\t(2)\t處理：計算總額與分類",
                "\t\t(3)\t輸出月報表",
                "",
                "",
                "",
            ]
            + o(*call),
        ]
    )
    expected = [6, 6, 6, 6, 6, 10, 8, 10, 6, 6, 6, 6, 10, 6, 10, 10, 6, 6, 6, 6, 10, 10, 9, 13, 12, 7, 6, 14, 6, 14]
    if len(rows) != len(expected):
        raise RuntimeError(f"MCQ count mismatch {len(rows)} vs {len(expected)}")
    for i, (row, exp) in enumerate(zip(rows, expected, strict=True), start=1):
        if len(row) != exp:
            raise RuntimeError(f"MCQ#{i} len {len(row)} expected {exp}")
    return rows


def apply_mcq(doc: Document, blocks: list[tuple[int, int]], payload: list[list[str]]) -> None:
    if len(blocks) != len(payload):
        raise RuntimeError("blocks/payload mismatch")
    for (start, end), lines in zip(blocks, payload):
        span = end - start
        if span != len(lines):
            raise RuntimeError(f"span {span} != payload {len(lines)} at {start}")
        for j, text in enumerate(lines):
            set_paragraph_text_distribute(doc.paragraphs[start + j], text)


def build_part_b() -> list[str]:
    lines: list[str] = [""] * 110
    base = 313

    def put(offset: int, s: str) -> None:
        lines[offset - base] = s

    # Q1 — spreadsheet (template slots 313–322)
    put(
        313,
        "「星晴網店」以試算表記錄訂單，欄位包括：A=訂單日期、B=產品類別、C=單價、D=數量、"
        "E=會員等級（VIP／一般）。F 欄為總價。請回答下列問題。",
    )
    put(
        316,
        subpart(
            "a",
            "在 F2 寫出一條公式（可用 IF 與 AND），然後複製到 F3:F200。"
            "規則：若 E2 為「VIP」且 D2>10，則總價為 C2*D2 的 8 折；否則為原價 C2*D2。",
            5,
        ),
    )
    put(317, ANSWER_BLANK)
    put(
        319,
        subpart(
            "b",
            "描述如何建立樞紐分析表：以「列」顯示產品類別；以「值」顯示總銷售額（對 F 欄求 SUM）。",
            5,
        ),
    )
    put(320, ANSWER_BLANK)

    # Q2 — data control & privacy (323–334)
    put(323, topic_header("2.", "數據控制與私隱", 10))
    put(326, subpart("a", "列出「有效性檢驗」與「奇偶檢測」的分別，並各舉一個輸入數據的例子。", 3))
    put(327, ANSWER_BLANK)
    put(328, subpart("b", "比較「直接存取」與「順序存取」讀取檔案記錄的優缺點。", 4))
    put(329, ANSWER_BLANK)
    put(330, ANSWER_BLANK_LONG)
    put(331, subpart("c", "學校收集學生健康申報資料時，提出兩項保障數據私隱的做法。", 3))
    put(332, ANSWER_BLANK)
    put(333, ANSWER_BLANK)

    # Q3 — algorithm trace (336–360)
    put(336, topic_header("3.", "算法追蹤", 10))
    put(338, "考慮以下偽代碼（陣列 A 索引由 1 開始，長度 n≥2）：")
    put(340, code_line("largest ← A[1]"))
    put(341, code_line("second_largest ← A[1]"))
    put(342, code_line("FOR i ← 2 TO n"))
    put(343, code_line("    IF A[i] > largest THEN", depth=2))
    put(344, code_line("        second_largest ← largest", depth=3))
    put(345, code_line("        largest ← A[i]", depth=3))
    put(346, code_line("    ELSE IF A[i] > second_largest AND A[i] < largest THEN", depth=2))
    put(347, code_line("        second_largest ← A[i]", depth=3))
    put(348, code_line("    ENDIF", depth=2))
    put(349, code_line("ENDFOR"))
    put(
        354,
        subpart(
            "a",
            "設 A = [10, 5, 20, 8, 20, 15]（n=6）。完成追蹤表，展示每次迭代後 largest 與 second_largest 的值。",
            5,
        ),
    )
    put(355, ANSWER_BLANK)
    put(356, ANSWER_BLANK_LONG)
    put(
        358,
        subpart(
            "b",
            "若陣列中最大值出現多次（如上例的 20），此算法是否仍能正確找出「次大值」？解釋你的答案。",
            5,
        ),
    )
    put(359, ANSWER_BLANK)
    put(360, ANSWER_BLANK_LONG)
    return lines


def build_part_c() -> list[str]:
    lines = [""] * 202
    b = 423

    def put(i: int, s: str) -> None:
        lines[i - b] = s

    put(423, "丙部 (40 分)：選修單元問答題（數據庫）")
    put(
        425,
        "某「社區中心活動報名系統」描述如下：一位會員可報名多個工作坊；一個工作坊亦可被多位會員報名。"
        "每次報名產生一筆報名記錄，包含報名日期與付款狀態。",
    )
    put(
        427,
        stem(
            "繪製實體關係圖（ERD）：須包含「會員 Member」「工作坊 Workshop」「報名 Registration」三個實體，"
            "標示多對多並以關聯實體拆解；並在圖中寫出主鍵（PK）與外鍵（FK）欄位名稱。",
            8,
            spaced_marks=True,
        ),
    )
    put(429, DIAGRAM_BLANK)

    put(440, "設 Member、Workshop、Registration 資料表如下：")
    put(442, subpart("b", "根據以下資料表："))
    put(444, sql_line("Member(MemberID, MemberName, JoinDate)"))
    put(445, sql_line("Workshop(WorkshopID, Title, Instructor, Fee)"))
    put(446, sql_line("Registration(RegID, MemberID, WorkshopID, RegDate, PayStatus)"))
    put(
        450,
        subpart(
            "i",
            "寫出一條 SQL，列出所有報名了 Instructor = 'Peter' 的工作坊之會員姓名（MemberName）。",
            4,
            depth=2,
        ),
    )
    put(451, ANSWER_BLANK)
    put(452, ANSWER_BLANK_LONG)
    put(
        453,
        subpart(
            "ii",
            "寫出一條 SQL，找出報名人數超過 30 的工作坊名稱（Title）及其報名人數。"
            "（需使用 COUNT、GROUP BY、HAVING）",
            6,
            depth=2,
        ),
    )
    put(454, ANSWER_BLANK)
    put(455, ANSWER_BLANK_LONG)
    put(
        456,
        subpart(
            "iii",
            "寫出一條 SQL，找出從未報名任何工作坊的會員姓名。（可用 NOT IN 或 NOT EXISTS）",
            6,
            depth=2,
        ),
    )

    put(458, topic_header("2.", "正規化與資料完整性", 10))
    put(
        460,
        "某社團登記表把「學號、姓名、社團名稱、社團會址、會費」全部放在同一工作表的一列中，"
        "且同一社團會址重複出現於多行。",
    )
    put(
        462,
        subpart(
            "a",
            "指出此設計違反哪一條正規化規則，並說明會造成什麼更新異常（update anomaly）。",
            5,
        ),
    )
    put(463, ANSWER_BLANK)
    put(464, ANSWER_BLANK_LONG)
    put(
        466,
        subpart("b", "建議如何拆分為至少兩個實體／資料表，並寫出各表的主要屬性。", 5),
    )
    put(467, ANSWER_BLANK)
    put(468, ANSWER_BLANK_LONG)

    put(482, topic_header("3.", "進階 SQL 與資料操作", 10))
    put(485, stem("沿用上述 Member、Workshop、Registration 資料表。"))
    put(
        487,
        subpart(
            "a",
            "寫出一條 SQL，列出曾報名兩個或以上不同工作坊的會員姓名"
            "（需 DISTINCT／COUNT／GROUP BY／HAVING）。",
            5,
        ),
    )
    put(488, ANSWER_BLANK)
    put(489, ANSWER_BLANK_LONG)
    put(
        491,
        subpart(
            "b",
            "寫出一條 SQL，把 WorkshopID = 'WS101' 且 PayStatus = '已付款' 的報名記錄 "
            "PayStatus 更新為「已確認」。（UPDATE … WHERE）",
            5,
        ),
    )
    put(492, ANSWER_BLANK)
    put(493, ANSWER_BLANK_LONG)
    return lines


def _apply_cover(doc: Document) -> None:
    """Patch 24_25 EN-layout cover for 2025-2026 S5 ICT."""
    cell = doc.tables[0].cell(0, 0)
    instr = [cell.paragraphs[i].text for i in range(17, min(22, len(cell.paragraphs)))]
    apply_cmp_cover_zh_on_en_layout(
        cell,
        ZhCoverPatch(
            year_term="2025 – 2026 下學期考試",
            level="中五級 資訊及通訊科技",
            paper="試題簿",
            total_line="總分：100",
        ),
        instructions=instr,
    )


def clear_answer_pages(doc: Document) -> None:
    note = "（本擬題稿不附標準答案；教師可自行增刪。）"
    for i in range(635, len(doc.paragraphs)):
        doc.paragraphs[i].text = ""
    if len(doc.paragraphs) > 635:
        doc.paragraphs[635].text = note


def generate(
    template: Path,
    output: Path,
    *,
    footer_meta: dict | None = None,
    rng: object | None = None,
) -> tuple[list[list[str]], str]:
    """Render DOCX; return (final MCQ rows, mcq answer key)."""
    shutil.copy(template, output)
    doc = Document(str(output))
    _apply_cover(doc)
    blocks = mcq_blocks(doc)
    payload = build_mcq_payload()
    block_map = {i + 1: row for i, row in enumerate(payload)}
    shuffled, mcq_key = build_random_mcq_key(
        block_map,
        correct_indices=F5_MCQ_CORRECT_INDEX,
        rng=rng,
    )
    apply_mcq(doc, blocks, [shuffled[i] for i in range(1, len(shuffled) + 1)])
    replace_span(doc, 313, 422, build_part_b())
    replace_span(doc, 423, 624, build_part_c())
    clear_answer_pages(doc)
    doc.save(str(output))
    if footer_meta:
        apply_footer_meta(output, footer_meta)
    final_rows = [shuffled[i] for i in range(1, len(shuffled) + 1)]
    return final_rows, mcq_key


def main(argv: list[str] | None = None) -> int:
    root = repo_root()
    default_out = root / "Subjects/S5-ICT/past-papers/2025-2026/Term 02"
    template_dir = root / "Subjects/S5-ICT/past-papers/2024-2025/Term 02"
    ap = argparse.ArgumentParser(
        description="F5 ICT blueprint: build exam spec → compare → optional DOCX render.",
    )
    ap.add_argument(
        "--template",
        type=Path,
        default=template_dir / "WrittenExam/24_25_S5_ICT_Exam02.docx",
    )
    ap.add_argument(
        "--output",
        type=Path,
        default=default_out / "WrittenExam/25_26_S5_ICT_Exam02.docx",
    )
    ap.add_argument(
        "--spec",
        type=Path,
        default=default_out / "_generation/25_26_S5_ICT_Exam02.spec.json",
        help="Exam spec JSON (primary artifact for compare)",
    )
    ap.add_argument("--subject", default="F5 ICT", help="Subject folder for past-paper scan")
    ap.add_argument("--years", type=int, default=3, help="Past academic years to compare")
    ap.add_argument(
        "--duplicate-report",
        type=Path,
        help="Write duplicate report JSON (regenerate_ids for batch fix)",
    )
    ap.add_argument("--skip-check", action="store_true", help="Skip spec similarity check")
    ap.add_argument(
        "--render",
        action="store_true",
        help="Render DOCX after check (default: only render when check passes)",
    )
    ap.add_argument(
        "--render-anyway",
        action="store_true",
        help="Render DOCX even if duplicates found",
    )
    ap.add_argument("--strict", action="store_true", help="Exit 2 on duplicates (legacy hard fail)")
    args = ap.parse_args(argv)

    template = args.template.expanduser().resolve()
    output = args.output.expanduser().resolve()
    spec_path = args.spec.expanduser().resolve()
    dup_report = args.duplicate_report or spec_path.with_suffix(".duplicates.json")

    output.parent.mkdir(parents=True, exist_ok=True)
    footer = {
        "academic_year": "2025-2026",
        "level": "中五級",
        "term_exam": "下學期考試",
        "subject": "資訊及通訊科技",
    }
    if args.skip_check:
        mcq_rows, mcq_key = generate(template, output, footer_meta=footer)
        spec = build_f5_ict_exam_spec(mcq_rows=mcq_rows, mcq_answers=mcq_key)
        save_spec(spec_path, spec)
        print(f"Wrote spec: {spec_path}")
        print(f"Wrote DOCX: {output}")
        return 0

    import random

    check_code = 1
    for attempt in range(1, 41):
        mcq_rows, mcq_key = generate(
            template, output, footer_meta=footer, rng=random.Random(2025_2026 + attempt)
        )
        spec = build_f5_ict_exam_spec(mcq_rows=mcq_rows, mcq_answers=mcq_key)
        save_spec(spec_path, spec)
        check_code = run_spec_check(
            candidate_spec=spec_path,
            template=template,
            candidate_docx=output,
            years=args.years,
            subject_subpath=args.subject,
            json_report=dup_report,
            strict=args.strict,
        )
        if check_code == 0:
            print(f"Wrote spec: {spec_path}")
            print(f"Wrote DOCX: {output}")
            print(f"MCQ key: {mcq_key} (attempt {attempt})")
            break
    else:
        print("Could not pass spec duplicate check after 40 attempts.")
        return check_code

    import importlib.util

    _qdir = Path(__file__).resolve().parents[1] / "question-quality-check"
    _pdir = Path(__file__).resolve().parents[1] / "paper-quality-check"
    if str(_qdir) not in sys.path:
        sys.path.insert(0, str(_qdir))
    if str(_pdir) not in sys.path:
        sys.path.append(str(_pdir))
    _qdocx = importlib.util.spec_from_file_location(
        "qqc_check_docx", _qdir / "check_docx.py"
    )
    assert _qdocx and _qdocx.loader
    _qmod = importlib.util.module_from_spec(_qdocx)
    _qdocx.loader.exec_module(_qmod)
    question_docx_main = _qmod.main
    from check_paper import format_paper_report_text, report_exit_code, run_paper_check

    docx_argv = [
        "--candidate",
        str(output),
        "--candidate-spec",
        str(spec_path),
        "--template",
        str(template),
        "--subject",
        args.subject,
        "--years",
        str(args.years),
    ]
    docx_code = int(question_docx_main(docx_argv))
    p_report = run_paper_check(
        output,
        candidate_spec_path=spec_path,
        template_docx_path=template,
    )
    print("\n--- Paper quality check (DOCX) ---")
    print(format_paper_report_text(p_report))
    final_code = max(check_code, docx_code, report_exit_code(p_report))
    if final_code == 0:
        print("All checks passed.")
    return final_code


if __name__ == "__main__":
    raise SystemExit(main())
