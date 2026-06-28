"""Helpers for marking Word (.docx) submissions."""

import re
from zipfile import ZipFile

from docx import Document

_IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".emf", ".wmf")


def _is_docx_image_path(name: str) -> bool:
    lower = name.lower()
    if not any(lower.endswith(ext) for ext in _IMAGE_EXTS):
        return False
    return lower.startswith("word/media/") or lower.startswith("media/")


def extract_docx_images(filepath):
    """Return list of (filename, bytes) for every image embedded in the docx."""
    images = []
    with ZipFile(filepath) as archive:
        for name in sorted(archive.namelist()):
            if _is_docx_image_path(name):
                images.append((name.split("/")[-1], archive.read(name)))
    return images


def read_docx_text(filepath):
    """Extract all paragraph and table cell text from a docx file."""
    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)
    return "\n".join(parts)


def count_docx_images(filepath):
    with ZipFile(filepath) as archive:
        return sum(1 for name in archive.namelist() if _is_docx_image_path(name))


def normalize_text(text):
    text = text.lower()
    text = text.replace("“", '"').replace("”", '"')
    text = text.replace("‘", "'").replace("’", "'")
    return re.sub(r"\s+", " ", text)
