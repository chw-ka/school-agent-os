"""
Reusable exam paper similarity comparison (DOCX / PDF).

Rule: similarity **greater than 60%** → treated as the same question (duplicate).
Comparisons are **cross-number**: MCQ#3 on the new paper can match MCQ#15 on a reference.
"""
from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass, field
from difflib import SequenceMatcher
from pathlib import Path
from typing import Iterable, Optional

try:
    from docx import Document
except ImportError as e:  # pragma: no cover
    raise ImportError("question-quality-check requires python-docx. Install: pip install python-docx") from e


# Similarity **>** this value ⇒ duplicate (「多於 60%」)
THRESH_DUPLICATE = 0.60
# 乙／丙 structured: whole question / scenario stem (整條)
THRESH_WRITTEN_STEM = 0.60
# 乙／丙 sub-questions (a)(b)(i)… — may mix parts from several DSE years
THRESH_WRITTEN_SUBPART = 0.85
# Cross-section pairs (e.g. 甲部 MCQ vs 丁部填充) — lower bar; near-paraphrase leaks answers
THRESH_INTRA_CROSS = 0.48

# Sub-part line: tab + (a) / (i) / (ii) …
_WRITTEN_SUBPART_START = re.compile(
    r"^\s*(?:\t+\s*)?\(([a-z]{1,2}|[ivx]{1,4})\)\s*[\t ]",
    re.IGNORECASE,
)

SECTION_BOUNDARIES: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("乙部", ("丙部",)),
    ("丙部", ("丁部",)),
    ("丁部", ("戊部",)),
    ("戊部", ("試卷完", "附表", "下期考試答案")),
)

STOP_MARKERS = ("試卷完", "附表", "下期考試答案", "本擬題稿", "2024-2025 年度", "－　試卷完")

EXCLUDE_NAME_PARTS = (
    "answer",
    "ans",
    "mock",
    "practical",
    "sample",
    "backup",
    "任務",
    "sp_dse",
    "full_sp",
)

INCLUDE_NAME_PARTS = ("exam", "試題", "writtenexam", "paper")

# Legacy aliases for CLI help text
THRESH_EXACT = THRESH_DUPLICATE
THRESH_SIMILAR = THRESH_DUPLICATE
THRESH_PARTIAL = THRESH_DUPLICATE


@dataclass(frozen=True)
class Match:
    reference: str
    match_type: str  # mcq_stem | mcq_full | written | line | section
    similarity: float
    candidate_label: str
    reference_label: str
    candidate_snippet: str
    reference_snippet: str

    def is_duplicate(self) -> bool:
        return self.similarity > THRESH_DUPLICATE

    def severity(self) -> str:
        return "duplicate" if self.is_duplicate() else "ok"


@dataclass
class CompareReport:
    candidate: str
    template: Optional[str]
    references_checked: list[str] = field(default_factory=list)
    matches: list[Match] = field(default_factory=list)
    threshold: float = THRESH_DUPLICATE

    @property
    def duplicate_count(self) -> int:
        return sum(1 for m in self.matches if m.is_duplicate())

    @property
    def exact_count(self) -> int:
        return self.duplicate_count

    @property
    def similar_count(self) -> int:
        return self.duplicate_count

    def to_dict(self) -> dict:
        d = asdict(self)
        d["matches"] = [asdict(m) for m in self.matches]
        d["summary"] = {
            "threshold": self.threshold,
            "duplicate": self.duplicate_count,
            "rule": f"similarity > {self.threshold:.0%} treated as same question",
        }
        return d

    def exit_code(self, *, fail_on_exact: bool = False, fail_on_similar: bool = False) -> int:
        """
        0 = no duplicates
        1 = duplicates found (full report; default, non-blocking)
        2 = strict abort (fail_on_exact or fail_on_similar)
        """
        if self.duplicate_count == 0:
            return 0
        if fail_on_exact or fail_on_similar:
            return 2
        return 1


def normalize_text(s: str) -> str:
    s = re.sub(r"\s+", "", s)
    s = re.sub(r"[（）()【】\[\]「」""''\'\t]", "", s)
    return s.lower()


def text_similarity(a: str, b: str) -> float:
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, normalize_text(a), normalize_text(b)).ratio()


def is_option_line(t: str) -> bool:
    s = t.lstrip()
    return len(s) >= 2 and s[0] in "ABCD" and s[1] == "."


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError("PDF support requires pypdf. Install: pip install pypdf") from e
    reader = PdfReader(str(path))
    return "\n".join((p.extract_text() or "") for p in reader.pages)


def extract_lines(path: Path) -> list[str]:
    path = path.expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(path)

    if path.suffix.lower() == ".docx":
        doc = Document(str(path))
        lines: list[str] = []
        for p in doc.paragraphs:
            t = p.text.rstrip()
            if t.strip():
                lines.append(t)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    t = cell.text.rstrip()
                    if t.strip():
                        lines.append(t)
        return _trim_exam_body(lines)

    if path.suffix.lower() == ".pdf":
        raw = _read_pdf(path)
        lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
        return _trim_exam_body(lines)

    raise ValueError(f"Unsupported format: {path.suffix}")


def _trim_exam_body(lines: list[str]) -> list[str]:
    out: list[str] = []
    for t in lines:
        if any(m in t for m in STOP_MARKERS):
            break
        out.append(t)
    return out


def extract_mcq_stems(lines: list[str]) -> list[dict]:
    """Each MCQ: {index, stem, full} — index is position in paper only (for display)."""
    stems: list[dict] = []
    i = 0
    qn = 0
    while i < len(lines):
        if "乙部" in lines[i]:
            break
        if "甲部" in lines[i] or "多項選擇" in lines[i]:
            i += 1
            continue
        if is_option_line(lines[i]) or len(lines[i]) < 6:
            i += 1
            continue
        parts: list[str] = []
        while i < len(lines) and not is_option_line(lines[i]) and "乙部" not in lines[i]:
            if lines[i] and "甲部" not in lines[i]:
                parts.append(lines[i])
            i += 1
        opts: list[str] = []
        while i < len(lines) and is_option_line(lines[i]):
            opts.append(lines[i])
            i += 1
        if parts and opts:
            qn += 1
            stem = "\n".join(parts)
            stems.append({"index": qn, "stem": stem, "full": "\n".join(parts + opts)})
    return stems


def _section_lines(lines: list[str], start_kw: str, end_kws: tuple[str, ...]) -> list[str]:
    buf: list[str] = []
    active = False
    for t in lines:
        if start_kw in t:
            active = True
            continue
        if active and any(k in t for k in end_kws):
            break
        if active:
            buf.append(t)
    return buf


def split_written_stem_and_subparts(text: str) -> tuple[str, list[str]]:
    """
    Split 乙／丙 slot text into scenario stem (整條) and sub-question blocks ((a)(b)(i)…).
    """
    lines = (text or "").splitlines()
    stem_lines: list[str] = []
    subparts: list[str] = []
    current: list[str] = []
    in_sub = False

    def _flush_sub() -> None:
        nonlocal current, in_sub
        block = "\n".join(current).strip()
        if len(normalize_text(block)) >= 12:
            subparts.append(block)
        current = []
        in_sub = False

    for line in lines:
        if _WRITTEN_SUBPART_START.match(line):
            if in_sub:
                _flush_sub()
            elif current:
                stem_lines.extend(current)
                current = []
            current = [line]
            in_sub = True
        else:
            current.append(line)

    if in_sub:
        _flush_sub()
    elif current:
        stem_lines.extend(current)

    stem = "\n".join(stem_lines).strip()
    return stem, subparts


def extract_written_units(lines: list[str]) -> list[dict]:
    """
    Extract 乙/丙 question units for cross-number comparison.
    Splits on numbered main questions (1. 2.) or keeps (a)(b) groups with context.
    """
    units: list[dict] = []
    uid = 0

    for sec_name, end_kws in SECTION_BOUNDARIES:
        sec_lines = _section_lines(lines, sec_name, end_kws)
        if not sec_lines:
            continue

        current: list[str] = []
        label = ""

        def flush() -> None:
            nonlocal uid, current, label
            text = "\n".join(current).strip()
            if len(normalize_text(text)) < 20:
                current = []
                return
            uid += 1
            units.append(
                {
                    "index": uid,
                    "section": sec_name,
                    "label": label or f"{sec_name}#{uid}",
                    "text": text,
                }
            )
            current = []

        for t in sec_lines:
            if _is_boilerplate_line(t) and "分)" not in t and "分）" not in t:
                continue
            if re.match(r"^\d+\.\s", t):
                flush()
                label = f"{sec_name}-{t.split(chr(9))[0].strip()[:20]}"
                current = [t]
            elif re.match(r"^\([a-z]\)", t, re.I) and len(t) > 25:
                if current and len(current) > 3:
                    flush()
                label = f"{sec_name}-{t[:24]}"
                current.append(t)
            else:
                if t:
                    current.append(t)

        flush()

    return units


def _is_boilerplate_line(t: str) -> bool:
    if len(t) < 12:
        return True
    if is_option_line(t):
        return True
    # Generic MCQ option patterns — not substantive unique content
    s = t.lstrip()
    if re.match(r"^[ABCD]\.\s*只有\s*\(", s):
        return True
    boiler = (
        "甲部",
        "多項選擇",
        "請選擇最合適",
        "HB 鉛筆",
        "乙部",
        "丙部",
        "丁部",
        "戊部",
        "配對題",
        "填充題",
        "是非題",
        "短答題",
        "核心單元",
        "選修單元",
        "試題答題簿",
        "考生須知",
        "評分準則",
        "選修 A",
        "選修 B",
        "請判斷下列",
        "請在空格",
        "將最合適的答案",
        "從字庫中選擇",
        "班別：",
        "姓名：",
        "________________",
    )
    if any(b in t for b in boiler) and len(t) < 50:
        return True
  # Cover / instruction blocks (shared across term papers from same template)
    if any(
        x in t
        for x in (
            "迦密聖道中學",
            "試題簿",
            "考生須知",
            "時限:",
            "頁數:",
            "總分:",
            "學期考試",
            "評分準則",
        )
    ):
        return True
    if re.match(r"^\d{4}\s*[–\-]\s*\d{4}", t.strip()):
        return True
    return False


def _section_label(sec_name: str) -> str:
    return sec_name[0] if sec_name else "?"


def _mcq_key_from_lines(
    lines: list[str],
    *,
    expected_count: int,
    letters: tuple[str, ...] = ("A", "B", "C", "D"),
) -> list[str]:
    """Parse MCQ answer-key line near the end of the exam body."""
    candidates: list[tuple[int, str]] = []
    key_re = re.compile(r"^[ABCD](?:\s*[ABCD])+\s*$", re.IGNORECASE)
    for line in lines:
        compact = re.sub(r"\s+", "", line)
        if len(compact) < 8 or not all(c.upper() in letters for c in compact):
            continue
        if key_re.match(line.strip()) or len(compact) >= 10:
            candidates.append((len(compact), line))
    if not candidates:
        return []
    candidates.sort(key=lambda x: -x[0])
    parsed: list[str] = []
    for ch in candidates[0][1].upper():
        if ch in letters:
            parsed.append(ch)
    if expected_count and len(parsed) > expected_count:
        parsed = parsed[:expected_count]
    return parsed


def _parse_mcq_option_text(full: str, letter: str) -> str:
    target = letter.strip().upper()
    for line in full.splitlines():
        s = line.lstrip()
        if len(s) >= 2 and s[0] == target and s[1] == ".":
            return s[2:].strip()
    return ""


_ANSWER_LEAK_STOP = frozenset(
    {
        "只有",
        "可以",
        "應該",
        "使用",
        "程式",
        "下列",
        "哪項",
        "較合",
        "合理",
        "是否",
        "需要",
        "不能",
        "完全",
        "通常",
        "主要",
        "為了",
        "因為",
        "例如",
        "課堂",
        "學生",
        "老師",
        "同學",
        "測驗",
        "功能",
        "步驟",
        "影像",
        "語音",
        "文字",
        "檔案",
        "模型",
        "播放",
        "轉成",
        "轉為",
        "改成",
        "刪除",
        "安裝",
        "設定",
        "調整",
        "測試",
        "修改",
        "完美",
        "一次",
    }
)


def _token_is_question_topic(token: str, text: str) -> bool:
    """Token names the skill being tested (not an leaked answer)."""
    if f"「{token}" in text or f"{token}（" in text:
        return True
    return False


def _is_section_header(text: str) -> bool:
    s = text.strip()
    return bool(re.match(r"^[甲乙丙丁戊]部[：:–\-]", s)) and len(s) < 48


def _answer_leak_tokens(answer_text: str) -> list[str]:
    tokens = re.findall(r"[\u4e00-\u9fff]{2,}", answer_text)
    out: list[str] = []
    for token in tokens:
        if token in _ANSWER_LEAK_STOP:
            continue
        if len(token) < 2:
            continue
        if token not in out:
            out.append(token)
    return out


def extract_comparable_units(lines: list[str]) -> list[dict]:
    """
    Comparable question units for intra-exam duplicate / answer-leak checks.
    Covers 甲部 MCQ stems plus 乙–戊 substantive prompts (T/F, fill, short answer).
    """
    units: list[dict] = []
    for m in extract_mcq_stems(lines):
        units.append(
            {
                "label": f"甲#{m['index']}",
                "text": m["stem"],
                "full": m["full"],
                "kind": "mcq",
            }
        )

    for sec_name, end_kws in SECTION_BOUNDARIES:
        sec_lines = _section_lines(lines, sec_name, end_kws)
        if not sec_lines:
            continue
        sec_tag = _section_label(sec_name)
        uid = 0
        for t in sec_lines:
            if _is_boilerplate_line(t):
                continue
            norm = normalize_text(t)
            if len(norm) < 12:
                continue

            kind = "written"
            if "________" in t or re.search(r"_{3,}", t):
                kind = "fill"
            elif sec_name == "丙部" and (t.endswith("。") or len(t) > 28):
                kind = "tf"
            elif sec_name == "戊部" and ("？" in t or "?" in t):
                kind = "sa"
            elif re.match(r"^\d+\.\s", t):
                kind = "written"
            elif sec_name == "乙部":
                continue
            else:
                continue

            uid += 1
            units.append(
                {
                    "label": f"{sec_tag}#{uid}",
                    "text": t.strip(),
                    "kind": kind,
                    "section": sec_name,
                }
            )
    return units


def compare_intra_exam_lines(
    lines: list[str],
    *,
    mcq_answers: Optional[list[str]] = None,
    threshold_same: float = THRESH_DUPLICATE,
    threshold_cross: float = THRESH_INTRA_CROSS,
) -> list[Match]:
    """Detect repeated / leaking content within one exam paper."""
    units = extract_comparable_units(lines)
    mcq_stems = extract_mcq_stems(lines)
    matches: list[Match] = []

    for i, a in enumerate(units):
        for b in units[i + 1 :]:
            if a.get("kind") == "fill" and b.get("kind") == "fill":
                continue
            cross = a.get("kind") != b.get("kind")
            thresh = threshold_cross if cross else threshold_same
            sim = text_similarity(a["text"], b["text"])
            if sim > thresh:
                matches.append(
                    Match(
                        reference="(within exam)",
                        match_type="intra_exam",
                        similarity=sim,
                        candidate_label=a["label"],
                        reference_label=b["label"],
                        candidate_snippet=a["text"].replace("\n", " ")[:160],
                        reference_snippet=b["text"].replace("\n", " ")[:160],
                    )
                )

    if mcq_answers and len(mcq_answers) == len(mcq_stems):
        mcq_by_label = {f"甲#{m['index']}": m for m in mcq_stems}
        for m, letter in zip(mcq_stems, mcq_answers, strict=True):
            opt = _parse_mcq_option_text(m["full"], letter)
            if not opt:
                continue
            for token in _answer_leak_tokens(opt):
                for u in units:
                    if u.get("kind") == "mcq":
                        continue
                    if token not in u["text"]:
                        continue
                    if _token_is_question_topic(token, u["text"]):
                        continue
                    matches.append(
                        Match(
                            reference="(within exam)",
                            match_type="answer_leak",
                            similarity=1.0,
                            candidate_label=f"甲#{m['index']}({letter})",
                            reference_label=u["label"],
                            candidate_snippet=f"MCQ answer «{token}» from: {opt[:80]}",
                            reference_snippet=u["text"].replace("\n", " ")[:160],
                        )
                    )
                    break

    best: dict[tuple, Match] = {}
    for m in matches:
        key = (m.match_type, m.candidate_label, m.reference_label)
        if key not in best or m.similarity > best[key].similarity:
            best[key] = m
    return sorted(best.values(), key=lambda x: (-x.similarity, x.candidate_label))


def compare_intra_exam(
    candidate_path: Path,
    *,
    mcq_answers: Optional[list[str]] = None,
) -> list[Match]:
    lines = extract_lines(candidate_path)
    stems = extract_mcq_stems(lines)
    if mcq_answers is None and stems:
        mcq_answers = _mcq_key_from_lines(lines, expected_count=len(stems))
    return compare_intra_exam_lines(lines, mcq_answers=mcq_answers or None)


def _add_match(matches: list[Match], m: Match) -> None:
    matches.append(m)


def compare_documents(
    candidate_path: Path,
    reference_path: Path,
    *,
    reference_label: Optional[str] = None,
    template_mode: bool = False,
) -> list[Match]:
    ref_label = reference_label or str(reference_path)
    cand_lines = extract_lines(candidate_path)
    ref_lines = extract_lines(reference_path)

    matches: list[Match] = []

    cand_mcq = extract_mcq_stems(cand_lines)
    ref_mcq = extract_mcq_stems(ref_lines)

    # MCQ: cross-number (all pairs, not same index only)
    for c in cand_mcq:
        for r in ref_mcq:
            for field_name, match_type in (("stem", "mcq_stem"), ("full", "mcq_full")):
                sim = text_similarity(c[field_name], r[field_name])
                if sim > THRESH_DUPLICATE:
                    _add_match(
                        matches,
                        Match(
                            reference=ref_label,
                            match_type=match_type,
                            similarity=sim,
                            candidate_label=f"MCQ#{c['index']}",
                            reference_label=f"MCQ#{r['index']}",
                            candidate_snippet=c[field_name].replace("\n", " ")[:160],
                            reference_snippet=r[field_name].replace("\n", " ")[:160],
                        ),
                    )

    # 乙/丙 written units: cross-number
    if not template_mode:
        cand_w = extract_written_units(cand_lines)
        ref_w = extract_written_units(ref_lines)
        for c in cand_w:
            for r in ref_w:
                sim = text_similarity(c["text"], r["text"])
                if sim > THRESH_DUPLICATE:
                    _add_match(
                        matches,
                        Match(
                            reference=ref_label,
                            match_type="written",
                            similarity=sim,
                            candidate_label=c["label"],
                            reference_label=r["label"],
                            candidate_snippet=c["text"].replace("\n", " ")[:160],
                            reference_snippet=r["text"].replace("\n", " ")[:160],
                        ),
                    )

    # Line-level: cross-position (different line numbers OK)
    for i, cl in enumerate(cand_lines):
        if _is_boilerplate_line(cl):
            continue
        for j, rl in enumerate(ref_lines):
            if _is_boilerplate_line(rl):
                continue
            sim = text_similarity(cl, rl)
            if sim > THRESH_DUPLICATE:
                _add_match(
                    matches,
                    Match(
                        reference=ref_label,
                        match_type="line",
                        similarity=sim,
                        candidate_label=f"line:{i}",
                        reference_label=f"line:{j}",
                        candidate_snippet=cl[:120],
                        reference_snippet=rl[:120],
                    ),
                )

    # Section-level (whole 乙–戊) — only for past papers, not template layout
    if not template_mode:
        for sec, end in SECTION_BOUNDARIES:
            cs = "\n".join(_section_lines(cand_lines, sec, end))
            rs = "\n".join(_section_lines(ref_lines, sec, end))
            if cs and rs:
                sim = text_similarity(cs, rs)
                if sim > THRESH_DUPLICATE:
                    _add_match(
                        matches,
                        Match(
                            reference=ref_label,
                            match_type="section",
                            similarity=sim,
                            candidate_label=sec,
                            reference_label=sec,
                            candidate_snippet=cs[:120],
                            reference_snippet=rs[:120],
                        ),
                    )

    # Dedupe: keep strongest match per (candidate_label, reference file, match_type)
    best: dict[tuple, Match] = {}
    for m in matches:
        key = (m.match_type, m.candidate_label, m.reference)
        if key not in best or m.similarity > best[key].similarity:
            best[key] = m
    return sorted(best.values(), key=lambda m: (-m.similarity, m.reference))


def parse_academic_year(folder_name: str) -> Optional[int]:
    m = re.match(r"^(\d{4})-(\d{4})$", folder_name)
    if not m:
        return None
    return int(m.group(1))


def discover_past_papers(
    root: Path,
    *,
    years: int = 3,
    subject_subpath: Optional[str] = None,
    include_candidate: Optional[Path] = None,
) -> list[Path]:
    root = root.expanduser().resolve()
    if not root.is_dir():
        return []

    year_dirs: list[tuple[int, Path]] = []
    for p in sorted(root.iterdir()):
        if not p.is_dir():
            continue
        y = parse_academic_year(p.name)
        if y is not None:
            year_dirs.append((y, p))

    if not year_dirs:
        past_paper_roots = sorted({p for p in root.rglob("past-papers") if p.is_dir()})
        if past_paper_roots:
            year_dirs = []
            for ppr in past_paper_roots:
                for p in sorted(ppr.iterdir()):
                    if not p.is_dir():
                        continue
                    y = parse_academic_year(p.name)
                    if y is not None:
                        year_dirs.append((y, p))

    if not year_dirs:
        candidates = _filter_exam_files(root.rglob("*"))
    else:
        year_dirs.sort(key=lambda x: -x[0])
        selected_years = {y for y, _ in year_dirs[:years]}
        files: list[Path] = []
        for y, ydir in year_dirs:
            if y not in selected_years:
                continue
            files.extend(_filter_exam_files(ydir.rglob("*")))
        candidates = files

    if subject_subpath:
        sub = subject_subpath.replace("\\", "/").strip("/")
        candidates = [p for p in candidates if sub.lower() in str(p).replace("\\", "/").lower()]

    by_stem: dict[str, Path] = {}
    for p in sorted(candidates, key=lambda x: (x.suffix != ".docx", str(x))):
        stem = p.stem.lower()
        if stem not in by_stem:
            by_stem[stem] = p

    out = sorted(by_stem.values(), key=lambda p: str(p))
    if include_candidate:
        out = [p for p in out if p.resolve() != include_candidate.resolve()]
    return out


def _filter_exam_files(paths: Iterable[Path]) -> list[Path]:
    result: list[Path] = []
    for p in paths:
        if not p.is_file():
            continue
        if p.suffix.lower() not in (".docx", ".pdf"):
            continue
        name = p.name.lower()
        if any(x in name for x in EXCLUDE_NAME_PARTS):
            continue
        if not any(x in name for x in INCLUDE_NAME_PARTS):
            continue
        result.append(p)
    return result


def infer_subject_subpath(path: Path) -> Optional[str]:
    parts = path.parts
    for part in parts:
        m = re.match(r"^S([2-6])-(CMP|ICT)$", part, re.I)
        if m:
            return f"S{m.group(1)}-{m.group(2).upper()}"
        if re.match(r"^F[2-6]\s", part) or re.match(r"^S[2-6]\s", part) or part in (
            "F5 ICT",
            "S5 ICT",
            "S2 CMP",
            "S3 CMP",
        ):
            return part
    return None


def run_full_check(
    candidate: Path,
    *,
    template: Optional[Path] = None,
    past_papers_root: Optional[Path] = None,
    years: int = 3,
    subject_subpath: Optional[str] = None,
    extra_references: Optional[list[Path]] = None,
) -> CompareReport:
    candidate = candidate.expanduser().resolve()
    if subject_subpath is None:
        subject_subpath = infer_subject_subpath(candidate)

    refs: list[Path] = []
    if template:
        refs.append(template.expanduser().resolve())
    if past_papers_root:
        refs.extend(
            discover_past_papers(
                past_papers_root,
                years=years,
                subject_subpath=subject_subpath,
                include_candidate=candidate,
            )
        )
    if extra_references:
        refs.extend([r.expanduser().resolve() for r in extra_references])

    seen: set[str] = set()
    unique_refs: list[Path] = []
    for r in refs:
        key = str(r.resolve())
        if key in seen or key == str(candidate):
            continue
        seen.add(key)
        if r.exists():
            unique_refs.append(r)

    report = CompareReport(
        candidate=str(candidate),
        template=str(template) if template else None,
        references_checked=[str(r) for r in unique_refs],
    )

    template_resolved = str(template.expanduser().resolve()) if template else None

    for ref in unique_refs:
        try:
            is_template = template_resolved is not None and str(ref.resolve()) == template_resolved
            report.matches.extend(
                compare_documents(
                    candidate,
                    ref,
                    reference_label=str(ref),
                    template_mode=is_template,
                )
            )
        except Exception as e:  # pragma: no cover
            report.matches.append(
                Match(
                    reference=str(ref),
                    match_type="error",
                    similarity=0.0,
                    candidate_label="",
                    reference_label="",
                    candidate_snippet=str(e),
                    reference_snippet="",
                )
            )

    try:
        report.matches.extend(compare_intra_exam(candidate))
    except Exception as e:  # pragma: no cover
        report.matches.append(
            Match(
                reference="(within exam)",
                match_type="error",
                similarity=0.0,
                candidate_label="",
                reference_label="",
                candidate_snippet=str(e),
                reference_snippet="",
            )
        )

    return report


def format_report_text(report: CompareReport, *, min_similarity: float = THRESH_DUPLICATE) -> str:
    pct = int(THRESH_DUPLICATE * 100)
    lines = [
        f"Candidate: {report.candidate}",
        f"Template: {report.template or '(none)'}",
        f"References checked: {len(report.references_checked)}",
        f"Rule: similarity > {pct}% → same question (duplicate)",
        f"Cross-section (within exam): similarity > {int(THRESH_INTRA_CROSS * 100)}% also flagged",
        f"Duplicates found: {report.duplicate_count}",
        "Note: question numbers may differ (cross-number matching).",
        "Includes intra-exam overlap (甲–戊) and MCQ answer leaks into other sections.",
        "",
    ]
    shown = [
        m
        for m in report.matches
        if m.is_duplicate() and m.similarity >= min_similarity and m.match_type != "error"
    ]
    shown.sort(key=lambda m: (-m.similarity, m.reference))

    if not shown:
        lines.append(f"No duplicates (all pairs ≤ {pct}% similarity).")
        return "\n".join(lines)

    lines.append("Duplicates:")
    for m in shown:
        lines.append(
            f"\n[DUPLICATE {m.similarity:.0%}] {m.match_type} "
            f"{m.candidate_label} ↔ {m.reference_label}"
        )
        lines.append(f"  Ref file: {Path(m.reference).name}")
        lines.append(f"  Cand: {m.candidate_snippet}")
        lines.append(f"  Ref:  {m.reference_snippet}")
    return "\n".join(lines)


def write_report_json(report: CompareReport, path: Path) -> None:
    path.write_text(json.dumps(report.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8")
