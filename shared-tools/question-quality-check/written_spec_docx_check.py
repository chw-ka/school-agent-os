"""Verify 乙／丙 spec slot text matches rendered DOCX paragraphs (spec-driven render)."""
from __future__ import annotations

import sys
from dataclasses import dataclass, field
from pathlib import Path

from quality_lib import extract_lines, normalize_text, text_similarity

_FMT = Path(__file__).resolve().parents[1] / "paper-formatter"
if str(_FMT) not in sys.path:
    sys.path.insert(0, str(_FMT))
from written_slot_ranges import WRITTEN_SLOT_PARAGRAPHS  # noqa: E402


@dataclass
class SlotConsistency:
    slot_id: str
    similarity: float
    ok: bool
    spec_preview: str = ""
    docx_preview: str = ""


@dataclass
class WrittenSpecDocxResult:
    docx_path: str
    spec_path: str
    threshold: float
    ok: bool
    slots: list[SlotConsistency] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "ok": self.ok,
            "docx": self.docx_path,
            "spec": self.spec_path,
            "threshold": self.threshold,
            "slots": [
                {
                    "slot": s.slot_id,
                    "similarity": round(s.similarity, 4),
                    "ok": s.ok,
                    "spec_preview": s.spec_preview,
                    "docx_preview": s.docx_preview,
                }
                for s in self.slots
            ],
        }


def docx_slot_text(docx_path: Path, start: int, end: int) -> str:
    from docx import Document

    doc = Document(str(docx_path))
    parts = []
    for i in range(start, end + 1):
        if i < len(doc.paragraphs):
            parts.append(doc.paragraphs[i].text)
    return "\n".join(parts)


def check_written_spec_docx(
    spec: dict,
    docx_path: Path,
    *,
    threshold: float = 0.92,
) -> WrittenSpecDocxResult:
    """Compare each written pick item in spec to the same paragraph span in DOCX."""
    from exam_spec import spec_items

    docx_path = docx_path.expanduser().resolve()
    by_id = {it.id: it for it in spec_items(spec) if it.id in WRITTEN_SLOT_PARAGRAPHS}
    rows: list[SlotConsistency] = []

    for slot_id, (start, end, _m) in WRITTEN_SLOT_PARAGRAPHS.items():
        item = by_id.get(slot_id)
        if item is None:
            continue
        spec_t = normalize_text(item.text or "")
        docx_t = normalize_text(docx_slot_text(docx_path, start, end))
        sim = text_similarity(spec_t, docx_t) if spec_t and docx_t else 0.0
        rows.append(
            SlotConsistency(
                slot_id=slot_id,
                similarity=sim,
                ok=sim >= threshold if spec_t else True,
                spec_preview=(item.text or "")[:100].replace("\n", " "),
                docx_preview=docx_slot_text(docx_path, start, end)[:100].replace("\n", " "),
            )
        )

    ok = all(s.ok for s in rows)
    return WrittenSpecDocxResult(
        docx_path=str(docx_path),
        spec_path=str(spec.get("meta", {}).get("title", "")),
        threshold=threshold,
        ok=ok,
        slots=rows,
    )


def format_written_spec_docx_report(result: WrittenSpecDocxResult) -> str:
    lines = [
        f"Written spec ↔ DOCX: {'OK' if result.ok else 'MISMATCH'}",
        f"Threshold: {result.threshold:.0%} normalized text similarity per slot",
        "",
    ]
    for s in sorted(result.slots, key=lambda x: x.similarity):
        flag = "OK" if s.ok else "FAIL"
        lines.append(f"  [{flag}] {s.slot_id}  {s.similarity:.0%}")
        if not s.ok:
            lines.append(f"    spec: {s.spec_preview}")
            lines.append(f"    docx: {s.docx_preview}")
    return "\n".join(lines)
