# read all docx in a attachments folder
import os
import pandas as pd
from docx import Document

# get the list of submissions in a attachments/xxx folder


def get_submissions():
    submissions = pd.DataFrame()
    submissions['class'] = ""
    submissions['classnumber'] = ""
    submissions['section_id'] = ""
    submissions['assignment_id'] = ""
    submissions['uid'] = ""
    submissions['enrollment_id'] = ""
    submissions['path'] = ""
    submissions['marks'] = 0
    submissions['comments'] = ""

    for root, dirs, files in os.walk("attachments/WS1: Introduction of automation"):
        idx = 0
        for file in files:
            if file.endswith(".docx"):
                idx += 1
                print("Reading", file)
                path = os.path.join(root, file)
                filename = file.split(".")[0]
                submissions.loc[idx, "class"] = filename[0:2]
                submissions.loc[idx, "classnumber"] = filename[2:4]
                submissions.loc[idx, "section_id"] = filename.split("-")[1]
                submissions.loc[idx, "assignment_id"] = filename.split("-")[2]
                submissions.loc[idx, "uid"] = filename.split("-")[3]
                submissions.loc[idx, "enrollment_id"] = filename.split("-")[3]
                submissions.loc[idx, "path"] = path
                submissions.loc[idx, "marks"] = 0
                submissions.loc[idx, "comments"] = ""
    return submissions


submissions = get_submissions()
for idx, row in submissions.iterrows():
    docx = Document(row["path"])
    t = 0
    print("=========================================")
    print(row["class"], row["classnumber"])
    print(docx.paragraphs[4].text, ";", docx.paragraphs[5].text, ";", docx.paragraphs[6].text)
    print(docx.tables[0].rows[2].cells[1].text)
    print(docx.tables[0].rows[3].cells[1].text)
    print(docx.tables[1].rows[2].cells[1].text)
    print(docx.tables[1].rows[3].cells[1].text)
    print("=========================================")

    # print each cell in each table in the docx file
    for table in docx.tables:
        r = 0
        for row in table.rows:
            i = 0
            for cell in row.cells:
                print(t, r, i, cell.text)
                i += 1
            r += 1
        t += 1

print(submissions)
