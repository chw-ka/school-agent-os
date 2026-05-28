from __future__ import annotations

from pathlib import Path


def _add_heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def main() -> None:
    try:
        from docx import Document
    except Exception as e:  # pragma: no cover
        raise SystemExit(
            "Missing dependency: python-docx. Install with: pip install python-docx"
        ) from e

    # This file lives at:
    # <repo>/Subjects/S3-CMP/past-papers/2025-2026/Term 02/_generation/<this_file>
    # parents[6] => <repo>
    repo_root = Path(__file__).resolve().parents[6]
    out_dir = repo_root / "Subjects" / "S3-CMP" / "past-papers" / "2025-2026" / "Term 02" / "PracticalAssessment"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "2526_S3_CMP_Term02_PracticalAssessment.docx"
    alt_out_path = out_dir / "2526_S3_CMP_Term02_PracticalAssessment_UPDATED.docx"

    doc = Document()

    doc.add_heading("中三級 電腦科（2025-2026）第二學期 — 實習評估", level=0)
    p = doc.add_paragraph()
    p.add_run("姓名：").bold = True
    p.add_run("____________________________   ")
    p.add_run("班別：").bold = True
    p.add_run("__________   ")
    p.add_run("學號：").bold = True
    p.add_run("_____")

    doc.add_paragraph("")
    doc.add_paragraph("題目總數：3 題")
    doc.add_paragraph("AI 工具規則：只限任務三可使用 Gemini。任務一及任務二必須在不使用任何 AI 工具下完成。")

    _add_heading(doc, "兩個時段安排（監考／教師用）", level=1)
    doc.add_paragraph("時段 A：只做任務一至任務二；鎖住任務三。")
    doc.add_paragraph("時段 B：只做任務三（可用 Gemini）；鎖住任務二。")

    _add_heading(doc, "已提供檔案", level=1)
    doc.add_paragraph("任務一：task1_2526.py")
    doc.add_paragraph("任務二：task2_2526.py")
    doc.add_paragraph("任務三：task3_2526.py、vibe_spec.md、canteen_sales.json、canteen_sales_sample.json、gemini_prompts.txt")

    _add_heading(doc, "提交", level=1)
    doc.add_paragraph("請提交以下檔案：")
    doc.add_paragraph("- task1_2526.py", style="List Bullet")
    doc.add_paragraph("- task2_2526.py", style="List Bullet")
    doc.add_paragraph("- task3_2526.py", style="List Bullet")
    doc.add_paragraph("- sales_report.txt", style="List Bullet")
    doc.add_paragraph("- gemini_prompts.txt（至少 2 則完整提示）", style="List Bullet")

    _add_heading(doc, "任務一 — 智能播音員（禁止使用任何 AI 工具）", level=1)
    doc.add_paragraph("請完成 task1_2526.py 的填空。要求：")
    doc.add_paragraph('只詢問一次語言："en" 或 "zh"（預設："en"）。', style="List Bullet")
    doc.add_paragraph('重覆輸入訊息；若輸入 "Q" 則退出。', style="List Bullet")
    doc.add_paragraph("每次播音：產生 mp3 → 播放 → 刪除 mp3 檔案。", style="List Bullet")
    doc.add_paragraph("提示：本試卷不需要使用 def（未學亦可完成）。", style="List Bullet")

    _add_heading(doc, "任務二 — 物件追蹤（禁止使用任何 AI 工具）", level=1)
    doc.add_paragraph("請完成 task2_2526.py 的填空。要求：")
    doc.add_paragraph("開啟影片檔（例如：tracking_video.mp4）。", style="List Bullet")
    doc.add_paragraph("在第一格畫面選取要追蹤的物件（ROI）。", style="List Bullet")
    doc.add_paragraph("使用 MIL tracker 追蹤並以紅色方框顯示。", style="List Bullet")
    doc.add_paragraph("左上角顯示 tracker 類型；按 ESC 退出。", style="List Bullet")
    doc.add_paragraph("提示：本試卷不需要使用 def（未學亦可完成）。", style="List Bullet")

    _add_heading(doc, "任務三 — Vibe Coding：《小食部銷售報告》（只限 Gemini）", level=1)
    doc.add_paragraph("請細讀 vibe_spec.md，並使用 Gemini 協助完成 task3_2526.py。")
    doc.add_paragraph("重點要求：")
    doc.add_paragraph("讀取 canteen_sales.json，生成 sales_report.txt（總收入、最賺錢食品、各班別收入）。", style="List Bullet")
    doc.add_paragraph("把至少 2 則完整提示貼到 gemini_prompts.txt。", style="List Bullet")
    doc.add_paragraph("提示：你可以要求 Gemini 用「不用 def」的寫法。", style="List Bullet")
    doc.add_paragraph("可參考 sales_report_sample.txt 了解輸出格式。", style="List Bullet")

    _add_heading(doc, "常用例子（供任務一及任務二抄寫參考）", level=1)
    doc.add_paragraph("以下表格提供常用函數庫的安裝、匯入及示例語句。你可以按需要抄寫到自己的程式中。")

    # Table A (like PracticalMock): libraries + pip + import
    table_a = doc.add_table(rows=1, cols=3)
    table_a.style = "Table Grid"
    hdr = table_a.rows[0].cells
    hdr[0].text = "函數庫"
    hdr[1].text = "pip 安裝指令"
    hdr[2].text = "匯入（import）例子"

    rows_a = [
        ("gTTS", "gTTS", "from gtts import gTTS"),
        ("playsound", "playsound==1.2.2", "from playsound import playsound"),
        ("os", "（已內建，不用安裝）", "import os"),
        ("OpenCV", "opencv-contrib-python", "import cv2"),
    ]
    for lib, pip_cmd, import_ex in rows_a:
        r = table_a.add_row().cells
        r[0].text = lib
        r[1].text = pip_cmd
        r[2].text = import_ex

    doc.add_paragraph("")

    # Table B (like PracticalMock): library + function + example
    table_b = doc.add_table(rows=1, cols=3)
    table_b.style = "Table Grid"
    hdr = table_b.rows[0].cells
    hdr[0].text = "函數庫"
    hdr[1].text = "功能"
    hdr[2].text = "示例語句"

    rows_b = [
        ("gTTS", "建立 gTTS 物件", 'audio = gTTS(text="I Love You", lang="en")'),
        ("gTTS", "儲存語音檔案", 'audio.save("audio.mp3")'),
        ("playsound", "播放聲音檔", 'playsound("audio.mp3")'),
        ("os", "刪除暫存聲音檔", 'os.remove("output.mp3")'),
        ("OpenCV", "讀取圖片", 'image = cv2.imread("face.jpg")'),
        ("OpenCV", "開啟影片（或 webcam）", 'cap = cv2.VideoCapture("video.mp4")  /  cv2.VideoCapture(0)'),
        ("OpenCV", "讀取每一格畫面", "isTrue, frame = cap.read()"),
        ("OpenCV", "轉灰階", "cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)"),
        ("OpenCV", "畫方框", "cv2.rectangle(frame, (x, y), (x + w, y + h), (255, 0, 0), 2)"),
        ("OpenCV", "顯示畫面", "cv2.imshow('Video', frame)"),
        ("OpenCV", "等待按鍵", "cv2.waitKey(1)"),
    ]
    for lib, fn, ex in rows_b:
        r = table_b.add_row().cells
        r[0].text = lib
        r[1].text = fn
        r[2].text = ex

    try:
        doc.save(out_path)
        print(f"Wrote {out_path}")
    except PermissionError:
        # Likely opened in Word; write an alternate file instead.
        doc.save(alt_out_path)
        print(f"Wrote {alt_out_path} (original file is locked)")


if __name__ == "__main__":
    main()

